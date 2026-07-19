import re
from io import BytesIO
from pathlib import Path
from datetime import datetime, timedelta

import pandas as pd
import pdfplumber
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from PIL import Image


st.set_page_config(page_title="Geotechnical Draw Intelligence", layout="wide")

DAILY_FILE = Path("gdi_history.csv")
WEEKEND_FILE = Path("gdi_weekend_cumulative.csv")
WEEKEND_DERIVED_FILE = Path("gdi_weekend_derived_daily.csv")
WEEKLY_FILE = Path("gdi_weekly_history.csv")
DRAWPOINT_MAPPING_FILE = Path("gdi_drawpoint_polygon_mapping.csv")
POLYGON_HISTORY_FILE = Path("gdi_polygon_history.csv")
DEFAULT_MAP_FILE = Path("lift2east_polygon_map_clean_v2.png")

DEFAULT_EP_THRESHOLD = 10000.0
DEFAULT_ET_THRESHOLD = 10.0


def find_value(pattern, text):
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None


def pdf_to_text(uploaded_pdf):
    text = ""
    with pdfplumber.open(uploaded_pdf) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def load_csv(path):
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


def save_csv(df, path):
    df.to_csv(path, index=False)


def save_or_update(path, record, key_cols):
    df = load_csv(path)
    record_df = pd.DataFrame([record])

    if df.empty:
        df = record_df
    else:
        df = pd.concat([df, record_df], ignore_index=True)
        df = df.drop_duplicates(subset=key_cols, keep="last")

    save_csv(df, path)




def weekend_actual_date(weekend_start, derived_day):
    """Return the real calendar date for Friday/Saturday/Sunday/Monday derived data."""
    offset_map = {
        "Friday": 0,
        "Saturday": 1,
        "Sunday": 2,
        "Monday": 3,
    }

    try:
        base_date = datetime.strptime(str(weekend_start), "%Y-%m-%d")
    except ValueError:
        return str(weekend_start)

    return (base_date + timedelta(days=offset_map.get(derived_day, 0))).strftime("%Y-%m-%d")


def safe_sort_history(df):
    if df.empty:
        return df
    df = df.copy()
    sort_candidates = ["Report Date", "Weekend Start", "Week Range"]
    for col in sort_candidates:
        if col in df.columns:
            return df.sort_values(col, kind="stable")
    return df


def to_float(value, default=0.0):
    try:
        if value is None or str(value).strip() == "":
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def to_int(value, default=0):
    try:
        if value is None or str(value).strip() == "":
            return default
        return int(float(value))
    except (TypeError, ValueError):
        return default


def row_value(row, column, default=""):
    if column not in row.index:
        return default
    value = row[column]
    if pd.isna(value):
        return default
    return value


def recalc_draw_record(record, ep_threshold, et_threshold):
    energy = to_float(record.get("Energy J", 0.0))
    potency = to_float(record.get("Potency m3", 0.0))
    tonnes = to_float(record.get("Tonnes", 0.0))

    ep = energy / potency if potency > 0 else 0.0
    et = energy / tonnes if tonnes > 0 else None
    status = get_status(ep, et, ep_threshold, et_threshold)
    health = cave_health_score(status)
    action = recommended_action(status)

    record["E/P"] = ep
    record["E/T"] = et if et is not None else ""
    events = to_int(record.get("New Events", record.get("Events", 0)))
    score = draw_compliance_score(ep, et, events, potency, ep_threshold, et_threshold)

    record["Status"] = status
    record["Cave Health Score"] = health
    record["Draw Compliance Score"] = score
    record["Score Label"] = score_label(score)
    record["Recommended Action"] = action
    return record



def draw_compliance_score(ep, et, events=0, potency=0.0, ep_threshold=DEFAULT_EP_THRESHOLD, et_threshold=DEFAULT_ET_THRESHOLD):
    """Return a management-friendly 0-100 draw compliance score."""
    if et is None:
        return 0

    ep_ratio = ep / ep_threshold if ep_threshold else 0
    et_ratio = et / et_threshold if et_threshold else 0

    ep_component = max(0, 40 - min(40, ep_ratio * 25))
    et_component = max(0, 40 - min(40, et_ratio * 25))

    event_component = 10
    if events >= 100:
        event_component = 2
    elif events >= 50:
        event_component = 5
    elif events >= 20:
        event_component = 8

    potency_component = 10
    if potency <= 0:
        potency_component = 0

    return int(round(max(0, min(100, ep_component + et_component + event_component + potency_component))))


def score_label(score):
    if score >= 80:
        return "Good Draw"
    if score >= 60:
        return "Watch"
    if score >= 40:
        return "Concern"
    return "Intervention Required"


def draw_advisor(score, status):
    if score >= 80:
        return "Draw response is favourable. Maintain current draw strategy and continue routine monitoring."
    if score >= 60:
        return "Draw response is acceptable but should be watched. Check whether the trend is improving or deteriorating before changing draw."
    if score >= 40:
        return "Draw response requires attention. Review tonnes distribution, adjacent drawpoints and recent seismic clustering."
    if status == "Waiting for tonnes":
        return "Enter tonnes drawn to finalise the draw compliance score and recommendation."
    return "Escalate for geotechnical review. Treat this as a potential draw-control concern until trends improve."


def management_summary(poly, events, energy, potency, tonnes, ep, et, status, score):
    et_text = f"{et:,.2f}" if et is not None else "waiting for tonnes"
    return (
        f"{poly} recorded {events} seismic events with total energy of {energy:,.0f} J and potency of {potency:.3f} m³. "
        f"Tonnes drawn were {tonnes:,.0f}. The calculated E/P is {ep:,.2f} and E/T is {et_text}. "
        f"The area is classified as {status}, with a Draw Compliance Score of {score}/100 ({score_label(score)}). "
        f"{draw_advisor(score, status)}"
    )


def add_score_columns(df, ep_threshold, et_threshold):
    if df.empty:
        return df
    df = df.copy()
    for col in ["E/P", "E/T", "New Events", "Events", "Potency m3"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    scores = []
    labels = []
    for _, row in df.iterrows():
        ep = to_float(row.get("E/P", 0.0))
        et_value = row.get("E/T", None)
        et = None if pd.isna(et_value) or str(et_value).strip() == "" else to_float(et_value)
        events = to_int(row.get("New Events", row.get("Events", 0)))
        potency = to_float(row.get("Potency m3", 0.0))
        score = draw_compliance_score(ep, et, events, potency, ep_threshold, et_threshold)
        scores.append(score)
        labels.append(score_label(score))
    df["Draw Compliance Score"] = scores
    df["Score Label"] = labels
    return df


def simple_forecast(df, column):
    if df.empty or column not in df.columns:
        return None
    values = pd.to_numeric(df[column], errors="coerce").dropna().tail(7)
    if values.empty:
        return None
    return float(values.mean())

def clean_number(value):
    if value is None:
        return None
    value = str(value).replace(",", "").strip()
    try:
        return float(value)
    except ValueError:
        return None


def clean_report_date(last_event):
    if not last_event:
        return datetime.now().strftime("%Y-%m-%d")

    match = re.search(r"(\d{1,2})\s+([A-Za-z]{3})", last_event)
    if not match:
        return datetime.now().strftime("%Y-%m-%d")

    day = int(match.group(1))
    month_text = match.group(2).title()

    month_map = {
        "Jan": 1, "Feb": 2, "Mar": 3, "Apr": 4,
        "May": 5, "Jun": 6, "Jul": 7, "Aug": 8,
        "Sep": 9, "Oct": 10, "Nov": 11, "Dec": 12
    }

    month = month_map.get(month_text, datetime.now().month)
    year = datetime.now().year
    return f"{year}-{month:02d}-{day:02d}"


def parse_poly_sections(text):
    poly_matches = list(re.finditer(r"Poly:\s*([A-Za-z0-9_]+)", text))
    results = {}

    for i, match in enumerate(poly_matches):
        poly = match.group(1)
        start = match.start()
        end = poly_matches[i + 1].start() if i + 1 < len(poly_matches) else len(text)
        block = text[start:end]

        last_event = find_value(r"Last event Date and Time:\s*(.+)", block)
        events = find_value(r"Number of new events:\s*(\d+)", block)
        energy = find_value(r"Total Energy:\s*([0-9.eE+-]+)", block)
        potency = find_value(r"Total Potency:\s*([0-9.eE+-]+)", block)
        current_rate = find_value(r"Current normalized activity rate:\s*([0-9.]+)", block)
        medium_rate = find_value(r"Medium term normalized activity rate:\s*([0-9.]+)", block)

        comment = ""
        comment_match = re.search(
            r"Medium term normalized activity rate:\s*[0-9.]+\s*(.*?)(?:\n\d+\s+of\s+\d+|\Z)",
            block,
            re.DOTALL
        )
        if comment_match:
            comment = " ".join(comment_match.group(1).split())

        if events or energy or potency:
            results[poly] = {
                "Poly": poly,
                "Last Event": last_event,
                "Events": int(events) if events else 0,
                "Energy J": clean_number(energy) or 0.0,
                "Potency m3": clean_number(potency) or 0.0,
                "Current Activity Rate": clean_number(current_rate),
                "Medium Activity Rate": clean_number(medium_rate),
                "Comment": comment,
            }

    return results


def parse_week_range(text):
    match = re.search(r"Seismic Report:\s*(.+)", text)
    if match:
        return match.group(1).strip()
    return ""


def parse_mine_total(text):
    total = find_value(r"Total number of events\s*=\s*(\d+)", text)
    return int(total) if total else 0


def parse_largest_magnitude(text):
    mags = re.findall(r"\b(?:ml|ML|M)(-?\d+(?:\.\d+)?)", text)
    values = [float(m) for m in mags if m is not None]
    return max(values) if values else None


def get_status(ep, et, ep_threshold, et_threshold):
    if et is None:
        return "Waiting for tonnes"
    if ep < ep_threshold and et < et_threshold:
        return "Efficient Draw"
    if ep >= ep_threshold and et < et_threshold:
        return "Brittle but Efficient"
    if ep < ep_threshold and et >= et_threshold:
        return "Ductile but Costly"
    return "High-Risk Draw"


def cave_health_score(status):
    return {
        "Efficient Draw": 90,
        "Brittle but Efficient": 65,
        "Ductile but Costly": 55,
        "High-Risk Draw": 30,
        "Waiting for tonnes": 0,
    }.get(status, 0)


def recommended_action(status):
    return {
        "Efficient Draw": "Maintain current draw strategy and continue routine monitoring.",
        "Brittle but Efficient": "Continue draw, but monitor seismic response closely due to elevated E/P.",
        "Ductile but Costly": "Review draw distribution and tonnes efficiency because seismic cost per tonne is elevated.",
        "High-Risk Draw": "Escalate for geotechnical review. Monitor draw, seismicity and nearby workplaces closely.",
        "Waiting for tonnes": "Enter tonnes drawn to finalise the draw compliance classification.",
    }.get(status, "")


def plot_matrix(ep, et, history, ep_threshold, et_threshold):
    max_x = max(et_threshold * 2, et * 1.4)
    max_y = max(ep_threshold * 2, ep * 1.4)

    fig = go.Figure()

    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=0, y1=ep_threshold,
                  fillcolor="green", opacity=0.12, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=0, y1=ep_threshold,
                  fillcolor="orange", opacity=0.14, line_width=0)
    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=ep_threshold, y1=max_y,
                  fillcolor="gold", opacity=0.16, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=ep_threshold, y1=max_y,
                  fillcolor="red", opacity=0.12, line_width=0)

    fig.add_shape(type="line", x0=et_threshold, x1=et_threshold, y0=0, y1=max_y,
                  line=dict(dash="dash", width=2))
    fig.add_shape(type="line", x0=0, x1=max_x, y0=ep_threshold, y1=ep_threshold,
                  line=dict(dash="dash", width=2))

    if not history.empty and "E/P" in history.columns and "E/T" in history.columns:
        hist = history.copy()
        hist["E/P"] = pd.to_numeric(hist["E/P"], errors="coerce")
        hist["E/T"] = pd.to_numeric(hist["E/T"], errors="coerce")
        hist = hist.dropna(subset=["E/P", "E/T"])

        if not hist.empty:
            fig.add_trace(go.Scatter(
                x=hist["E/T"],
                y=hist["E/P"],
                mode="markers",
                name="Saved Records",
                text=hist["Report Date"] if "Report Date" in hist.columns else None,
                hovertemplate="Date: %{text}<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
            ))

    fig.add_trace(go.Scatter(
        x=[et],
        y=[ep],
        mode="markers+text",
        name="Current",
        text=["Current"],
        textposition="top center",
        marker=dict(size=22),
        hovertemplate="Current<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
    ))

    fig.add_annotation(x=et_threshold * 0.45, y=ep_threshold * 0.5, text="Efficient Draw", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 0.5, text="Ductile but Costly", showarrow=False)
    fig.add_annotation(x=et_threshold * 0.5, y=ep_threshold * 1.45, text="Brittle but Efficient", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 1.45, text="High-Risk Draw", showarrow=False)

    fig.update_layout(
        title="Draw Compliance Matrix",
        xaxis_title="Energy / Tonnes (E/T)",
        yaxis_title="Energy / Potency (E/P)",
        xaxis=dict(range=[0, max_x]),
        yaxis=dict(range=[0, max_y]),
        height=650,
    )

    return fig



def status_colour(status):
    return {
        "Efficient Draw": "green",
        "Brittle but Efficient": "orange",
        "Ductile but Costly": "orange",
        "High-Risk Draw": "red",
        "Waiting for tonnes": "lightgrey",
        "Waiting for seismic data": "lightgrey",
    }.get(status, "lightgrey")


def normalise_drawpoint(value):
    return re.sub(r"\s+", "", str(value).strip().upper())


def infer_drawbell(drawpoint):
    """Return the physical drawbell for a Lift II drawpoint.

    Examples:
        T2E01  + T1W01  -> A01
        T3E01  + T2W01  -> B01
        T4E01  + T3W01  -> C01

    For an east drawpoint TnE##, the bell letter is based on n-2.
    For a west drawpoint TnW##, the bell letter is based on n-1.
    """
    value = normalise_drawpoint(drawpoint)
    match = re.fullmatch(r"T(\d+)([EW])(\d{2})", value)
    if not match:
        return ""

    tunnel = int(match.group(1))
    side = match.group(2)
    bell_number = match.group(3)

    letter_index = tunnel - 2 if side == "E" else tunnel - 1
    if not 0 <= letter_index < 26:
        return ""

    return f"{chr(ord('A') + letter_index)}{bell_number}"


def read_table(uploaded_file):
    if uploaded_file is None:
        return pd.DataFrame()
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file)
    return pd.read_excel(uploaded_file)


def build_mapping_template(drawpoints):
    unique = sorted({normalise_drawpoint(v) for v in drawpoints if str(v).strip()})
    return pd.DataFrame({
        "Drawpoint": unique,
        "Drawbell": [infer_drawbell(v) for v in unique],
        "Polygon": [""] * len(unique),
    })


def build_drawbell_mapping(template, saved_mapping):
    """Create one editable polygon assignment per drawbell and propagate it to both drawpoints."""
    bell_map = template[["Drawbell"]].drop_duplicates().copy()
    bell_map = bell_map[bell_map["Drawbell"].astype(str).str.strip() != ""]
    bell_map["Polygon"] = ""

    if not saved_mapping.empty:
        saved = saved_mapping.copy()
        if "Drawbell" not in saved.columns and "Drawpoint" in saved.columns:
            saved["Drawbell"] = saved["Drawpoint"].map(infer_drawbell)
        saved["Drawbell"] = saved["Drawbell"].astype(str).str.strip().str.upper()
        saved["Polygon"] = saved.get("Polygon", "").astype(str).str.strip().str.upper()
        saved = saved[saved["Polygon"].isin(list("ABCDEF"))]
        saved = saved.drop_duplicates(subset=["Drawbell"], keep="last")[["Drawbell", "Polygon"]]
        bell_map = bell_map.drop(columns=["Polygon"]).merge(saved, on="Drawbell", how="left")
        bell_map["Polygon"] = bell_map["Polygon"].fillna("")

    return bell_map.sort_values("Drawbell", kind="stable").reset_index(drop=True)


def expand_drawbell_mapping(template, bell_mapping):
    """Attach the drawbell polygon assignment to every drawpoint in that drawbell."""
    clean = bell_mapping[["Drawbell", "Polygon"]].copy()
    clean["Drawbell"] = clean["Drawbell"].astype(str).str.strip().str.upper()
    clean["Polygon"] = clean["Polygon"].astype(str).str.strip().str.upper()
    return template.drop(columns=["Polygon"]).merge(clean, on="Drawbell", how="left")


def aggregate_polygon_tonnes(tonnes_df, drawpoint_col, tonnes_col, mapping_df):
    work = tonnes_df[[drawpoint_col, tonnes_col]].copy()
    work.columns = ["Drawpoint", "Tonnes"]
    work["Drawpoint"] = work["Drawpoint"].map(normalise_drawpoint)
    work["Tonnes"] = pd.to_numeric(work["Tonnes"], errors="coerce").fillna(0.0)

    mapping = mapping_df.copy()
    mapping["Drawpoint"] = mapping["Drawpoint"].map(normalise_drawpoint)
    mapping["Polygon"] = mapping["Polygon"].astype(str).str.strip().str.upper()
    mapping["Drawbell"] = mapping["Drawbell"].astype(str).str.strip().str.upper()

    merged = work.merge(mapping, on="Drawpoint", how="left")
    drawbell_totals = (
        merged.dropna(subset=["Polygon"])
        .groupby(["Polygon", "Drawbell"], as_index=False)["Tonnes"].sum()
    )
    polygon_totals = (
        merged.dropna(subset=["Polygon"])
        .groupby("Polygon", as_index=False)["Tonnes"].sum()
        .rename(columns={"Tonnes": "Polygon Tonnes"})
    )
    return merged, drawbell_totals, polygon_totals


def polygon_health_score(ep, et, events, potency, tonnes, target_tonnes, ep_threshold, et_threshold):
    """Return a transparent 0-100 polygon health score."""
    if potency <= 0 or tonnes <= 0 or et is None:
        return 0

    ep_ratio = ep / ep_threshold if ep_threshold else 0
    et_ratio = et / et_threshold if et_threshold else 0
    matrix_score = max(0, 60 - min(60, 25 * ep_ratio + 25 * et_ratio))

    production_score = 20
    if target_tonnes > 0:
        production_score = min(20, max(0, 20 * tonnes / target_tonnes))

    event_score = 15
    if events >= 100:
        event_score = 3
    elif events >= 60:
        event_score = 7
    elif events >= 30:
        event_score = 11

    completeness_score = 5
    return int(round(max(0, min(100, matrix_score + production_score + event_score + completeness_score))))


def polygon_health_rating(score):
    if score <= 0:
        return "No data"
    if score >= 80:
        return "Healthy"
    if score >= 60:
        return "Watch"
    if score >= 40:
        return "Concern"
    return "Critical"


def polygon_map_figure(results_df, map_image, selected_polygon="A"):
    """Status map using polygon paths aligned to the actual red boundaries."""
    image = Image.open(map_image) if not isinstance(map_image, Image.Image) else map_image
    width, height = image.size

    fig = go.Figure()
    fig.add_layout_image(dict(
        source=image, x=0, y=0, sizex=width, sizey=height,
        xref="x", yref="y", sizing="stretch", layer="below"
    ))

    # Pixel coordinates traced from the actual Lift II East polygon boundaries.
    polygon_points = {
        "D": [(642, 457), (784, 457), (784, 603), (642, 603)],
        "A": [(784, 460), (932, 460), (932, 603), (784, 603)],
        "E": [(642, 603), (784, 603), (784, 744), (642, 744)],
        "B": [(784, 603), (932, 603), (932, 744), (784, 744)],
        "F": [(642, 744), (784, 744), (784, 891), (716, 876), (642, 855)],
        "C": [(784, 744), (932, 744), (932, 891), (784, 891)],
    }

    rows = {
        str(r["Polygon"]).upper(): r
        for _, r in results_df.iterrows()
    } if not results_df.empty else {}

    for poly, points in polygon_points.items():
        row = rows.get(poly, {})
        status = row.get("Status", "Waiting for seismic data") if hasattr(row, "get") else "Waiting for seismic data"
        colour = status_colour(status)
        selected = poly == selected_polygon

        path = "M " + " L ".join(f"{x},{y}" for x, y in points) + " Z"
        fig.add_shape(
            type="path",
            path=path,
            fillcolor=colour,
            opacity=0.20 if not selected else 0.34,
            line=dict(color=colour, width=6 if selected else 3),
            layer="above",
        )

        centre_x = sum(x for x, _ in points) / len(points)
        centre_y = sum(y for _, y in points) / len(points)
        hover = f"<b>Polygon {poly}</b><br>Status: {status}"
        if hasattr(row, "get"):
            hover += (
                f"<br>Tonnes: {to_float(row.get('Polygon Tonnes', 0)):,.0f} t"
                f"<br>Energy: {to_float(row.get('Energy J', 0)):,.0f} J"
                f"<br>Potency: {to_float(row.get('Potency m3', 0)):,.3f} m³"
                f"<br>E/P: {to_float(row.get('E/P', 0)):,.2f}"
            )
            etv = row.get("E/T", None)
            hover += f"<br>E/T: {to_float(etv):,.2f}" if etv not in [None, ""] and not pd.isna(etv) else "<br>E/T: Waiting"

        fig.add_trace(go.Scatter(
            x=[centre_x], y=[centre_y], mode="markers",
            marker=dict(size=82, opacity=0),
            hovertext=[hover], hoverinfo="text", showlegend=False,
        ))

    fig.update_xaxes(range=[0, width], visible=False, fixedrange=True)
    fig.update_yaxes(range=[height, 0], visible=False, scaleanchor="x", fixedrange=True)
    fig.update_layout(
        height=700,
        margin=dict(l=0, r=0, t=0, b=0),
        title=None,
        showlegend=False,
        dragmode=False,
        hovermode="closest",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
    )
    return fig

def pie_chart(df, title):
    fig = px.pie(df, names="Zone", values="Events", title=title)
    fig.update_traces(textposition="inside", textinfo="percent+label")
    return fig


def bar_chart(df, title):
    fig = px.bar(df.sort_values("Events", ascending=True), x="Events", y="Zone", orientation="h", title=title)
    return fig


def build_word_report(title, comment, record):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("Summary Data", level=1)

    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Comment", level=1)
    doc.add_paragraph(comment)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.sidebar.header("Settings")
ep_threshold = st.sidebar.number_input("E/P Threshold", min_value=0.0, value=DEFAULT_EP_THRESHOLD)
et_threshold = st.sidebar.number_input("E/T Threshold", min_value=0.0, value=DEFAULT_ET_THRESHOLD)

st.markdown("""
<style>
:root {--navy:#082b5c;--blue:#0f4c8a;--border:#d8e0ea;--soft:#f6f9fc;}
.block-container {padding-top:0.35rem; padding-bottom:1.5rem; max-width:1900px;}
[data-testid="stSidebar"] {background:#f4f7fb; border-right:1px solid var(--border);}
[data-testid="stMetric"] {background:#ffffff; border:1px solid var(--border); padding:0.72rem 0.85rem; border-radius:8px; box-shadow:0 1px 2px rgba(15,23,42,.04);}
[data-testid="stMetricLabel"] {font-weight:700; color:#334155;}
[data-testid="stMetricValue"] {font-weight:800; color:#0f172a;}
.status-badge {display:inline-block;padding:.45rem .85rem;border-radius:999px;color:white;font-weight:800;margin:.15rem 0 .7rem 0;}
.gdi-header {background:linear-gradient(90deg,#062650,#0b3d78); color:white; padding:.65rem 1rem; border-radius:0 0 8px 8px; margin:-.35rem -1rem .65rem -1rem; display:flex; align-items:center; justify-content:space-between;}
.gdi-header-title {font-size:1.35rem;font-weight:800;line-height:1.1;}
.gdi-header-sub {font-size:.8rem;opacity:.9;margin-top:.15rem;}
.map-title {font-size:1.25rem;font-weight:800;color:#0f172a;margin:.15rem 0 .35rem 0;}
.legend-row {display:flex;gap:1.05rem;align-items:center;flex-wrap:wrap;margin:.1rem 0 .45rem 0;font-size:.86rem;}
.legend-item {display:flex;gap:.35rem;align-items:center;}
.legend-box {width:13px;height:13px;border-radius:2px;display:inline-block;}
.section-bar {background:#0b3d78;color:white;padding:.42rem .65rem;border-radius:6px 6px 0 0;font-weight:800;margin-top:.1rem;}
div[data-testid="stDataFrame"] {border:1px solid var(--border); border-radius:7px; overflow:hidden;}
.stTabs [data-baseweb="tab-list"] {gap:.25rem;}
.stTabs [data-baseweb="tab"] {font-weight:700;}
</style>
<div class="gdi-header">
  <div><div class="gdi-header-title">GDI — GEOTECHNICAL DIGITAL INTELLIGENCE PLATFORM</div><div class="gdi-header-sub">Draw Compliance & Seismic Decision Support System</div></div>
  <div style="font-weight:700;">Lift II East</div>
</div>
""", unsafe_allow_html=True)

tab_polygon, tab_daily, tab_weekend, tab_weekly, tab_insights, tab_history = st.tabs([
    "🗺️ Mine Dashboard",
    "📈 Daily Analysis",
    "📅 Weekend Analysis",
    "📊 Weekly Dashboard",
    "💡 Insights",
    "🗄️ History / Admin"
])



with tab_polygon:
    polygon_date = st.date_input("Reporting Date", value=datetime.now().date(), key="polygon_date")

    if "selected_polygon" not in st.session_state:
        st.session_state.selected_polygon = "A"

    mapping_df = load_csv(DRAWPOINT_MAPPING_FILE)
    merged_tonnes = pd.DataFrame()
    drawbell_totals = pd.DataFrame(columns=["Polygon", "Drawbell", "Tonnes"])
    polygon_totals = pd.DataFrame(columns=["Polygon", "Polygon Tonnes"])

    source_mode = st.radio(
        "Production input",
        ["Manual drawpoint entry", "Upload spreadsheet"],
        horizontal=True,
        key="polygon_production_mode",
    )

    if source_mode == "Upload spreadsheet":
        tonnes_upload = st.file_uploader(
            "Upload production spreadsheet — drawpoint tonnes",
            type=["xlsx", "xls", "csv"],
            key="polygon_tonnes_upload",
        )
        tonnes_df = read_table(tonnes_upload) if tonnes_upload else pd.DataFrame()

        if not tonnes_df.empty:
            p1, p2 = st.columns(2)
            drawpoint_col = p1.selectbox("Drawpoint column", tonnes_df.columns, key="polygon_drawpoint_col")
            tonnes_col = p2.selectbox("Tonnes column", tonnes_df.columns, key="polygon_tonnes_col")
            template = build_mapping_template(tonnes_df[drawpoint_col])
            bell_mapping = build_drawbell_mapping(template, mapping_df)

            with st.expander("Drawbell → polygon master mapping", expanded=bool(mapping_df.empty)):
                edited_bell_mapping = st.data_editor(
                    bell_mapping, width="stretch", hide_index=True, disabled=["Drawbell"],
                    column_config={
                        "Polygon": st.column_config.SelectboxColumn(
                            "Polygon", options=["", "A", "B", "C", "D", "E", "F"]
                        )
                    },
                    key="polygon_drawbell_mapping_editor",
                )
                edited_mapping = expand_drawbell_mapping(template, edited_bell_mapping)
                if st.button("Save Master Mapping", key="save_polygon_mapping"):
                    save_csv(edited_mapping, DRAWPOINT_MAPPING_FILE)
                    st.success("Master mapping saved.")

            merged_tonnes, drawbell_totals, polygon_totals = aggregate_polygon_tonnes(
                tonnes_df, drawpoint_col, tonnes_col, edited_mapping
            )
        else:
            st.caption("Upload a spreadsheet to calculate polygon tonnes automatically.")

    else:
        if mapping_df.empty:
            st.warning("Manual entry needs the master drawpoint mapping. Upload one production spreadsheet first, assign drawbells to polygons, and save the mapping.")
        else:
            manual_template = mapping_df[["Drawpoint", "Drawbell", "Polygon"]].copy()
            manual_template["Tonnes"] = 0.0
            st.caption("Enter tonnes per drawpoint. The app adds drawpoints → drawbells → polygons automatically.")
            manual_input = st.data_editor(
                manual_template,
                width="stretch",
                hide_index=True,
                disabled=["Drawpoint", "Drawbell", "Polygon"],
                column_config={"Tonnes": st.column_config.NumberColumn("Tonnes", min_value=0.0, step=1.0)},
                key="manual_drawpoint_tonnes",
            )
            merged_tonnes = manual_input.copy()
            drawbell_totals = (
                manual_input.groupby(["Polygon", "Drawbell"], as_index=False)["Tonnes"].sum()
            )
            polygon_totals = (
                manual_input.groupby("Polygon", as_index=False)["Tonnes"].sum()
                .rename(columns={"Tonnes": "Polygon Tonnes"})
            )

    with st.expander("Polygon seismic input", expanded=True):
        seismic_seed = pd.DataFrame({
            "Polygon": list("ABCDEF"), "Events": [0] * 6,
            "Energy J": [0.0] * 6, "Potency m3": [0.0] * 6,
            "Target Tonnes": [0.0] * 6,
        })
        seismic_input = st.data_editor(
            seismic_seed, width="stretch", hide_index=True, disabled=["Polygon"],
            key="polygon_seismic_editor_v4",
        )

    results = seismic_input.merge(polygon_totals, on="Polygon", how="left")
    for col in ["Polygon Tonnes", "Energy J", "Potency m3", "Events", "Target Tonnes"]:
        results[col] = pd.to_numeric(results[col], errors="coerce").fillna(0.0)

    results["E/P"] = results.apply(lambda r: r["Energy J"] / r["Potency m3"] if r["Potency m3"] > 0 else 0.0, axis=1)
    results["E/T"] = results.apply(lambda r: r["Energy J"] / r["Polygon Tonnes"] if r["Polygon Tonnes"] > 0 else None, axis=1)
    results["Status"] = results.apply(
        lambda r: get_status(r["E/P"], r["E/T"], ep_threshold, et_threshold) if r["Potency m3"] > 0 else "Waiting for seismic data",
        axis=1,
    )
    results["Map Rating"] = results["Status"].map(
        lambda s: "Green" if s == "Efficient Draw" else "Red" if s == "High-Risk Draw" else "Amber" if s in ["Brittle but Efficient", "Ductile but Costly"] else "Grey"
    )
    results["Health Score"] = results.apply(
        lambda r: polygon_health_score(r["E/P"], r["E/T"], r["Events"], r["Potency m3"], r["Polygon Tonnes"], r["Target Tonnes"], ep_threshold, et_threshold),
        axis=1,
    )
    results["Health Rating"] = results["Health Score"].map(polygon_health_rating)
    results["Achievement %"] = results.apply(
        lambda r: (100 * r["Polygon Tonnes"] / r["Target Tonnes"]) if r["Target Tonnes"] > 0 else None,
        axis=1,
    )

    rating_counts = results["Map Rating"].value_counts()
    g1, g2, g3, g4 = st.columns(4)
    g1.metric("Good", int(rating_counts.get("Green", 0)))
    g2.metric("Average / watch", int(rating_counts.get("Amber", 0)))
    g3.metric("Poor", int(rating_counts.get("Red", 0)))
    g4.metric("Waiting", int(rating_counts.get("Grey", 0)))

    left, right = st.columns([1.15, 1], gap="medium")
    with left:
        st.markdown('<div class="map-title">Lift II East — Polygon Status Map</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="legend-row">'
            '<span class="legend-item"><span class="legend-box" style="background:#15803d"></span>Good</span>'
            '<span class="legend-item"><span class="legend-box" style="background:#d97706"></span>Watch</span>'
            '<span class="legend-item"><span class="legend-box" style="background:#dc2626"></span>Poor</span>'
            '<span class="legend-item"><span class="legend-box" style="background:#cbd5e1"></span>No data</span>'
            '</div>', unsafe_allow_html=True
        )
        if DEFAULT_MAP_FILE.exists():
            st.plotly_chart(
                polygon_map_figure(results, DEFAULT_MAP_FILE, st.session_state.selected_polygon),
                width="stretch", key="polygon_status_map_v4",
                config={"displayModeBar": False, "displaylogo": False},
            )
            st.caption("Select a polygon using the controls on the right. The map colour updates from the live E/P–E/T classification.")
        else:
            st.error("Keep lift2east_polygon_map_clean.png in the same folder as this app.")

    with right:
        st.markdown('<div class="section-bar">SELECTED POLYGON</div>', unsafe_allow_html=True)
        bcols = st.columns(3)
        for index, poly in enumerate("ABCDEF"):
            if bcols[index % 3].button(poly, key=f"select_poly_v4_{poly}", use_container_width=True):
                st.session_state.selected_polygon = poly
                st.rerun()

        selected_poly = st.session_state.selected_polygon
        selected = results[results["Polygon"] == selected_poly].iloc[0]
        rating = selected["Map Rating"]
        badge_colour = {"Green": "#15803d", "Amber": "#d97706", "Red": "#dc2626", "Grey": "#64748b"}.get(rating, "#64748b")
        st.markdown(f"## POLYGON {selected_poly}")
        st.markdown(f'<span class="status-badge" style="background:{badge_colour}">{selected["Status"]}</span>', unsafe_allow_html=True)

        h1, h2 = st.columns([1, 1])
        h1.metric("Health Score", f"{int(selected['Health Score'])}/100")
        h2.metric("Health Rating", selected["Health Rating"])
        st.progress(int(selected["Health Score"]))

        k1, k2 = st.columns(2)
        k1.metric("Energy", f"{selected['Energy J']:,.0f} J")
        k2.metric("Potency", f"{selected['Potency m3']:,.3f} m³")
        k3, k4 = st.columns(2)
        k3.metric("Polygon tonnes", f"{selected['Polygon Tonnes']:,.0f} t")
        k4.metric("Events", f"{to_int(selected['Events'])}")
        k5, k6 = st.columns(2)
        k5.metric("E/P", f"{selected['E/P']:,.2f}")
        et_display = selected["E/T"]
        k6.metric("E/T", f"{et_display:,.2f}" if et_display is not None and not pd.isna(et_display) else "Waiting")

        if selected["Target Tonnes"] > 0:
            st.metric("Production achievement", f"{selected['Achievement %']:.1f}%", f"Target {selected['Target Tonnes']:,.0f} t")

        if selected["Potency m3"] > 0 and et_display is not None and not pd.isna(et_display):
            selected_history = load_csv(POLYGON_HISTORY_FILE)
            if not selected_history.empty and "Polygon" in selected_history.columns:
                selected_history = selected_history[selected_history["Polygon"].astype(str) == selected_poly]
            st.plotly_chart(
                plot_matrix(selected["E/P"], et_display, selected_history, ep_threshold, et_threshold),
                width="stretch", key=f"polygon_matrix_v4_{selected_poly}",
            )
        else:
            st.info("Enter energy, potency and tonnes to display the matrix position.")

    st.markdown(f'<div class="section-bar">DRAWBELLS & DRAWPOINTS — POLYGON {st.session_state.selected_polygon}</div>', unsafe_allow_html=True)
    if not drawbell_totals.empty:
        selected_poly = st.session_state.selected_polygon
        poly_drawbells = drawbell_totals[drawbell_totals["Polygon"] == selected_poly].copy()
        poly_points = merged_tonnes[merged_tonnes["Polygon"] == selected_poly][["Drawpoint", "Drawbell", "Tonnes"]].copy()
        d1, d2 = st.columns([1, 1.3])
        d1.dataframe(poly_drawbells, width="stretch", hide_index=True)
        d2.dataframe(poly_points, width="stretch", hide_index=True)
    else:
        st.caption("Drawbell and drawpoint detail appears after tonnes are entered or uploaded.")

    st.markdown('<div class="section-bar">POLYGON SUMMARY</div>', unsafe_allow_html=True)
    summary_cols = ["Polygon", "Status", "Health Score", "Health Rating", "Polygon Tonnes", "Target Tonnes", "Achievement %", "E/P", "E/T", "Events"]
    st.dataframe(results[summary_cols], width="stretch", hide_index=True)

    if st.button("Save Polygon Results", type="primary", key="save_polygon_results_v4"):
        save_df = results.copy()
        save_df.insert(0, "Report Date", str(polygon_date))
        existing = load_csv(POLYGON_HISTORY_FILE)
        combined = pd.concat([existing, save_df], ignore_index=True) if not existing.empty else save_df
        combined = combined.drop_duplicates(subset=["Report Date", "Polygon"], keep="last")
        save_csv(combined, POLYGON_HISTORY_FILE)
        st.success("Polygon results saved.")

with tab_daily:
    st.header("Daily Draw Intelligence")

    uploaded_pdf = st.file_uploader("Upload Daily IMS PDF", type=["pdf"], key="daily_pdf")
    uploaded_excel = st.file_uploader("Optional: Upload Production Tonnes Excel", type=["xlsx", "xls", "csv"], key="daily_excel")

    tonnes_from_excel = None

    if uploaded_excel:
        tonnes_df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith(".csv") else pd.read_excel(uploaded_excel)
        st.dataframe(tonnes_df.head(), width="stretch")
        tonnes_col = st.selectbox("Select Tonnes Column", tonnes_df.columns, key="daily_tonnes_col")
        tonnes_from_excel = pd.to_numeric(tonnes_df[tonnes_col], errors="coerce").sum()
        st.success(f"Tonnes from Excel: {tonnes_from_excel:,.2f}")

    if uploaded_pdf:
        full_text = pdf_to_text(uploaded_pdf)
        poly_data = parse_poly_sections(full_text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index)
            data = poly_data[target_poly]

            st.success(f"{target_poly} section found.")

            c1, c2, c3 = st.columns(3)
            c1.metric("New Events", data["Events"])
            c1.metric("Total Energy (J)", f"{data['Energy J']:,.2f}")
            c2.metric("Total Potency (m³)", f"{data['Potency m3']:,.3f}")
            c2.metric("Current Activity Rate", data["Current Activity Rate"])
            c3.metric("Medium Activity Rate", data["Medium Activity Rate"])
            c3.metric("Last Event", data["Last Event"])

            st.header("Tonnes Input")
            input_mode = st.radio("Tonnes Source", ["Manual Entry", "Excel Upload"], horizontal=True, key="daily_mode")

            if input_mode == "Excel Upload" and tonnes_from_excel is not None:
                tonnes = tonnes_from_excel
            else:
                tonnes = st.number_input("Tonnes Drawn", min_value=0.0, value=0.0, step=1.0, key="daily_tonnes")

            energy = data["Energy J"]
            potency = data["Potency m3"]

            if potency > 0:
                ep = energy / potency
                et = energy / tonnes if tonnes > 0 else None
                status = get_status(ep, et, ep_threshold, et_threshold)
                health = cave_health_score(status)
                action = recommended_action(status)

                st.header("Draw Intelligence Results")
                r1, r2, r3, r4 = st.columns(4)
                r1.metric("E/P", f"{ep:,.2f}")
                r2.metric("E/T", f"{et:,.2f}" if et is not None else "Waiting")
                r3.metric("Status", status)
                r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                score = draw_compliance_score(ep, et, data["Events"], potency, ep_threshold, et_threshold)
                s1, s2 = st.columns(2)
                s1.metric("Draw Compliance Score", f"{score}/100")
                s2.metric("Score Label", score_label(score))
                st.info(draw_advisor(score, status))

                st.info(action)

                if et is not None:
                    st.plotly_chart(plot_matrix(ep, et, load_csv(DAILY_FILE), ep_threshold, et_threshold), width="stretch")

                report_date = st.text_input("Report Date", value=clean_report_date(data["Last Event"]), key="daily_date")

                comment = management_summary(
                    target_poly, data["Events"], energy, potency, tonnes, ep, et, status, score
                )

                st.text_area("Management Comment", value=comment, height=170)

                if st.button("Save or Update Daily Record"):
                    record = {
                        "Report Date": report_date,
                        "Poly": target_poly,
                        "New Events": data["Events"],
                        "Energy J": energy,
                        "Potency m3": potency,
                        "Tonnes": tonnes,
                        "E/P": ep,
                        "E/T": et if et is not None else "",
                        "Status": status,
                        "Cave Health Score": health,
                        "Draw Compliance Score": score,
                        "Score Label": score_label(score),
                        "Recommended Action": action,
                    }
                    save_or_update(DAILY_FILE, record, ["Report Date", "Poly"])
                    st.success("Daily record saved or updated.")


with tab_weekend:
    st.header("Weekend Cumulative Processor")
    st.write("Save Friday baseline, Fri-Sat, Fri-Sun and Fri-Mon cumulative PDFs. GDI derives each separate day.")

    uploaded_weekend_pdf = st.file_uploader("Upload Weekend Cumulative IMS PDF", type=["pdf"], key="weekend_pdf")

    weekend_start = st.text_input("Weekend Start Date / Friday Date", value=datetime.now().strftime("%Y-%m-%d"))
    stage = st.selectbox(
        "Cumulative Stage",
        ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"]
    )

    stage_order = {
        "Friday Baseline": 0,
        "Friday to Saturday": 1,
        "Friday to Sunday": 2,
        "Friday to Monday": 3,
    }

    stage_daily_name = {
        "Friday Baseline": "Friday",
        "Friday to Saturday": "Saturday",
        "Friday to Sunday": "Sunday",
        "Friday to Monday": "Monday",
    }

    if uploaded_weekend_pdf:
        text = pdf_to_text(uploaded_weekend_pdf)
        poly_data = parse_poly_sections(text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index, key="weekend_poly")
            data = poly_data[target_poly]

            st.subheader("Cumulative Values")
            c1, c2, c3 = st.columns(3)
            c1.metric("Cumulative Events", data["Events"])
            c2.metric("Cumulative Energy (J)", f"{data['Energy J']:,.2f}")
            c3.metric("Cumulative Potency (m³)", f"{data['Potency m3']:,.3f}")

            cumulative_record = {
                "Weekend Start": weekend_start,
                "Poly": target_poly,
                "Stage": stage,
                "Stage Order": stage_order[stage],
                "Events": data["Events"],
                "Energy J": data["Energy J"],
                "Potency m3": data["Potency m3"],
                "Last Event": data["Last Event"],
            }

            if st.button("Save Weekend Cumulative Record"):
                save_or_update(WEEKEND_FILE, cumulative_record, ["Weekend Start", "Poly", "Stage"])
                st.success("Weekend cumulative record saved.")

            weekend_df = load_csv(WEEKEND_FILE)

            if not weekend_df.empty:
                filtered = weekend_df[
                    (weekend_df["Weekend Start"].astype(str) == weekend_start)
                    & (weekend_df["Poly"].astype(str) == target_poly)
                ].copy()

                filtered["Stage Order"] = pd.to_numeric(filtered["Stage Order"], errors="coerce")
                filtered = filtered.sort_values("Stage Order")

                st.subheader("Saved Weekend Cumulative Records")
                st.dataframe(filtered, width="stretch")

                current_order = stage_order[stage]

                if current_order == 0:
                    derived_events = data["Events"]
                    derived_energy = data["Energy J"]
                    derived_potency = data["Potency m3"]
                else:
                    previous = filtered[filtered["Stage Order"] == current_order - 1]

                    if previous.empty:
                        st.warning("Previous cumulative stage is missing. Upload/save the previous stage first.")
                        derived_events = None
                        derived_energy = None
                        derived_potency = None
                    else:
                        prev = previous.iloc[-1]
                        derived_events = data["Events"] - int(prev["Events"])
                        derived_energy = data["Energy J"] - float(prev["Energy J"])
                        derived_potency = data["Potency m3"] - float(prev["Potency m3"])

                if derived_events is not None:
                    derived_day = stage_daily_name[stage]

                    st.subheader("Derived Separate Day Values")
                    derived_df = pd.DataFrame([{
                        "Derived Day": derived_day,
                        "Events": derived_events,
                        "Energy J": derived_energy,
                        "Potency m3": derived_potency,
                    }])
                    st.dataframe(derived_df, width="stretch")

                    if derived_energy < 0 or derived_potency < 0 or derived_events < 0:
                        st.error("Derived values are negative. Check the cumulative order.")
                    else:
                        st.success("Derived day calculated successfully.")

                        st.header("Weekend Derived Draw Compliance")
                        weekend_tonnes = st.number_input(
                            f"{derived_day} Tonnes Drawn",
                            min_value=0.0,
                            value=0.0,
                            step=1.0,
                            key="weekend_tonnes"
                        )

                        if derived_potency > 0:
                            ep = derived_energy / derived_potency
                            et = derived_energy / weekend_tonnes if weekend_tonnes > 0 else None
                            status = get_status(ep, et, ep_threshold, et_threshold)
                            health = cave_health_score(status)
                            action = recommended_action(status)

                            r1, r2, r3, r4 = st.columns(4)
                            r1.metric("Derived E/P", f"{ep:,.2f}")
                            r2.metric("Derived E/T", f"{et:,.2f}" if et is not None else "Waiting")
                            r3.metric("Status", status)
                            r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                            score = draw_compliance_score(ep, et, derived_events, derived_potency, ep_threshold, et_threshold)
                            ws1, ws2 = st.columns(2)
                            ws1.metric("Draw Compliance Score", f"{score}/100")
                            ws2.metric("Score Label", score_label(score))
                            st.info(draw_advisor(score, status))

                            st.info(action)

                            if et is not None:
                                st.plotly_chart(
                                    plot_matrix(ep, et, load_csv(WEEKEND_DERIVED_FILE), ep_threshold, et_threshold),
                                    width="stretch"
                                )

                            if st.button("Save Weekend Derived Daily Record"):
                                actual_report_date = weekend_actual_date(weekend_start, derived_day)

                                derived_record = {
                                    "Weekend Start": weekend_start,
                                    "Derived Day": derived_day,
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Draw Compliance Score": score,
                                    "Score Label": score_label(score),
                                    "Recommended Action": action,
                                }

                                daily_record = {
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "New Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Draw Compliance Score": score,
                                    "Score Label": score_label(score),
                                    "Recommended Action": action,
                                }

                                save_or_update(WEEKEND_DERIVED_FILE, derived_record, ["Weekend Start", "Derived Day", "Poly"])
                                save_or_update(DAILY_FILE, daily_record, ["Report Date", "Poly"])
                                st.success("Weekend derived daily record saved and Daily Draw History updated.")


with tab_weekly:
    st.header("Weekly IMS Dashboard")

    uploaded_weekly_pdf = st.file_uploader("Upload Monday-to-Monday Weekly IMS Summary", type=["pdf"], key="weekly_pdf")

    if uploaded_weekly_pdf:
        text = pdf_to_text(uploaded_weekly_pdf)
        poly_data = parse_poly_sections(text)

        week_range = parse_week_range(text)
        mine_total = parse_mine_total(text)
        largest_ml = parse_largest_magnitude(text)

        st.subheader("Weekly Overview")
        o1, o2, o3 = st.columns(3)
        o1.metric("Reporting Period", week_range or "Not found")
        o2.metric("Mine-wide Events", mine_total)
        o3.metric("Largest ML", largest_ml if largest_ml is not None else "Not found")

        if poly_data:
            weekly_df = pd.DataFrame(poly_data.values())

            st.subheader("Weekly Zone Table")
            st.dataframe(weekly_df, width="stretch")

            weekly_df["Lift Group"] = weekly_df["Poly"].apply(
                lambda x: "Lift II" if str(x).startswith("Lift_II")
                else "Lift I" if str(x).startswith("Lift_I_")
                else "Other"
            )

            lift_split = weekly_df.groupby("Lift Group", as_index=False)["Events"].sum()
            lift_split = lift_split.rename(columns={"Lift Group": "Zone"})

            lift_i = weekly_df[weekly_df["Lift Group"] == "Lift I"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})
            lift_ii = weekly_df[weekly_df["Lift Group"] == "Lift II"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})

            st.plotly_chart(pie_chart(lift_split, "Mine-wide Event Split"), width="stretch")

            c1, c2 = st.columns(2)
            with c1:
                if not lift_i.empty:
                    st.plotly_chart(pie_chart(lift_i, "Lift I Event Distribution"), width="stretch")
            with c2:
                if not lift_ii.empty:
                    st.plotly_chart(pie_chart(lift_ii, "Lift II Event Distribution"), width="stretch")

            ranking = weekly_df.sort_values("Events", ascending=False)[["Poly", "Events", "Energy J", "Potency m3", "Comment"]]
            st.subheader("Hotspot Ranking")
            st.dataframe(ranking, width="stretch")
            st.plotly_chart(bar_chart(ranking.rename(columns={"Poly": "Zone"}), "Weekly Hotspot Ranking"), width="stretch")

            dominant = ranking.iloc[0]["Poly"] if not ranking.empty else "Not found"
            second = ranking.iloc[1]["Poly"] if len(ranking) > 1 else "Not found"

            weekly_comment = (
                f"Weekly seismicity for the period {week_range} recorded {mine_total} mine-wide events. "
                f"The dominant hotspot was {dominant}, followed by {second}. "
                f"The largest recorded event was ML{largest_ml}. Continued monitoring is recommended."
            )

            st.subheader("Weekly Executive Summary")
            st.text_area("Copy-ready weekly summary", value=weekly_comment, height=170)

            st.header("Weekly Lift II East Draw Indicator")

            if "Lift_II_E" in poly_data:
                liie = poly_data["Lift_II_E"]

                w1, w2, w3 = st.columns(3)
                w1.metric("Lift II East Weekly Events", liie["Events"])
                w2.metric("Lift II East Weekly Energy (J)", f"{liie['Energy J']:,.2f}")
                w3.metric("Lift II East Weekly Potency (m³)", f"{liie['Potency m3']:,.3f}")

                weekly_tonnes = st.number_input(
                    "Weekly Lift II East Tonnes Drawn",
                    min_value=0.0,
                    value=0.0,
                    step=1.0,
                    key="weekly_tonnes"
                )

                if liie["Potency m3"] > 0:
                    weekly_ep = liie["Energy J"] / liie["Potency m3"]
                    weekly_et = liie["Energy J"] / weekly_tonnes if weekly_tonnes > 0 else None
                    weekly_status = get_status(weekly_ep, weekly_et, ep_threshold, et_threshold)
                    weekly_health = cave_health_score(weekly_status)
                    weekly_action = recommended_action(weekly_status)

                    q1, q2, q3, q4 = st.columns(4)
                    q1.metric("Weekly E/P", f"{weekly_ep:,.2f}")
                    q2.metric("Weekly E/T", f"{weekly_et:,.2f}" if weekly_et is not None else "Waiting")
                    q3.metric("Weekly Status", weekly_status)
                    q4.metric("Weekly Cave Health", f"{weekly_health}/100" if weekly_health else "Waiting")

                    st.info(weekly_action)

                    if weekly_et is not None:
                        st.plotly_chart(
                            plot_matrix(weekly_ep, weekly_et, load_csv(WEEKLY_FILE), ep_threshold, et_threshold),
                            width="stretch"
                        )

                    weekly_draw_comment = (
                        f"For the weekly period {week_range}, Lift II East recorded {liie['Events']} events, "
                        f"with total energy of {liie['Energy J']:,.0f} J and total potency of {liie['Potency m3']:.3f} m³. "
                        f"The weekly E/P ratio is {weekly_ep:,.2f}. "
                    )

                    if weekly_et is not None:
                        weekly_draw_comment += (
                            f"The weekly E/T ratio is {weekly_et:,.2f}, placing Lift II East in the "
                            f"'{weekly_status}' quadrant. {weekly_action}"
                        )
                    else:
                        weekly_draw_comment += "Weekly E/T cannot be finalised until weekly tonnes are entered."

                    st.text_area("Weekly Lift II East Draw Comment", value=weekly_draw_comment, height=150)

                    if st.button("Save Weekly Summary"):
                        weekly_record = {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Second Hotspot": second,
                            "Lift II East Events": liie["Events"],
                            "Lift II East Energy J": liie["Energy J"],
                            "Lift II East Potency m3": liie["Potency m3"],
                            "Weekly Tonnes": weekly_tonnes,
                            "Weekly E/P": weekly_ep,
                            "Weekly E/T": weekly_et if weekly_et is not None else "",
                            "Weekly Status": weekly_status,
                            "Weekly Cave Health Score": weekly_health,
                            "Weekly Comment": weekly_comment,
                            "Weekly Draw Comment": weekly_draw_comment,
                        }

                        save_or_update(WEEKLY_FILE, weekly_record, ["Week Range"])
                        st.success("Weekly summary saved.")

                    word_file = build_word_report(
                        "GDI Weekly IMS Summary",
                        weekly_comment + "\n\n" + weekly_draw_comment,
                        {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Lift II East Weekly E/P": round(weekly_ep, 2),
                            "Lift II East Weekly E/T": round(weekly_et, 2) if weekly_et is not None else "Waiting",
                            "Weekly Status": weekly_status,
                        }
                    )

                    st.download_button(
                        "Download Weekly Word Report",
                        data=word_file,
                        file_name="GDI_Weekly_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Lift_II_E section was not found in the weekly PDF.")


with tab_insights:
    st.header("GDI Insights")

    daily_insights = add_score_columns(safe_sort_history(load_csv(DAILY_FILE)), ep_threshold, et_threshold)
    weekly_insights = safe_sort_history(load_csv(WEEKLY_FILE))

    if daily_insights.empty:
        st.info("No daily records available yet. Save daily or weekend-derived records first.")
    else:
        latest = daily_insights.iloc[-1]
        latest_score = to_int(latest.get("Draw Compliance Score", 0))
        latest_status = str(latest.get("Status", ""))

        st.subheader("Current Draw Snapshot")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Latest Date", latest.get("Report Date", ""))
        c2.metric("Latest Poly", latest.get("Poly", ""))
        c3.metric("Compliance Score", f"{latest_score}/100")
        c4.metric("Status", latest_status)
        st.info(draw_advisor(latest_score, latest_status))

        st.subheader("30-Day Trend Forecast")
        forecast_ep = simple_forecast(daily_insights.tail(30), "E/P")
        forecast_et = simple_forecast(daily_insights.tail(30), "E/T")
        f1, f2 = st.columns(2)
        f1.metric("Expected E/P", f"{forecast_ep:,.2f}" if forecast_ep is not None else "Waiting")
        f2.metric("Expected E/T", f"{forecast_et:,.2f}" if forecast_et is not None else "Waiting")

        st.subheader("Draw Compliance Calendar")
        calendar_df = daily_insights.copy()
        calendar_df["Report Date"] = pd.to_datetime(calendar_df["Report Date"], errors="coerce")
        calendar_df = calendar_df.dropna(subset=["Report Date"])
        if not calendar_df.empty:
            calendar_df["Day"] = calendar_df["Report Date"].dt.strftime("%Y-%m-%d")
            fig = px.scatter(
                calendar_df,
                x="Day",
                y="Poly",
                size="Draw Compliance Score",
                color="Score Label",
                hover_data=["E/P", "E/T", "Status", "Tonnes"],
                title="Daily Draw Compliance Calendar"
            )
            fig.update_layout(height=500)
            st.plotly_chart(fig, width="stretch")

        st.subheader("Hotspot / Concern Ranking")
        ranking_df = daily_insights.copy()
        ranking_df["Concern Level"] = 100 - pd.to_numeric(ranking_df["Draw Compliance Score"], errors="coerce")
        ranking_df = ranking_df.sort_values("Concern Level", ascending=False)
        st.dataframe(
            ranking_df[[col for col in ["Report Date", "Poly", "New Events", "Energy J", "Potency m3", "Tonnes", "E/P", "E/T", "Status", "Draw Compliance Score", "Score Label"] if col in ranking_df.columns]].head(20),
            width="stretch"
        )

        st.subheader("Automatic Executive Summary")
        summary_text = management_summary(
            latest.get("Poly", "Selected area"),
            to_int(latest.get("New Events", latest.get("Events", 0))),
            to_float(latest.get("Energy J", 0.0)),
            to_float(latest.get("Potency m3", 0.0)),
            to_float(latest.get("Tonnes", 0.0)),
            to_float(latest.get("E/P", 0.0)),
            None if str(latest.get("E/T", "")).strip() == "" else to_float(latest.get("E/T", 0.0)),
            latest_status,
            latest_score,
        )
        st.text_area("Copy-ready management summary", value=summary_text, height=170)

    if not weekly_insights.empty:
        st.subheader("Saved Weekly Summaries")
        st.dataframe(weekly_insights.tail(10), width="stretch")


with tab_history:
    st.header("History / Admin")

    daily_history = safe_sort_history(load_csv(DAILY_FILE))
    weekend_history = safe_sort_history(load_csv(WEEKEND_FILE))
    weekend_derived_history = safe_sort_history(load_csv(WEEKEND_DERIVED_FILE))
    weekly_history = safe_sort_history(load_csv(WEEKLY_FILE))

    st.subheader("Daily Draw History")
    if daily_history.empty:
        st.info("No daily records saved yet.")
    else:
        st.dataframe(daily_history, width="stretch")

        selected_daily_row = st.selectbox(
            "Select Daily Record",
            options=list(range(len(daily_history))),
            format_func=lambda x: f"{daily_history.iloc[x].get('Report Date', '')} - {daily_history.iloc[x].get('Poly', '')}",
            key="daily_selected_row",
        )

        daily_row = daily_history.iloc[selected_daily_row]

        with st.expander("Edit Selected Daily Record", expanded=False):
            with st.form("edit_daily_form"):
                edit_report_date = st.text_input("Report Date", value=str(row_value(daily_row, "Report Date")), key="edit_daily_report_date")
                edit_poly = st.text_input("Poly", value=str(row_value(daily_row, "Poly")), key="edit_daily_poly")
                edit_events = st.number_input("New Events", min_value=0, value=to_int(row_value(daily_row, "New Events", 0)), step=1, key="edit_daily_events")
                edit_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(daily_row, "Energy J", 0.0)), step=1000.0, key="edit_daily_energy")
                edit_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(daily_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_daily_potency")
                edit_tonnes = st.number_input("Tonnes", min_value=0.0, value=to_float(row_value(daily_row, "Tonnes", 0.0)), step=1.0, key="edit_daily_tonnes")
                save_daily_edit = st.form_submit_button("Save Edited Daily Record")

            if save_daily_edit:
                updated_record = {
                    "Report Date": edit_report_date,
                    "Poly": edit_poly,
                    "New Events": edit_events,
                    "Energy J": edit_energy,
                    "Potency m3": edit_potency,
                    "Tonnes": edit_tonnes,
                }
                updated_record = recalc_draw_record(updated_record, ep_threshold, et_threshold)

                daily_history = daily_history.drop(daily_history.index[selected_daily_row])
                daily_history = pd.concat([daily_history, pd.DataFrame([updated_record])], ignore_index=True)
                daily_history = daily_history.drop_duplicates(subset=["Report Date", "Poly"], keep="last")
                daily_history = safe_sort_history(daily_history)
                save_csv(daily_history, DAILY_FILE)
                st.success("Daily record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Daily Record"):
            daily_history = daily_history.drop(daily_history.index[selected_daily_row])
            save_csv(daily_history, DAILY_FILE)
            st.success("Daily record deleted.")
            st.rerun()

        st.download_button(
            "Download Daily History CSV",
            daily_history.to_csv(index=False).encode("utf-8"),
            "gdi_history.csv",
            "text/csv",
        )

    st.subheader("Weekend Cumulative History")
    if weekend_history.empty:
        st.info("No weekend cumulative records saved yet.")
    else:
        st.dataframe(weekend_history, width="stretch")

        selected_weekend_row = st.selectbox(
            "Select Weekend Cumulative Record",
            options=list(range(len(weekend_history))),
            format_func=lambda x: f"{weekend_history.iloc[x].get('Weekend Start', '')} - {weekend_history.iloc[x].get('Poly', '')} - {weekend_history.iloc[x].get('Stage', '')}",
            key="weekend_selected_row",
        )

        weekend_row = weekend_history.iloc[selected_weekend_row]

        with st.expander("Edit Selected Weekend Cumulative Record", expanded=False):
            with st.form("edit_weekend_form"):
                edit_weekend_start = st.text_input("Weekend Start", value=str(row_value(weekend_row, "Weekend Start")), key="edit_weekend_start")
                edit_weekend_poly = st.text_input("Poly", value=str(row_value(weekend_row, "Poly")), key="edit_weekend_poly")
                edit_weekend_stage = st.selectbox(
                    "Stage",
                    ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"],
                    index=["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"].index(row_value(weekend_row, "Stage", "Friday Baseline")) if row_value(weekend_row, "Stage", "Friday Baseline") in ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"] else 0,
                    key="edit_weekend_stage",
                )
                edit_weekend_events = st.number_input("Events", min_value=0, value=to_int(row_value(weekend_row, "Events", 0)), step=1, key="edit_weekend_events")
                edit_weekend_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(weekend_row, "Energy J", 0.0)), step=1000.0, key="edit_weekend_energy")
                edit_weekend_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(weekend_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_weekend_potency")
                edit_weekend_last_event = st.text_input("Last Event", value=str(row_value(weekend_row, "Last Event")), key="edit_weekend_last_event")
                save_weekend_edit = st.form_submit_button("Save Edited Weekend Cumulative Record")

            if save_weekend_edit:
                stage_order_edit = {
                    "Friday Baseline": 0,
                    "Friday to Saturday": 1,
                    "Friday to Sunday": 2,
                    "Friday to Monday": 3,
                }
                updated_weekend_record = {
                    "Weekend Start": edit_weekend_start,
                    "Poly": edit_weekend_poly,
                    "Stage": edit_weekend_stage,
                    "Stage Order": stage_order_edit[edit_weekend_stage],
                    "Events": edit_weekend_events,
                    "Energy J": edit_weekend_energy,
                    "Potency m3": edit_weekend_potency,
                    "Last Event": edit_weekend_last_event,
                }

                weekend_history = weekend_history.drop(weekend_history.index[selected_weekend_row])
                weekend_history = pd.concat([weekend_history, pd.DataFrame([updated_weekend_record])], ignore_index=True)
                weekend_history = weekend_history.drop_duplicates(subset=["Weekend Start", "Poly", "Stage"], keep="last")
                weekend_history = safe_sort_history(weekend_history)
                save_csv(weekend_history, WEEKEND_FILE)
                st.success("Weekend cumulative record edited.")
                st.rerun()

        if st.button("Delete Selected Weekend Cumulative Record"):
            weekend_history = weekend_history.drop(weekend_history.index[selected_weekend_row])
            save_csv(weekend_history, WEEKEND_FILE)
            st.success("Weekend cumulative record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Cumulative CSV",
            weekend_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_cumulative.csv",
            "text/csv",
        )

    st.subheader("Weekend Derived Daily History")
    if weekend_derived_history.empty:
        st.info("No weekend derived records saved yet.")
    else:
        st.dataframe(weekend_derived_history, width="stretch")

        selected_derived_row = st.selectbox(
            "Select Weekend Derived Record",
            options=list(range(len(weekend_derived_history))),
            format_func=lambda x: f"{weekend_derived_history.iloc[x].get('Report Date', weekend_derived_history.iloc[x].get('Weekend Start', ''))} - {weekend_derived_history.iloc[x].get('Derived Day', '')} - {weekend_derived_history.iloc[x].get('Poly', '')}",
            key="derived_selected_row",
        )

        derived_row = weekend_derived_history.iloc[selected_derived_row]

        with st.expander("Edit Selected Weekend Derived Record", expanded=False):
            with st.form("edit_derived_form"):
                edit_derived_start = st.text_input("Weekend Start", value=str(row_value(derived_row, "Weekend Start")), key="edit_derived_start")
                edit_derived_day = st.selectbox(
                    "Derived Day",
                    ["Friday", "Saturday", "Sunday", "Monday"],
                    index=["Friday", "Saturday", "Sunday", "Monday"].index(row_value(derived_row, "Derived Day", "Friday")) if row_value(derived_row, "Derived Day", "Friday") in ["Friday", "Saturday", "Sunday", "Monday"] else 0,
                    key="edit_derived_day",
                )
                edit_derived_report_date = st.text_input("Report Date", value=str(row_value(derived_row, "Report Date", weekend_actual_date(edit_derived_start, edit_derived_day))), key="edit_derived_report_date")
                edit_derived_poly = st.text_input("Poly", value=str(row_value(derived_row, "Poly")), key="edit_derived_poly")
                edit_derived_events = st.number_input("Events", min_value=0, value=to_int(row_value(derived_row, "Events", 0)), step=1, key="edit_derived_events")
                edit_derived_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(derived_row, "Energy J", 0.0)), step=1000.0, key="edit_derived_energy")
                edit_derived_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(derived_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_derived_potency")
                edit_derived_tonnes = st.number_input("Tonnes", min_value=0.0, value=to_float(row_value(derived_row, "Tonnes", 0.0)), step=1.0, key="edit_derived_tonnes")
                sync_daily = st.checkbox("Also update Daily Draw History", value=True, key="sync_derived_to_daily")
                save_derived_edit = st.form_submit_button("Save Edited Weekend Derived Record")

            if save_derived_edit:
                updated_derived_record = {
                    "Weekend Start": edit_derived_start,
                    "Derived Day": edit_derived_day,
                    "Report Date": edit_derived_report_date,
                    "Poly": edit_derived_poly,
                    "Events": edit_derived_events,
                    "Energy J": edit_derived_energy,
                    "Potency m3": edit_derived_potency,
                    "Tonnes": edit_derived_tonnes,
                }
                updated_derived_record = recalc_draw_record(updated_derived_record, ep_threshold, et_threshold)

                weekend_derived_history = weekend_derived_history.drop(weekend_derived_history.index[selected_derived_row])
                weekend_derived_history = pd.concat([weekend_derived_history, pd.DataFrame([updated_derived_record])], ignore_index=True)
                weekend_derived_history = weekend_derived_history.drop_duplicates(subset=["Weekend Start", "Derived Day", "Poly"], keep="last")
                weekend_derived_history = safe_sort_history(weekend_derived_history)
                save_csv(weekend_derived_history, WEEKEND_DERIVED_FILE)

                if sync_daily:
                    daily_record = {
                        "Report Date": edit_derived_report_date,
                        "Poly": edit_derived_poly,
                        "New Events": edit_derived_events,
                        "Energy J": edit_derived_energy,
                        "Potency m3": edit_derived_potency,
                        "Tonnes": edit_derived_tonnes,
                    }
                    daily_record = recalc_draw_record(daily_record, ep_threshold, et_threshold)
                    save_or_update(DAILY_FILE, daily_record, ["Report Date", "Poly"])

                st.success("Weekend derived record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Weekend Derived Record"):
            weekend_derived_history = weekend_derived_history.drop(weekend_derived_history.index[selected_derived_row])
            save_csv(weekend_derived_history, WEEKEND_DERIVED_FILE)
            st.success("Weekend derived record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Derived CSV",
            weekend_derived_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_derived_daily.csv",
            "text/csv",
        )

    st.subheader("Weekly Summary History")
    if weekly_history.empty:
        st.info("No weekly records saved yet.")
    else:
        st.dataframe(weekly_history, width="stretch")

        selected_weekly_row = st.selectbox(
            "Select Weekly Record",
            options=list(range(len(weekly_history))),
            format_func=lambda x: f"{weekly_history.iloc[x].get('Week Range', '')}",
            key="weekly_selected_row",
        )

        weekly_row = weekly_history.iloc[selected_weekly_row]

        with st.expander("Edit Selected Weekly Record", expanded=False):
            with st.form("edit_weekly_form"):
                edit_week_range = st.text_input("Week Range", value=str(row_value(weekly_row, "Week Range")), key="edit_week_range")
                edit_mine_events = st.number_input("Mine-wide Events", min_value=0, value=to_int(row_value(weekly_row, "Mine-wide Events", 0)), step=1, key="edit_mine_events")
                edit_largest_ml = st.number_input("Largest ML", value=to_float(row_value(weekly_row, "Largest ML", 0.0)), step=0.1, key="edit_largest_ml")
                edit_dominant = st.text_input("Dominant Hotspot", value=str(row_value(weekly_row, "Dominant Hotspot")), key="edit_dominant")
                edit_second = st.text_input("Second Hotspot", value=str(row_value(weekly_row, "Second Hotspot")), key="edit_second")
                edit_liie_events = st.number_input("Lift II East Events", min_value=0, value=to_int(row_value(weekly_row, "Lift II East Events", 0)), step=1, key="edit_liie_events")
                edit_liie_energy = st.number_input("Lift II East Energy J", min_value=0.0, value=to_float(row_value(weekly_row, "Lift II East Energy J", 0.0)), step=1000.0, key="edit_liie_energy")
                edit_liie_potency = st.number_input("Lift II East Potency m3", min_value=0.0, value=to_float(row_value(weekly_row, "Lift II East Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_liie_potency")
                edit_weekly_tonnes = st.number_input("Weekly Tonnes", min_value=0.0, value=to_float(row_value(weekly_row, "Weekly Tonnes", 0.0)), step=1.0, key="edit_weekly_tonnes")
                edit_weekly_comment = st.text_area("Weekly Comment", value=str(row_value(weekly_row, "Weekly Comment")), height=120, key="edit_weekly_comment")
                edit_weekly_draw_comment = st.text_area("Weekly Draw Comment", value=str(row_value(weekly_row, "Weekly Draw Comment")), height=120, key="edit_weekly_draw_comment")
                save_weekly_edit = st.form_submit_button("Save Edited Weekly Record")

            if save_weekly_edit:
                weekly_ep = edit_liie_energy / edit_liie_potency if edit_liie_potency > 0 else 0.0
                weekly_et = edit_liie_energy / edit_weekly_tonnes if edit_weekly_tonnes > 0 else None
                weekly_status = get_status(weekly_ep, weekly_et, ep_threshold, et_threshold)
                weekly_health = cave_health_score(weekly_status)

                updated_weekly_record = {
                    "Week Range": edit_week_range,
                    "Mine-wide Events": edit_mine_events,
                    "Largest ML": edit_largest_ml,
                    "Dominant Hotspot": edit_dominant,
                    "Second Hotspot": edit_second,
                    "Lift II East Events": edit_liie_events,
                    "Lift II East Energy J": edit_liie_energy,
                    "Lift II East Potency m3": edit_liie_potency,
                    "Weekly Tonnes": edit_weekly_tonnes,
                    "Weekly E/P": weekly_ep,
                    "Weekly E/T": weekly_et if weekly_et is not None else "",
                    "Weekly Status": weekly_status,
                    "Weekly Cave Health Score": weekly_health,
                    "Weekly Comment": edit_weekly_comment,
                    "Weekly Draw Comment": edit_weekly_draw_comment,
                }

                weekly_history = weekly_history.drop(weekly_history.index[selected_weekly_row])
                weekly_history = pd.concat([weekly_history, pd.DataFrame([updated_weekly_record])], ignore_index=True)
                weekly_history = weekly_history.drop_duplicates(subset=["Week Range"], keep="last")
                weekly_history = safe_sort_history(weekly_history)
                save_csv(weekly_history, WEEKLY_FILE)
                st.success("Weekly record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Weekly Record"):
            weekly_history = weekly_history.drop(weekly_history.index[selected_weekly_row])
            save_csv(weekly_history, WEEKLY_FILE)
            st.success("Weekly record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekly CSV",
            weekly_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekly_history.csv",
            "text/csv",
        )
