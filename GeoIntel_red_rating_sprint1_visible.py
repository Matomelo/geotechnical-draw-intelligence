import re
from io import BytesIO
from pathlib import Path
from datetime import datetime, timedelta

import pandas as pd
import pdfplumber
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from PIL import Image


st.set_page_config(page_title="GeoIntel — Geotechnical Decision Intelligence", layout="wide")

DAILY_FILE = Path("gdi_history.csv")
WEEKEND_FILE = Path("gdi_weekend_cumulative.csv")
WEEKEND_DERIVED_FILE = Path("gdi_weekend_derived_daily.csv")
WEEKLY_FILE = Path("gdi_weekly_history.csv")
DRAWPOINT_MAPPING_FILE = Path("gdi_drawpoint_polygon_mapping.csv")
POLYGON_HISTORY_FILE = Path("gdi_polygon_history.csv")
MANUAL_TONNES_FILE = Path("gdi_manual_drawpoint_tonnes.csv")
SETTINGS_FILE = Path("gdi_settings.csv")
RISK_ENGINE_FILE = Path("gdi_risk_engine_settings.csv")
CONVERGENCE_FILE = Path("gdi_polygon_convergence.csv")
AUDIT_FILE = Path("gdi_audit_trail.csv")
APPROVAL_FILE = Path("gdi_approval_status.csv")
HANDOVER_FILE = Path("gdi_shift_handover.csv")
RED_RATING_FILE = Path("geointel_red_rating_tracker.csv")
DEFAULT_MAP_FILE = Path("lift2east_polygon_map_clean_v2.png")

DEFAULT_EP_THRESHOLD = 10000.0
DEFAULT_ET_THRESHOLD = 10.0


def find_value(pattern, text):
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None


def pdf_to_text(uploaded_pdf):
    text = ""
    with pdfplumber.open(uploaded_pdf) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def load_csv(path):
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


def save_csv(df, path):
    df.to_csv(path, index=False)


def save_or_update(path, record, key_cols):
    df = load_csv(path)
    record_df = pd.DataFrame([record])

    if df.empty:
        df = record_df
    else:
        df = pd.concat([df, record_df], ignore_index=True)
        df = df.drop_duplicates(subset=key_cols, keep="last")

    save_csv(df, path)






def calculate_red_rating_duration(start_time, end_time):
    """Return duration in minutes and hours, allowing the event to cross midnight."""
    start_dt = datetime.combine(datetime.today().date(), start_time)
    end_dt = datetime.combine(datetime.today().date(), end_time)
    if end_dt < start_dt:
        end_dt += timedelta(days=1)
    minutes = max(0, int((end_dt - start_dt).total_seconds() // 60))
    return minutes, minutes / 60.0


def save_red_rating_record(record):
    existing = load_csv(RED_RATING_FILE)
    new_row = pd.DataFrame([record])
    combined = pd.concat([existing, new_row], ignore_index=True) if not existing.empty else new_row
    save_csv(combined, RED_RATING_FILE)


def red_rating_summary(df):
    if df.empty:
        return {"incidents": 0, "hours": 0.0, "open": 0, "top_location": "—"}
    work = df.copy()
    work["Duration Hours"] = pd.to_numeric(work.get("Duration Hours", 0), errors="coerce").fillna(0.0)
    open_count = int(work.get("Status", pd.Series(dtype=str)).astype(str).isin(["Open", "Under Review"]).sum()) if "Status" in work.columns else 0
    top_location = "—"
    if "Locations" in work.columns and not work["Locations"].dropna().empty:
        expanded = work["Locations"].dropna().astype(str).str.split(",").explode().str.strip()
        expanded = expanded[expanded.ne("")]
        if not expanded.empty:
            top_location = expanded.value_counts().index[0]
    return {"incidents": len(work), "hours": float(work["Duration Hours"].sum()), "open": open_count, "top_location": top_location}


def prepare_red_rating_data(df):
    """Clean saved/imported red-rating records without inventing operational values."""
    if df.empty:
        return df.copy()
    work = df.copy()
    work["Date"] = pd.to_datetime(work.get("Date"), errors="coerce", dayfirst=True)
    work["Duration Hours"] = pd.to_numeric(work.get("Duration Hours", 0), errors="coerce").fillna(0.0)
    for col, default in {
        "Locations":"", "Category":"Unspecified", "Activity Affected":"Unspecified",
        "Department / Function":"Unspecified", "Shift":"Unspecified", "Status":"Closed",
        "Trigger":"", "Action / Comments":"", "Outcome":"", "Responsible Person":""
    }.items():
        if col not in work.columns:
            work[col] = default
        work[col] = work[col].fillna(default).astype(str).replace("nan", default)
    return work


def explode_red_locations(df):
    if df.empty or "Locations" not in df.columns:
        return pd.DataFrame(columns=["Location", "Duration Hours", "Date", "Status"])
    cols = [c for c in ["Locations", "Duration Hours", "Date", "Status"] if c in df.columns]
    loc = df[cols].copy()
    loc["Location"] = loc["Locations"].astype(str).str.split(",")
    loc = loc.explode("Location")
    loc["Location"] = loc["Location"].astype(str).str.strip()
    return loc[loc["Location"].ne("")]


def reliability_score(count, hours):
    """Simple 0-100 operational reliability score: repeated/long restrictions reduce reliability."""
    penalty = min(100.0, float(count) * 12.0 + float(hours) * 3.0)
    return int(round(max(0.0, 100.0 - penalty)))



def _excel_time_text(value):
    """Convert Excel time cells, datetime/time objects or text into HH:MM."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    if hasattr(value, "strftime"):
        try:
            return value.strftime("%H:%M")
        except Exception:
            pass
    if isinstance(value, (int, float)):
        total_minutes = int(round((float(value) % 1) * 24 * 60)) % (24 * 60)
        return f"{total_minutes // 60:02d}:{total_minutes % 60:02d}"
    text = str(value).strip()
    parsed = pd.to_datetime(text, errors="coerce")
    return parsed.strftime("%H:%M") if not pd.isna(parsed) else text


def read_red_rating_excel(uploaded_file):
    """Read the approved Red Rating Tracker workbook into GeoIntel's tracker format."""
    if uploaded_file is None:
        return pd.DataFrame(), []
    try:
        raw = pd.read_excel(uploaded_file, sheet_name="Red Rating Delay Register")
    except ValueError:
        return pd.DataFrame(), ["Sheet 'Red Rating Delay Register' was not found."]
    except Exception as exc:
        return pd.DataFrame(), [f"Could not read workbook: {exc}"]

    required = ["Date", "Crosscut / Location", "Time Red Started", "End Time"]
    missing = [column for column in required if column not in raw.columns]
    if missing:
        return pd.DataFrame(), ["Missing required column(s): " + ", ".join(missing)]

    rows, warnings = [], []
    for idx, row in raw.iterrows():
        date_value = pd.to_datetime(row.get("Date"), errors="coerce", dayfirst=True)
        locations = str(row.get("Crosscut / Location", "")).strip()
        if pd.isna(date_value) or not locations or locations.lower() == "nan":
            continue

        start_text = _excel_time_text(row.get("Time Red Started"))
        end_text = _excel_time_text(row.get("End Time"))
        try:
            start_obj = datetime.strptime(start_text, "%H:%M").time()
            end_obj = datetime.strptime(end_text, "%H:%M").time()
            minutes, hours = calculate_red_rating_duration(start_obj, end_obj)
        except Exception:
            hours_cell = pd.to_numeric(pd.Series([row.get("Hours Lost")]), errors="coerce").iloc[0]
            hours = float(hours_cell * 24) if not pd.isna(hours_cell) and float(hours_cell) <= 1 else (float(hours_cell) if not pd.isna(hours_cell) else 0.0)
            minutes = int(round(hours * 60))
            warnings.append(f"Row {idx + 2}: duration was taken from Hours Lost because start/end time could not be read.")

        rating = str(row.get("Rating", "Red")).strip()
        reason = str(row.get("Reason for Red Rating", "")).strip()
        comments = str(row.get("Comments", "")).strip()
        status = str(row.get("Status", "Closed")).strip()
        activity = str(row.get("Activity Affected", "")).strip()
        if status.lower() in ["", "nan"]:
            status = "Closed"

        record_date = date_value.strftime("%Y-%m-%d")
        record_id = f"XLS-{date_value.strftime('%Y%m%d')}-{start_text.replace(':','')}-{idx + 2}"
        rows.append({
            "Record ID": record_id,
            "Date": record_date,
            "Start Time": start_text,
            "End Time": end_text,
            "Duration Minutes": minutes,
            "Duration Hours": round(hours, 3),
            "Locations": locations,
            "Category": rating if rating and rating.lower() != "nan" else "Red",
            "Activity Affected": "" if activity.lower() == "nan" else activity,
            "Trigger": "" if reason.lower() == "nan" else reason,
            "Action / Comments": "" if comments.lower() == "nan" else comments,
            "Responsible Person": str(row.get("Responsible Person", "")).strip() if str(row.get("Responsible Person", "")).lower() != "nan" else "",
            "Department / Function": str(row.get("Department / Function", row.get("Department", "Unspecified"))).strip(),
            "Shift": str(row.get("Shift", "Unspecified")).strip(),
            "Outcome": str(row.get("Outcome", "")).strip() if str(row.get("Outcome", "")).lower() != "nan" else "",
            "Status": status,
            "Entered By": "Excel import",
            "Timestamp": datetime.now().isoformat(timespec="seconds"),
        })

    return pd.DataFrame(rows), warnings


def merge_red_rating_import(import_df):
    """Merge imported workbook records without duplicating records already imported."""
    existing = load_csv(RED_RATING_FILE)
    combined = pd.concat([existing, import_df], ignore_index=True) if not existing.empty else import_df.copy()
    if "Record ID" in combined.columns:
        combined = combined.drop_duplicates(subset=["Record ID"], keep="last")
    save_csv(combined, RED_RATING_FILE)
    return combined


def weekend_actual_date(weekend_start, derived_day):
    """Return the real calendar date for Friday/Saturday/Sunday/Monday derived data."""
    offset_map = {
        "Friday": 0,
        "Saturday": 1,
        "Sunday": 2,
        "Monday": 3,
    }

    try:
        base_date = datetime.strptime(str(weekend_start), "%Y-%m-%d")
    except ValueError:
        return str(weekend_start)

    return (base_date + timedelta(days=offset_map.get(derived_day, 0))).strftime("%Y-%m-%d")


def safe_sort_history(df):
    if df.empty:
        return df
    df = df.copy()
    sort_candidates = ["Report Date", "Weekend Start", "Week Range"]
    for col in sort_candidates:
        if col in df.columns:
            return df.sort_values(col, kind="stable")
    return df


def to_float(value, default=0.0):
    try:
        if value is None or str(value).strip() == "":
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def to_int(value, default=0):
    try:
        if value is None or str(value).strip() == "":
            return default
        return int(float(value))
    except (TypeError, ValueError):
        return default


def row_value(row, column, default=""):
    if column not in row.index:
        return default
    value = row[column]
    if pd.isna(value):
        return default
    return value


def recalc_draw_record(record, ep_threshold, et_threshold):
    energy = to_float(record.get("Energy J", 0.0))
    potency = to_float(record.get("Potency m3", 0.0))
    tonnes = to_float(record.get("Tonnes", 0.0))

    ep = energy / potency if potency > 0 else 0.0
    et = energy / tonnes if tonnes > 0 else None
    status = get_status(ep, et, ep_threshold, et_threshold)
    health = cave_health_score(status)
    action = recommended_action(status)

    record["E/P"] = ep
    record["E/T"] = et if et is not None else ""
    events = to_int(record.get("New Events", record.get("Events", 0)))
    score = draw_compliance_score(ep, et, events, potency, ep_threshold, et_threshold)

    record["Status"] = status
    record["Cave Health Score"] = health
    record["Draw Compliance Score"] = score
    record["Score Label"] = score_label(score)
    record["Recommended Action"] = action
    return record



def draw_compliance_score(ep, et, events=0, potency=0.0, ep_threshold=DEFAULT_EP_THRESHOLD, et_threshold=DEFAULT_ET_THRESHOLD):
    """Return a management-friendly 0-100 draw compliance score."""
    if et is None:
        return 0

    ep_ratio = ep / ep_threshold if ep_threshold else 0
    et_ratio = et / et_threshold if et_threshold else 0

    ep_component = max(0, 40 - min(40, ep_ratio * 25))
    et_component = max(0, 40 - min(40, et_ratio * 25))

    event_component = 10
    if events >= 100:
        event_component = 2
    elif events >= 50:
        event_component = 5
    elif events >= 20:
        event_component = 8

    potency_component = 10
    if potency <= 0:
        potency_component = 0

    return int(round(max(0, min(100, ep_component + et_component + event_component + potency_component))))


def score_label(score):
    if score >= 80:
        return "Good Draw"
    if score >= 60:
        return "Watch"
    if score >= 40:
        return "Concern"
    return "Intervention Required"


def draw_advisor(score, status):
    if score >= 80:
        return "Draw response is favourable. Maintain current draw strategy and continue routine monitoring."
    if score >= 60:
        return "Draw response is acceptable but should be watched. Check whether the trend is improving or deteriorating before changing draw."
    if score >= 40:
        return "Draw response requires attention. Review tonnes distribution, adjacent drawpoints and recent seismic clustering."
    if status == "Waiting for tonnes":
        return "Enter tonnes drawn to finalise the draw compliance score and recommendation."
    return "Escalate for geotechnical review. Treat this as a potential draw-control concern until trends improve."


def management_summary(poly, events, energy, potency, tonnes, ep, et, status, score):
    et_text = f"{et:,.2f}" if et is not None else "waiting for tonnes"
    return (
        f"{poly} recorded {events} seismic events with total energy of {energy:,.0f} J and potency of {potency:.3f} m³. "
        f"Tonnes drawn were {tonnes:,.0f}. The calculated E/P is {ep:,.2f} and E/T is {et_text}. "
        f"The area is classified as {status}, with a Draw Compliance Score of {score}/100 ({score_label(score)}). "
        f"{draw_advisor(score, status)}"
    )


def add_score_columns(df, ep_threshold, et_threshold):
    if df.empty:
        return df
    df = df.copy()
    for col in ["E/P", "E/T", "New Events", "Events", "Potency m3"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    scores = []
    labels = []
    for _, row in df.iterrows():
        ep = to_float(row.get("E/P", 0.0))
        et_value = row.get("E/T", None)
        et = None if pd.isna(et_value) or str(et_value).strip() == "" else to_float(et_value)
        events = to_int(row.get("New Events", row.get("Events", 0)))
        potency = to_float(row.get("Potency m3", 0.0))
        score = draw_compliance_score(ep, et, events, potency, ep_threshold, et_threshold)
        scores.append(score)
        labels.append(score_label(score))
    df["Draw Compliance Score"] = scores
    df["Score Label"] = labels
    return df


def simple_forecast(df, column):
    if df.empty or column not in df.columns:
        return None
    values = pd.to_numeric(df[column], errors="coerce").dropna().tail(7)
    if values.empty:
        return None
    return float(values.mean())

def clean_number(value):
    if value is None:
        return None
    value = str(value).replace(",", "").strip()
    try:
        return float(value)
    except ValueError:
        return None


def clean_report_date(last_event):
    if not last_event:
        return datetime.now().strftime("%Y-%m-%d")

    match = re.search(r"(\d{1,2})\s+([A-Za-z]{3})", last_event)
    if not match:
        return datetime.now().strftime("%Y-%m-%d")

    day = int(match.group(1))
    month_text = match.group(2).title()

    month_map = {
        "Jan": 1, "Feb": 2, "Mar": 3, "Apr": 4,
        "May": 5, "Jun": 6, "Jul": 7, "Aug": 8,
        "Sep": 9, "Oct": 10, "Nov": 11, "Dec": 12
    }

    month = month_map.get(month_text, datetime.now().month)
    year = datetime.now().year
    return f"{year}-{month:02d}-{day:02d}"


def parse_poly_sections(text):
    poly_matches = list(re.finditer(r"Poly:\s*([A-Za-z0-9_]+)", text))
    results = {}

    for i, match in enumerate(poly_matches):
        poly = match.group(1)
        start = match.start()
        end = poly_matches[i + 1].start() if i + 1 < len(poly_matches) else len(text)
        block = text[start:end]

        last_event = find_value(r"Last event Date and Time:\s*(.+)", block)
        events = find_value(r"Number of new events:\s*(\d+)", block)
        energy = find_value(r"Total Energy:\s*([0-9.eE+-]+)", block)
        potency = find_value(r"Total Potency:\s*([0-9.eE+-]+)", block)
        current_rate = find_value(r"Current normalized activity rate:\s*([0-9.]+)", block)
        medium_rate = find_value(r"Medium term normalized activity rate:\s*([0-9.]+)", block)

        comment = ""
        comment_match = re.search(
            r"Medium term normalized activity rate:\s*[0-9.]+\s*(.*?)(?:\n\d+\s+of\s+\d+|\Z)",
            block,
            re.DOTALL
        )
        if comment_match:
            comment = " ".join(comment_match.group(1).split())

        if events or energy or potency:
            results[poly] = {
                "Poly": poly,
                "Last Event": last_event,
                "Events": int(events) if events else 0,
                "Energy J": clean_number(energy) or 0.0,
                "Potency m3": clean_number(potency) or 0.0,
                "Current Activity Rate": clean_number(current_rate),
                "Medium Activity Rate": clean_number(medium_rate),
                "Comment": comment,
            }

    return results


def parse_week_range(text):
    match = re.search(r"Seismic Report:\s*(.+)", text)
    if match:
        return match.group(1).strip()
    return ""


def parse_mine_total(text):
    total = find_value(r"Total number of events\s*=\s*(\d+)", text)
    return int(total) if total else 0


def parse_largest_magnitude(text):
    mags = re.findall(r"\b(?:ml|ML|M)(-?\d+(?:\.\d+)?)", text)
    values = [float(m) for m in mags if m is not None]
    return max(values) if values else None


def get_status(ep, et, ep_threshold, et_threshold):
    if et is None:
        return "Waiting for tonnes"
    if ep < ep_threshold and et < et_threshold:
        return "Efficient Draw"
    if ep >= ep_threshold and et < et_threshold:
        return "Brittle but Efficient"
    if ep < ep_threshold and et >= et_threshold:
        return "Ductile but Costly"
    return "High-Risk Draw"


def cave_health_score(status):
    return {
        "Efficient Draw": 90,
        "Brittle but Efficient": 65,
        "Ductile but Costly": 55,
        "High-Risk Draw": 30,
        "Waiting for tonnes": 0,
    }.get(status, 0)


def recommended_action(status):
    return {
        "Efficient Draw": "Maintain current draw strategy and continue routine monitoring.",
        "Brittle but Efficient": "Continue draw, but monitor seismic response closely due to elevated E/P.",
        "Ductile but Costly": "Review draw distribution and tonnes efficiency because seismic cost per tonne is elevated.",
        "High-Risk Draw": "Escalate for geotechnical review. Monitor draw, seismicity and nearby workplaces closely.",
        "Waiting for tonnes": "Enter tonnes drawn to finalise the draw compliance classification.",
    }.get(status, "")


def plot_matrix(ep, et, history, ep_threshold, et_threshold):
    max_x = max(et_threshold * 2, et * 1.4)
    max_y = max(ep_threshold * 2, ep * 1.4)

    fig = go.Figure()

    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=0, y1=ep_threshold,
                  fillcolor="green", opacity=0.12, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=0, y1=ep_threshold,
                  fillcolor="orange", opacity=0.14, line_width=0)
    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=ep_threshold, y1=max_y,
                  fillcolor="gold", opacity=0.16, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=ep_threshold, y1=max_y,
                  fillcolor="red", opacity=0.12, line_width=0)

    fig.add_shape(type="line", x0=et_threshold, x1=et_threshold, y0=0, y1=max_y,
                  line=dict(dash="dash", width=2))
    fig.add_shape(type="line", x0=0, x1=max_x, y0=ep_threshold, y1=ep_threshold,
                  line=dict(dash="dash", width=2))

    if not history.empty and "E/P" in history.columns and "E/T" in history.columns:
        hist = history.copy()
        hist["E/P"] = pd.to_numeric(hist["E/P"], errors="coerce")
        hist["E/T"] = pd.to_numeric(hist["E/T"], errors="coerce")
        hist = hist.dropna(subset=["E/P", "E/T"])

        if not hist.empty:
            fig.add_trace(go.Scatter(
                x=hist["E/T"],
                y=hist["E/P"],
                mode="markers",
                name="Saved Records",
                text=hist["Report Date"] if "Report Date" in hist.columns else None,
                hovertemplate="Date: %{text}<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
            ))

    fig.add_trace(go.Scatter(
        x=[et],
        y=[ep],
        mode="markers+text",
        name="Current",
        text=["Current"],
        textposition="top center",
        marker=dict(size=22),
        hovertemplate="Current<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
    ))

    fig.add_annotation(x=et_threshold * 0.45, y=ep_threshold * 0.5, text="Efficient Draw", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 0.5, text="Ductile but Costly", showarrow=False)
    fig.add_annotation(x=et_threshold * 0.5, y=ep_threshold * 1.45, text="Brittle but Efficient", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 1.45, text="High-Risk Draw", showarrow=False)

    fig.update_layout(
        title="Draw Compliance Matrix",
        xaxis_title="Energy / Tonnes (E/T)",
        yaxis_title="Energy / Potency (E/P)",
        xaxis=dict(range=[0, max_x]),
        yaxis=dict(range=[0, max_y]),
        height=650,
    )

    return fig



def status_colour(status):
    return {
        "Efficient Draw": "green",
        "Brittle but Efficient": "orange",
        "Ductile but Costly": "orange",
        "High-Risk Draw": "red",
        "Waiting for tonnes": "lightgrey",
        "Waiting for seismic data": "lightgrey",
    }.get(status, "lightgrey")


def normalise_drawpoint(value):
    return re.sub(r"\s+", "", str(value).strip().upper())


def infer_drawbell(drawpoint):
    """Return the physical drawbell for a Lift II drawpoint.

    Examples:
        T2E01  + T1W01  -> A01
        T3E01  + T2W01  -> B01
        T4E01  + T3W01  -> C01

    For an east drawpoint TnE##, the bell letter is based on n-2.
    For a west drawpoint TnW##, the bell letter is based on n-1.
    """
    value = normalise_drawpoint(drawpoint)
    match = re.fullmatch(r"T(\d+)([EW])(\d{2})", value)
    if not match:
        return ""

    tunnel = int(match.group(1))
    side = match.group(2)
    bell_number = match.group(3)

    letter_index = tunnel - 2 if side == "E" else tunnel - 1
    if not 0 <= letter_index < 26:
        return ""

    return f"{chr(ord('A') + letter_index)}{bell_number}"


def read_table(uploaded_file):
    if uploaded_file is None:
        return pd.DataFrame()
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file)
    return pd.read_excel(uploaded_file)


def build_mapping_template(drawpoints):
    unique = sorted({normalise_drawpoint(v) for v in drawpoints if str(v).strip()})
    return pd.DataFrame({
        "Drawpoint": unique,
        "Drawbell": [infer_drawbell(v) for v in unique],
        "Polygon": [""] * len(unique),
    })


def build_drawbell_mapping(template, saved_mapping):
    """Create one editable polygon assignment per drawbell and propagate it to both drawpoints."""
    bell_map = template[["Drawbell"]].drop_duplicates().copy()
    bell_map = bell_map[bell_map["Drawbell"].astype(str).str.strip() != ""]
    bell_map["Polygon"] = ""

    if not saved_mapping.empty:
        saved = saved_mapping.copy()
        if "Drawbell" not in saved.columns and "Drawpoint" in saved.columns:
            saved["Drawbell"] = saved["Drawpoint"].map(infer_drawbell)
        saved["Drawbell"] = saved["Drawbell"].astype(str).str.strip().str.upper()
        saved["Polygon"] = saved.get("Polygon", "").astype(str).str.strip().str.upper()
        saved = saved[saved["Polygon"].isin(list("ABCDEF"))]
        saved = saved.drop_duplicates(subset=["Drawbell"], keep="last")[["Drawbell", "Polygon"]]
        bell_map = bell_map.drop(columns=["Polygon"]).merge(saved, on="Drawbell", how="left")
        bell_map["Polygon"] = bell_map["Polygon"].fillna("")

    return bell_map.sort_values("Drawbell", kind="stable").reset_index(drop=True)


def expand_drawbell_mapping(template, bell_mapping):
    """Attach the drawbell polygon assignment to every drawpoint in that drawbell."""
    clean = bell_mapping[["Drawbell", "Polygon"]].copy()
    clean["Drawbell"] = clean["Drawbell"].astype(str).str.strip().str.upper()
    clean["Polygon"] = clean["Polygon"].astype(str).str.strip().str.upper()
    return template.drop(columns=["Polygon"]).merge(clean, on="Drawbell", how="left")


def aggregate_polygon_tonnes(tonnes_df, drawpoint_col, tonnes_col, mapping_df):
    work = tonnes_df[[drawpoint_col, tonnes_col]].copy()
    work.columns = ["Drawpoint", "Tonnes"]
    work["Drawpoint"] = work["Drawpoint"].map(normalise_drawpoint)
    work["Tonnes"] = pd.to_numeric(work["Tonnes"], errors="coerce").fillna(0.0)

    mapping = mapping_df.copy()
    mapping["Drawpoint"] = mapping["Drawpoint"].map(normalise_drawpoint)
    mapping["Polygon"] = mapping["Polygon"].astype(str).str.strip().str.upper()
    mapping["Drawbell"] = mapping["Drawbell"].astype(str).str.strip().str.upper()

    merged = work.merge(mapping, on="Drawpoint", how="left")
    drawbell_totals = (
        merged.dropna(subset=["Polygon"])
        .groupby(["Polygon", "Drawbell"], as_index=False)["Tonnes"].sum()
    )
    polygon_totals = (
        merged.dropna(subset=["Polygon"])
        .groupby("Polygon", as_index=False)["Tonnes"].sum()
        .rename(columns={"Tonnes": "Polygon Tonnes"})
    )
    return merged, drawbell_totals, polygon_totals


def polygon_health_score(ep, et, events, potency, tonnes, target_tonnes, ep_threshold, et_threshold):
    """Return a transparent 0-100 polygon health score."""
    if potency <= 0 or tonnes <= 0 or et is None:
        return 0

    ep_ratio = ep / ep_threshold if ep_threshold else 0
    et_ratio = et / et_threshold if et_threshold else 0
    matrix_score = max(0, 60 - min(60, 25 * ep_ratio + 25 * et_ratio))

    production_score = 20
    if target_tonnes > 0:
        production_score = min(20, max(0, 20 * tonnes / target_tonnes))

    event_score = 15
    if events >= 100:
        event_score = 3
    elif events >= 60:
        event_score = 7
    elif events >= 30:
        event_score = 11

    completeness_score = 5
    return int(round(max(0, min(100, matrix_score + production_score + event_score + completeness_score))))


def polygon_health_rating(score):
    if score <= 0:
        return "No data"
    if score >= 80:
        return "Healthy"
    if score >= 60:
        return "Watch"
    if score >= 40:
        return "Concern"
    return "Critical"


def polygon_trend(history_df, polygon, current_health):
    """Compare the current health score with the most recent saved score."""
    if history_df.empty or "Polygon" not in history_df.columns or "Health Score" not in history_df.columns:
        return "No history", 0.0, "→"

    hist = history_df[history_df["Polygon"].astype(str).str.upper() == str(polygon).upper()].copy()
    if hist.empty:
        return "No history", 0.0, "→"

    if "Report Date" in hist.columns:
        hist["_date"] = pd.to_datetime(hist["Report Date"], errors="coerce")
        hist = hist.sort_values("_date", kind="stable")

    previous_scores = pd.to_numeric(hist["Health Score"], errors="coerce").dropna()
    if previous_scores.empty:
        return "No history", 0.0, "→"

    previous = float(previous_scores.iloc[-1])
    change = float(current_health) - previous
    if change >= 5:
        return "Improving", change, "↑"
    if change <= -5:
        return "Deteriorating", change, "↓"
    return "Stable", change, "→"


def polygon_explanation(row, ep_threshold, et_threshold):
    """Plain-language proof explaining why a polygon received its colour."""
    status = str(row.get("Status", "Waiting for data"))
    ep = to_float(row.get("E/P", 0.0))
    et_value = row.get("E/T", None)
    et = None if et_value in [None, ""] or pd.isna(et_value) else to_float(et_value)

    if status == "Waiting for seismic data":
        return "The polygon is grey because energy and potency have not yet been entered."
    if et is None:
        return "The polygon is grey because tonnes are still missing, so E/T cannot be calculated."
    if status == "Efficient Draw":
        return f"Green because E/P ({ep:,.2f}) is below {ep_threshold:,.2f} and E/T ({et:,.2f}) is below {et_threshold:,.2f}."
    if status == "Brittle but Efficient":
        return f"Amber because E/P ({ep:,.2f}) exceeds {ep_threshold:,.2f}, while E/T ({et:,.2f}) remains below {et_threshold:,.2f}."
    if status == "Ductile but Costly":
        return f"Amber because E/P ({ep:,.2f}) is below {ep_threshold:,.2f}, but E/T ({et:,.2f}) exceeds {et_threshold:,.2f}."
    return f"Red because both E/P ({ep:,.2f}) and E/T ({et:,.2f}) exceed their thresholds."


def polygon_alerts(row, trend_label, trend_change, ep_threshold, et_threshold):
    alerts = []
    status = str(row.get("Status", ""))
    health = to_float(row.get("Health Score", 0.0))
    achievement = row.get("Achievement %", None)

    if status == "High-Risk Draw":
        alerts.append("High-risk matrix position: both E/P and E/T are above threshold.")
    elif status == "Brittle but Efficient":
        alerts.append("Elevated E/P: monitor the seismic response closely.")
    elif status == "Ductile but Costly":
        alerts.append("Elevated E/T: review draw distribution and tonnes efficiency.")

    if health > 0 and health < 40:
        alerts.append("Health score is critical and requires geotechnical review.")
    if trend_label == "Deteriorating":
        alerts.append(f"Health score deteriorated by {abs(trend_change):.0f} points from the latest saved record.")
    if achievement is not None and not pd.isna(achievement) and achievement < 80:
        alerts.append(f"Production achievement is only {achievement:.1f}% of target.")

    return alerts


def selected_polygon_report(row, trend_label, trend_change, explanation, alerts):
    poly = str(row.get("Polygon", ""))
    report_data = {
        "Polygon": poly,
        "Status": row.get("Status", ""),
        "Health Score": f"{to_int(row.get('Health Score', 0))}/100",
        "Health Rating": row.get("Health Rating", ""),
        "Events": to_int(row.get("Events", 0)),
        "Energy J": f"{to_float(row.get('Energy J', 0)):,.0f}",
        "Potency m3": f"{to_float(row.get('Potency m3', 0)):.3f}",
        "Polygon Tonnes": f"{to_float(row.get('Polygon Tonnes', 0)):,.0f}",
        "E/P": f"{to_float(row.get('E/P', 0)):,.2f}",
        "E/T": "Waiting" if row.get("E/T", None) in [None, ""] or pd.isna(row.get("E/T", None)) else f"{to_float(row.get('E/T')):,.2f}",
        "Trend": f"{trend_label} ({trend_change:+.0f})",
        "Why this colour": explanation,
        "Alerts": "; ".join(alerts) if alerts else "No active alerts",
    }
    summary = (
        f"Polygon {poly} is classified as {row.get('Status', '')}. {explanation} "
        f"The health trend is {trend_label.lower()}. "
        + ("Active alerts: " + "; ".join(alerts) if alerts else "No active alerts are currently identified.")
    )
    return build_word_report(f"GDI Polygon {poly} Decision Support Report", summary, report_data)


def polygon_data_quality(results_df, merged_tonnes, mapping_df):
    """Return transparent data-quality checks for the current polygon dataset."""
    issues = []
    if results_df.empty:
        return ["No polygon results are available."]

    for _, row in results_df.iterrows():
        poly = str(row.get("Polygon", "?"))
        energy = to_float(row.get("Energy J", 0.0))
        potency = to_float(row.get("Potency m3", 0.0))
        tonnes = to_float(row.get("Polygon Tonnes", 0.0))
        events = to_int(row.get("Events", 0))

        if energy > 0 and potency <= 0:
            issues.append(f"Polygon {poly}: energy is entered but potency is zero.")
        if potency > 0 and tonnes <= 0:
            issues.append(f"Polygon {poly}: seismic data exists but tonnes are missing.")
        if events > 0 and energy <= 0:
            issues.append(f"Polygon {poly}: events are entered but total energy is zero.")
        if energy < 0 or potency < 0 or tonnes < 0 or events < 0:
            issues.append(f"Polygon {poly}: negative values were detected.")

    if not merged_tonnes.empty:
        if "Polygon" in merged_tonnes.columns:
            missing = merged_tonnes[
                merged_tonnes["Polygon"].isna()
                | (merged_tonnes["Polygon"].astype(str).str.strip() == "")
            ]
            if not missing.empty:
                issues.append(f"{missing['Drawpoint'].nunique()} drawpoint(s) are not assigned to a polygon.")
        if "Drawpoint" in merged_tonnes.columns:
            duplicate_count = int(merged_tonnes["Drawpoint"].astype(str).duplicated().sum())
            if duplicate_count:
                issues.append(f"{duplicate_count} duplicate drawpoint row(s) were detected in the production data.")

    if mapping_df.empty:
        issues.append("The master drawpoint mapping has not been saved yet.")
    elif "Polygon" in mapping_df.columns:
        bad_mapping = mapping_df[~mapping_df["Polygon"].astype(str).str.upper().isin(list("ABCDEF"))]
        if not bad_mapping.empty:
            issues.append(f"{len(bad_mapping)} mapping row(s) have an invalid or blank polygon.")

    return issues


def prepare_history_snapshot(history_df, report_date, ep_threshold, et_threshold):
    """Prepare one saved date for map replay, including any missing derived columns."""
    if history_df.empty or "Report Date" not in history_df.columns:
        return pd.DataFrame()
    snapshot = history_df[history_df["Report Date"].astype(str) == str(report_date)].copy()
    if snapshot.empty:
        return snapshot

    for col in ["Energy J", "Potency m3", "Polygon Tonnes", "Events", "Target Tonnes"]:
        if col not in snapshot.columns:
            snapshot[col] = 0.0
        snapshot[col] = pd.to_numeric(snapshot[col], errors="coerce").fillna(0.0)

    snapshot["E/P"] = snapshot.apply(
        lambda r: r["Energy J"] / r["Potency m3"] if r["Potency m3"] > 0 else 0.0, axis=1
    )
    snapshot["E/T"] = snapshot.apply(
        lambda r: r["Energy J"] / r["Polygon Tonnes"] if r["Polygon Tonnes"] > 0 else None, axis=1
    )
    snapshot["Status"] = snapshot.apply(
        lambda r: get_status(r["E/P"], r["E/T"], ep_threshold, et_threshold)
        if r["Potency m3"] > 0 else "Waiting for seismic data", axis=1
    )
    snapshot["Map Rating"] = snapshot["Status"].map(
        lambda s: "Green" if s == "Efficient Draw" else "Red" if s == "High-Risk Draw"
        else "Amber" if s in ["Brittle but Efficient", "Ductile but Costly"] else "Grey"
    )
    snapshot["Health Score"] = snapshot.apply(
        lambda r: polygon_health_score(
            r["E/P"], r["E/T"], r["Events"], r["Potency m3"],
            r["Polygon Tonnes"], r["Target Tonnes"], ep_threshold, et_threshold
        ), axis=1
    )
    snapshot["Health Rating"] = snapshot["Health Score"].map(polygon_health_rating)
    snapshot["Achievement %"] = snapshot.apply(
        lambda r: 100 * r["Polygon Tonnes"] / r["Target Tonnes"] if r["Target Tonnes"] > 0 else None,
        axis=1,
    )
    return snapshot


def executive_summary_text(results_df):
    """Create a concise management overview from the six polygon results."""
    if results_df.empty:
        return "No polygon results are available."
    work = results_df.copy()
    work["Health Score"] = pd.to_numeric(work.get("Health Score", 0), errors="coerce").fillna(0)
    work["Polygon Tonnes"] = pd.to_numeric(work.get("Polygon Tonnes", 0), errors="coerce").fillna(0)
    best = work.sort_values("Health Score", ascending=False).iloc[0]
    risk = work.sort_values("Health Score", ascending=True).iloc[0]
    total_tonnes = work["Polygon Tonnes"].sum()
    red_count = int((work.get("Map Rating", "") == "Red").sum())
    amber_count = int((work.get("Map Rating", "") == "Amber").sum())
    return (
        f"Lift II East recorded {total_tonnes:,.0f} t across the six polygons. "
        f"Polygon {best['Polygon']} is the strongest current performer with a health score of "
        f"{to_int(best['Health Score'])}/100. Polygon {risk['Polygon']} requires the most attention "
        f"with a health score of {to_int(risk['Health Score'])}/100. "
        f"The dashboard currently shows {red_count} poor and {amber_count} watch polygon(s)."
    )


def build_dashboard_report(report_date, results_df, quality_issues, executive_text):
    """Generate a management Word report covering all polygons."""
    doc = Document()
    doc.add_heading("GDI Lift II East Management Report", 0)
    doc.add_paragraph(f"Reporting date: {report_date}")
    doc.add_heading("Executive Overview", level=1)
    doc.add_paragraph(executive_text)

    doc.add_heading("Data Quality", level=1)
    if quality_issues:
        for issue in quality_issues:
            doc.add_paragraph(issue, style="List Bullet")
    else:
        doc.add_paragraph("All core data-quality checks passed.")

    doc.add_heading("Polygon Summary", level=1)
    columns = ["Polygon", "Status", "Health Score", "Polygon Tonnes", "E/P", "E/T", "Events"]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in results_df.iterrows():
        cells = table.add_row().cells
        values = [
            row.get("Polygon", ""), row.get("Status", ""), to_int(row.get("Health Score", 0)),
            f"{to_float(row.get('Polygon Tonnes', 0)):,.0f}", f"{to_float(row.get('E/P', 0)):,.2f}",
            "Waiting" if row.get("E/T", None) in [None, ""] or pd.isna(row.get("E/T", None))
            else f"{to_float(row.get('E/T')):,.2f}", to_int(row.get("Events", 0)),
        ]
        for i, value in enumerate(values):
            cells[i].text = str(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




def load_saved_settings():
    defaults = {"E/P Threshold": DEFAULT_EP_THRESHOLD, "E/T Threshold": DEFAULT_ET_THRESHOLD, "Health Drop Alert": 15.0}
    df = load_csv(SETTINGS_FILE)
    if df.empty or "Setting" not in df.columns or "Value" not in df.columns:
        return defaults
    for _, row in df.iterrows():
        key = str(row.get("Setting", ""))
        if key in defaults:
            defaults[key] = to_float(row.get("Value"), defaults[key])
    return defaults


def save_settings(ep_threshold, et_threshold, health_drop_alert):
    save_csv(pd.DataFrame({
        "Setting": ["E/P Threshold", "E/T Threshold", "Health Drop Alert"],
        "Value": [ep_threshold, et_threshold, health_drop_alert],
    }), SETTINGS_FILE)


def load_manual_tonnes_for_date(report_date, mapping_df):
    base = mapping_df[["Drawpoint", "Drawbell", "Polygon"]].copy()
    base["Tonnes"] = 0.0
    saved = load_csv(MANUAL_TONNES_FILE)
    if saved.empty or "Report Date" not in saved.columns:
        return base
    day = saved[saved["Report Date"].astype(str) == str(report_date)].copy()
    if day.empty:
        return base
    day["Drawpoint"] = day["Drawpoint"].map(normalise_drawpoint)
    base["Drawpoint"] = base["Drawpoint"].map(normalise_drawpoint)
    base = base.drop(columns=["Tonnes"]).merge(day[["Drawpoint", "Tonnes"]], on="Drawpoint", how="left")
    base["Tonnes"] = pd.to_numeric(base["Tonnes"], errors="coerce").fillna(0.0)
    return base


def save_manual_tonnes(report_date, manual_df):
    save_df = manual_df[["Drawpoint", "Drawbell", "Polygon", "Tonnes"]].copy()
    save_df.insert(0, "Report Date", str(report_date))
    existing = load_csv(MANUAL_TONNES_FILE)
    combined = pd.concat([existing, save_df], ignore_index=True) if not existing.empty else save_df
    combined = combined.drop_duplicates(subset=["Report Date", "Drawpoint"], keep="last")
    save_csv(combined, MANUAL_TONNES_FILE)


def polygon_trend_figure(history_df, polygon):
    if history_df.empty or "Polygon" not in history_df.columns or "Report Date" not in history_df.columns:
        return None
    hist = history_df[history_df["Polygon"].astype(str).str.upper() == str(polygon).upper()].copy()
    if hist.empty:
        return None
    hist["Report Date"] = pd.to_datetime(hist["Report Date"], errors="coerce")
    hist = hist.dropna(subset=["Report Date"]).sort_values("Report Date").tail(30)
    for col in ["E/P", "E/T", "Polygon Tonnes", "Health Score"]:
        if col in hist.columns:
            hist[col] = pd.to_numeric(hist[col], errors="coerce")
    available = [c for c in ["Health Score", "Polygon Tonnes", "E/P", "E/T"] if c in hist.columns]
    if not available:
        return None
    metric = st.selectbox("Trend metric", available, key=f"trend_metric_{polygon}")
    fig = go.Figure(go.Scatter(x=hist["Report Date"], y=hist[metric], mode="lines+markers", name=metric))
    fig.update_layout(height=310, margin=dict(l=10, r=10, t=35, b=10), title=f"Polygon {polygon} — {metric} trend", showlegend=False)
    return fig


def build_printable_html_report(report_date, results_df, executive_text, quality_issues):
    rows = []
    for _, r in results_df.iterrows():
        et = r.get("E/T", None)
        et_text = "Waiting" if et is None or pd.isna(et) else f"{to_float(et):,.2f}"
        rows.append(
            "<tr>"
            f"<td>{r.get('Polygon','')}</td>"
            f"<td>{r.get('Status','')}</td>"
            f"<td>{to_int(r.get('Health Score',0))}</td>"
            f"<td>{to_float(r.get('Polygon Tonnes',0)):,.0f}</td>"
            f"<td>{to_float(r.get('E/P',0)):,.2f}</td>"
            f"<td>{et_text}</td>"
            f"<td>{to_int(r.get('Events',0))}</td>"
            "</tr>"
        )
    issues = "<li>All core data-quality checks passed.</li>" if not quality_issues else "".join(f"<li>{i}</li>" for i in quality_issues)
    html = (
        '<!doctype html><html><head><meta charset="utf-8"><title>GDI Report</title>'
        '<style>body{font-family:Arial;margin:28px;color:#172033}h1{color:#082b5c}'
        'table{border-collapse:collapse;width:100%}th,td{border:1px solid #ccd5e0;padding:7px;text-align:left}'
        'th{background:#082b5c;color:white}.note{background:#f4f7fb;padding:12px;border-left:5px solid #0b3d78}</style>'
        f'</head><body><h1>GDI — Lift II East Management Report</h1><p><b>Reporting date:</b> {report_date}</p>'
        f'<div class="note">{executive_text}</div><h2>Data Quality</h2><ul>{issues}</ul>'
        '<h2>Polygon Summary</h2><table><tr><th>Polygon</th><th>Status</th><th>Health</th><th>Tonnes</th><th>E/P</th><th>E/T</th><th>Events</th></tr>'
        + ''.join(rows) +
        '</table><p>Open this file in a browser and use Print → Save as PDF.</p></body></html>'
    )
    return html.encode("utf-8")




def auto_detect_production_columns(df):
    """Guess drawpoint and tonnes columns from common production spreadsheet headings."""
    if df.empty:
        return None, None
    names = {str(c).strip().lower(): c for c in df.columns}
    draw_candidates = ["drawpoint", "draw point", "dp", "draw_point", "location", "drawpoint name"]
    tonne_candidates = ["tonnes", "tons", "tonnage", "actual tonnes", "actual tons", "mass"]
    draw_col = next((names[n] for n in draw_candidates if n in names), None)
    tonnes_col = next((names[n] for n in tonne_candidates if n in names), None)
    if draw_col is None:
        for c in df.columns:
            sample = df[c].astype(str).head(50)
            if sample.str.match(r"^T\d+[EW]\d{2}$", case=False).mean() > 0.35:
                draw_col = c
                break
    if tonnes_col is None:
        numeric = [c for c in df.columns if pd.to_numeric(df[c], errors="coerce").notna().mean() > 0.7]
        tonnes_col = numeric[0] if numeric else None
    return draw_col, tonnes_col


def normalise_polygon_name(value):
    value = str(value).strip().upper()
    match = re.search(r"(?:POLY(?:GON)?[_\s-]*)?([A-F])$", value)
    return match.group(1) if match else ""


def parse_polygon_ims_pdf(uploaded_pdf):
    """Read one IMS PDF and return A-F polygon events, energy and potency where available."""
    if uploaded_pdf is None:
        return pd.DataFrame(columns=["Polygon", "Events", "Energy J", "Potency m3"])
    text = pdf_to_text(uploaded_pdf)
    parsed = parse_poly_sections(text)
    rows = []
    for key, data in parsed.items():
        poly = normalise_polygon_name(key)
        if poly:
            rows.append({
                "Polygon": poly,
                "Events": to_int(data.get("Events", 0)),
                "Energy J": to_float(data.get("Energy J", 0)),
                "Potency m3": to_float(data.get("Potency m3", 0)),
            })
    return pd.DataFrame(rows).drop_duplicates("Polygon", keep="last") if rows else pd.DataFrame(columns=["Polygon", "Events", "Energy J", "Potency m3"])


def write_audit(user, role, action, details=""):
    record = pd.DataFrame([{
        "Timestamp": datetime.now().isoformat(timespec="seconds"),
        "User": user or "Unknown", "Role": role or "Unknown",
        "Action": action, "Details": details,
    }])
    existing = load_csv(AUDIT_FILE)
    save_csv(pd.concat([existing, record], ignore_index=True) if not existing.empty else record, AUDIT_FILE)


def get_approval_status(report_date):
    df = load_csv(APPROVAL_FILE)
    if df.empty or "Report Date" not in df.columns:
        return "Draft", "", ""
    rows = df[df["Report Date"].astype(str) == str(report_date)]
    if rows.empty:
        return "Draft", "", ""
    row = rows.iloc[-1]
    return str(row.get("Status", "Draft")), str(row.get("User", "")), str(row.get("Timestamp", ""))


def save_approval_status(report_date, status, user, role):
    record = {"Report Date": str(report_date), "Status": status, "User": user, "Role": role, "Timestamp": datetime.now().isoformat(timespec="seconds")}
    save_or_update(APPROVAL_FILE, record, ["Report Date"])
    write_audit(user, role, f"Set workflow status to {status}", f"Report date {report_date}")


def convergence_risk_label(closure_mm, rate_mm_day, status):
    if closure_mm <= 0 and rate_mm_day <= 0:
        return "No convergence data"
    high_conv = closure_mm >= 15 or rate_mm_day >= 2
    watch_conv = closure_mm >= 8 or rate_mm_day >= 1
    if high_conv and status in ["High-Risk Draw", "Brittle but Efficient"]:
        return "Combined geotechnical concern"
    if high_conv:
        return "High convergence"
    if watch_conv:
        return "Convergence watch"
    return "Stable convergence"


def load_risk_engine_settings():
    """Load enabled components and user-defined weights for the risk engine."""
    defaults = {
        "Draw Response": {"enabled": True, "weight": 30.0},
        "Health Score": {"enabled": True, "weight": 20.0},
        "Production Achievement": {"enabled": True, "weight": 20.0},
        "Seismic Activity": {"enabled": True, "weight": 30.0},
        "Convergence": {"enabled": False, "weight": 0.0},
    }
    df = load_csv(RISK_ENGINE_FILE)
    if df.empty or not {"Component", "Enabled", "Weight"}.issubset(df.columns):
        return defaults
    for _, r in df.iterrows():
        component = str(r.get("Component", ""))
        if component in defaults:
            enabled_value = str(r.get("Enabled", True)).strip().lower()
            defaults[component]["enabled"] = enabled_value in ["true", "1", "yes", "on"]
            defaults[component]["weight"] = max(0.0, to_float(r.get("Weight", 0.0)))
    return defaults


def save_risk_engine_settings(config):
    rows = [
        {"Component": name, "Enabled": values["enabled"], "Weight": values["weight"]}
        for name, values in config.items()
    ]
    save_csv(pd.DataFrame(rows), RISK_ENGINE_FILE)


def risk_component_scores(row):
    """Return each component as an independent 0-100 risk score."""
    status = str(row.get("Status", ""))
    draw_risk = {
        "Efficient Draw": 5,
        "Brittle but Efficient": 55,
        "Ductile but Costly": 60,
        "High-Risk Draw": 95,
        "Waiting for tonnes": 50,
        "Waiting for seismic data": 45,
    }.get(status, 45)

    health = to_float(row.get("Health Score", 0))
    health_risk = 100 - max(0, min(100, health)) if health > 0 else 50

    achievement = row.get("Achievement %", None)
    if achievement is None or pd.isna(achievement):
        production_risk = 50
    else:
        ach = to_float(achievement)
        if ach < 50:
            production_risk = 95
        elif ach < 70:
            production_risk = 75
        elif ach < 85:
            production_risk = 50
        elif ach <= 120:
            production_risk = 10
        elif ach <= 140:
            production_risk = 45
        else:
            production_risk = 75

    events = to_int(row.get("Events", 0))
    energy = to_float(row.get("Energy J", 0))
    if events >= 100:
        seismic_risk = 90
    elif events >= 60:
        seismic_risk = 70
    elif events >= 30:
        seismic_risk = 50
    elif events >= 10:
        seismic_risk = 25
    else:
        seismic_risk = 10
    if energy <= 0 and events <= 0:
        seismic_risk = 35

    conv = str(row.get("Convergence Risk", "No convergence data"))
    convergence_risk = {
        "Combined geotechnical concern": 100,
        "High convergence": 80,
        "Convergence watch": 50,
        "Stable convergence": 5,
        "No convergence data": 50,
    }.get(conv, 50)

    return {
        "Draw Response": draw_risk,
        "Health Score": health_risk,
        "Production Achievement": production_risk,
        "Seismic Activity": seismic_risk,
        "Convergence": convergence_risk,
    }


def combined_risk_index(row, config):
    """Weighted 0-100 risk index using only components switched on by the user."""
    scores = risk_component_scores(row)
    enabled = {
        name: values for name, values in config.items()
        if values.get("enabled", False) and to_float(values.get("weight", 0)) > 0
    }
    if not enabled:
        return 0
    total_weight = sum(to_float(v.get("weight", 0)) for v in enabled.values())
    if total_weight <= 0:
        return 0
    weighted = sum(scores[name] * to_float(values.get("weight", 0)) for name, values in enabled.items())
    return int(round(max(0, min(100, weighted / total_weight))))


def risk_inputs_used(config):
    active = [name for name, values in config.items() if values.get("enabled") and to_float(values.get("weight", 0)) > 0]
    return ", ".join(active) if active else "No components enabled"


def combined_risk_label(score):
    if score >= 75:
        return "Critical"
    if score >= 50:
        return "High"
    if score >= 25:
        return "Moderate"
    return "Low"


def neighbouring_polygon_assessment(results_df):
    """Assess influence between polygons sharing an edge in the 2-column by 3-row layout."""
    adjacency = {
        "A": ["D", "B"], "D": ["A", "E"],
        "B": ["A", "E", "C"], "E": ["D", "B", "F"],
        "C": ["B", "F"], "F": ["E", "C"],
    }
    risk_map = {
        str(r["Polygon"]): to_int(r.get("Combined Risk Index", 0))
        for _, r in results_df.iterrows()
    }
    out = {}
    for poly, neighbours in adjacency.items():
        active = [n for n in neighbours if risk_map.get(n, 0) >= 50]
        if active:
            out[poly] = f"Neighbour influence: monitor {', '.join(active)}"
        else:
            out[poly] = "No elevated neighbouring influence"
    return out


def decision_recommendation(row):
    risk = str(row.get("Combined Risk", "Low"))
    status = str(row.get("Status", ""))
    conv = str(row.get("Convergence Risk", ""))
    achievement = row.get("Achievement %", None)
    actions = []
    if risk in ["Critical", "High"]:
        actions.append("Escalate for geotechnical and production review")
    if status in ["High-Risk Draw", "Brittle but Efficient"]:
        actions.append("review seismic response and nearby draw distribution")
    if status == "Ductile but Costly":
        actions.append("review energy cost per tonne and draw balance")
    if conv in ["Combined geotechnical concern", "High convergence"]:
        actions.append("increase convergence monitoring frequency")
    if achievement is not None and not pd.isna(achievement):
        ach = to_float(achievement)
        if ach < 75:
            actions.append("investigate underdraw and inactive drawpoints")
        elif ach > 130:
            actions.append("check for possible overdraw")
    if not actions:
        actions.append("maintain current draw strategy and routine monitoring")
    return "; ".join(actions).capitalize() + "."


def save_handover(report_date, shift, observation, action, owner, status, user):
    record = {
        "Report Date": str(report_date), "Shift": shift,
        "Observation": observation, "Action": action,
        "Responsible Person": owner, "Item Status": status,
        "Entered By": user, "Timestamp": datetime.now().isoformat(timespec="seconds"),
    }
    existing = load_csv(HANDOVER_FILE)
    combined = pd.concat([existing, pd.DataFrame([record])], ignore_index=True) if not existing.empty else pd.DataFrame([record])
    save_csv(combined, HANDOVER_FILE)


def polygon_map_figure(results_df, map_image, selected_polygon="A"):
    """Status map using polygon paths aligned to the actual red boundaries."""
    image = Image.open(map_image) if not isinstance(map_image, Image.Image) else map_image
    width, height = image.size

    fig = go.Figure()
    fig.add_layout_image(dict(
        source=image, x=0, y=0, sizex=width, sizey=height,
        xref="x", yref="y", sizing="stretch", layer="below"
    ))

    # Pixel coordinates traced from the actual Lift II East polygon boundaries.
    # Exact pixel boundaries traced from Draw_Polygons.pdf on the unchanged
    # 1600 x 1132 clean map. Do not crop, resize, or offset the background.
    polygon_points = {
        "D": [(709, 504), (869, 504), (869, 669), (709, 669)],
        "A": [(869, 504), (1030, 504), (1030, 669), (869, 669)],
        "E": [(709, 669), (869, 669), (869, 828), (709, 828)],
        "B": [(869, 669), (1030, 669), (1030, 828), (869, 828)],
        "F": [(709, 828), (869, 828), (869, 982), (709, 982)],
        "C": [(869, 828), (1030, 828), (1030, 982), (869, 982)],
    }

    rows = {
        str(r["Polygon"]).upper(): r
        for _, r in results_df.iterrows()
    } if not results_df.empty else {}

    for poly, points in polygon_points.items():
        row = rows.get(poly, {})
        status = row.get("Status", "Waiting for seismic data") if hasattr(row, "get") else "Waiting for seismic data"
        colour = status_colour(status)
        selected = poly == selected_polygon

        path = "M " + " L ".join(f"{x},{y}" for x, y in points) + " Z"
        # Keep every polygon visibly shaded while preserving the mine drawing below.
        # The selected polygon gets a stronger fill plus a dark-blue focus outline.
        fig.add_shape(
            type="path",
            path=path,
            fillcolor=colour,
            opacity=0.30 if not selected else 0.46,
            line=dict(color=colour, width=4 if not selected else 5),
            layer="above",
        )

        if selected:
            # Separate outline creates a clean selected-polygon highlight without
            # changing the status colour of the polygon itself.
            fig.add_shape(
                type="path",
                path=path,
                fillcolor="rgba(0,0,0,0)",
                opacity=1.0,
                line=dict(color="#0b3d78", width=8),
                layer="above",
            )

        centre_x = sum(x for x, _ in points) / len(points)
        centre_y = sum(y for _, y in points) / len(points)
        hover = f"<b>Polygon {poly}</b><br>Status: {status}"
        if hasattr(row, "get"):
            hover += (
                f"<br>Tonnes: {to_float(row.get('Polygon Tonnes', 0)):,.0f} t"
                f"<br>Energy: {to_float(row.get('Energy J', 0)):,.0f} J"
                f"<br>Potency: {to_float(row.get('Potency m3', 0)):,.3f} m³"
                f"<br>E/P: {to_float(row.get('E/P', 0)):,.2f}"
            )
            etv = row.get("E/T", None)
            hover += f"<br>E/T: {to_float(etv):,.2f}" if etv not in [None, ""] and not pd.isna(etv) else "<br>E/T: Waiting"

        fig.add_trace(go.Scatter(
            x=[centre_x], y=[centre_y], mode="markers",
            marker=dict(size=92, opacity=0.01, color="rgba(0,0,0,0.01)"),
            customdata=[[poly]],
            hovertext=[hover], hoverinfo="text", showlegend=False,
            selected=dict(marker=dict(opacity=0.01)),
            unselected=dict(marker=dict(opacity=0.01)),
        ))

    fig.update_xaxes(range=[0, width], visible=False, fixedrange=True)
    fig.update_yaxes(range=[height, 0], visible=False, scaleanchor="x", fixedrange=True)
    fig.update_layout(
        height=700,
        margin=dict(l=0, r=0, t=0, b=0),
        title=None,
        showlegend=False,
        dragmode=False,
        hovermode="closest",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
    )
    return fig

def pie_chart(df, title):
    fig = px.pie(df, names="Zone", values="Events", title=title)
    fig.update_traces(textposition="inside", textinfo="percent+label")
    return fig


def bar_chart(df, title):
    fig = px.bar(df.sort_values("Events", ascending=True), x="Events", y="Zone", orientation="h", title=title)
    return fig


def build_word_report(title, comment, record):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("Summary Data", level=1)

    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Comment", level=1)
    doc.add_paragraph(comment)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.sidebar.header("User Session")
current_user = st.sidebar.text_input("User name", value=st.session_state.get("gdi_user", ""))
role_options = ["Graduate", "Strata Control Officer", "Geotechnical Specialist", "Geotechnical Superintendent", "Manager"]
saved_role = st.session_state.get("gdi_role", "Graduate")
if saved_role not in role_options:
    saved_role = "Graduate"
current_role = st.sidebar.selectbox("Role", role_options, index=role_options.index(saved_role))
st.session_state.gdi_user = current_user
st.session_state.gdi_role = current_role
st.sidebar.caption("Local prototype login; use company authentication before production deployment.")

st.sidebar.header("Settings")
_saved_settings = load_saved_settings()
ep_threshold = st.sidebar.number_input("E/P Threshold", min_value=0.0, value=float(_saved_settings["E/P Threshold"]))
et_threshold = st.sidebar.number_input("E/T Threshold", min_value=0.0, value=float(_saved_settings["E/T Threshold"]))
health_drop_alert = st.sidebar.number_input("Health drop alert", min_value=1.0, value=float(_saved_settings["Health Drop Alert"]), help="Alert when the health score drops by this many points.")
if st.sidebar.button("Save Settings", use_container_width=True):
    save_settings(ep_threshold, et_threshold, health_drop_alert)
    st.sidebar.success("Settings saved permanently.")

st.sidebar.header("Configurable Risk Engine")
risk_engine_config = load_risk_engine_settings()
with st.sidebar.expander("Index components and weights", expanded=False):
    st.caption("Switch components on or off. Enabled weights are automatically normalised to 100% during calculation.")
    edited_risk_config = {}
    for component, values in risk_engine_config.items():
        enabled = st.checkbox(f"Include {component}", value=bool(values["enabled"]), key=f"risk_enabled_{component}")
        weight = st.number_input(
            f"{component} weight",
            min_value=0.0,
            max_value=100.0,
            value=float(values["weight"]),
            step=5.0,
            disabled=not enabled,
            key=f"risk_weight_{component}",
        )
        edited_risk_config[component] = {"enabled": enabled, "weight": weight if enabled else 0.0}
    active_total = sum(v["weight"] for v in edited_risk_config.values() if v["enabled"])
    st.caption(f"Active raw weight total: {active_total:.0f}. The engine normalises this automatically.")
    if st.button("Save Risk Engine", use_container_width=True):
        save_risk_engine_settings(edited_risk_config)
        st.success("Risk engine configuration saved.")
        st.rerun()
risk_engine_config = edited_risk_config

st.markdown("""
<style>
:root {--navy:#082b5c;--blue:#0f4c8a;--border:#d8e0ea;--soft:#f6f9fc;}
.block-container {padding-top:0.35rem; padding-bottom:1.5rem; max-width:1900px;}
[data-testid="stSidebar"] {background:#f4f7fb; border-right:1px solid var(--border);}
[data-testid="stMetric"] {background:#ffffff; border:1px solid var(--border); padding:0.72rem 0.85rem; border-radius:8px; box-shadow:0 1px 2px rgba(15,23,42,.04);}
[data-testid="stMetricLabel"] {font-weight:700; color:#334155;}
[data-testid="stMetricValue"] {font-weight:800; color:#0f172a;}
.status-badge {display:inline-block;padding:.45rem .85rem;border-radius:999px;color:white;font-weight:800;margin:.15rem 0 .7rem 0;}
.gdi-header {background:linear-gradient(90deg,#062650,#0b3d78); color:white; padding:.65rem 1rem; border-radius:0 0 8px 8px; margin:-.35rem -1rem .65rem -1rem; display:flex; align-items:center; justify-content:space-between;}
.gdi-header-title {font-size:1.35rem;font-weight:800;line-height:1.1;}
.gdi-header-sub {font-size:.8rem;opacity:.9;margin-top:.15rem;}
.map-title {font-size:1.25rem;font-weight:800;color:#0f172a;margin:.15rem 0 .35rem 0;}
.legend-row {display:flex;gap:1.05rem;align-items:center;flex-wrap:wrap;margin:.1rem 0 .45rem 0;font-size:.86rem;}
.legend-item {display:flex;gap:.35rem;align-items:center;}
.legend-box {width:13px;height:13px;border-radius:2px;display:inline-block;}
.section-bar {background:#0b3d78;color:white;padding:.42rem .65rem;border-radius:6px 6px 0 0;font-weight:800;margin-top:.1rem;}
.proof-card {background:#f8fbff;border:1px solid #bfd3ea;border-left:5px solid #0b3d78;border-radius:8px;padding:.8rem .9rem;margin:.35rem 0 .75rem 0;}
.proof-title {font-weight:800;color:#0b3d78;margin-bottom:.25rem;}
.trend-chip {display:inline-block;padding:.28rem .58rem;border-radius:999px;background:#e8f1fb;color:#0b3d78;font-weight:800;margin:.2rem 0 .55rem 0;}
.alert-box {background:#fff7ed;border:1px solid #fdba74;border-left:5px solid #ea580c;border-radius:8px;padding:.65rem .8rem;margin:.3rem 0;}
div[data-testid="stDataFrame"] {border:1px solid var(--border); border-radius:7px; overflow:hidden;}
.stTabs [data-baseweb="tab-list"] {gap:.25rem;}
.stTabs [data-baseweb="tab"] {font-weight:700;}
</style>
<div class="gdi-header">
  <div><div class="gdi-header-title">GeoIntel — GEOTECHNICAL DECISION INTELLIGENCE</div><div class="gdi-header-sub">Lift II East Draw, Seismic & Risk Decision Support</div></div>
  <div style="font-weight:700;">Lift II East</div>
</div>
""", unsafe_allow_html=True)

tab_polygon, tab_daily, tab_weekend, tab_weekly, tab_insights, tab_red_rating, tab_history = st.tabs([
    "🏠 GeoIntel Workspace",
    "📈 Daily Analysis",
    "📅 Weekend Analysis",
    "📊 Weekly Dashboard",
    "🧠 Intelligence",
    "🚨 Red Rating Tracker",
    "⚙️ History / Admin"
])



with tab_polygon:
    polygon_date = st.date_input("Reporting Date", value=datetime.now().date(), key="polygon_date")
    workflow_status, workflow_user, workflow_time = get_approval_status(polygon_date)
    is_locked = workflow_status == "Approved"
    wf1, wf2, wf3 = st.columns([1, 1, 2])
    wf1.metric("Workflow", workflow_status)
    wf2.metric("Record lock", "Locked" if is_locked else "Editable")
    wf3.caption(f"Last action: {workflow_user or '—'} {workflow_time or ''}")
    if is_locked:
        st.info("This reporting date is approved and locked against editing.")

    if "selected_polygon" not in st.session_state:
        st.session_state.selected_polygon = "A"

    mapping_df = load_csv(DRAWPOINT_MAPPING_FILE)
    merged_tonnes = pd.DataFrame()
    drawbell_totals = pd.DataFrame(columns=["Polygon", "Drawbell", "Tonnes"])
    polygon_totals = pd.DataFrame(columns=["Polygon", "Polygon Tonnes"])


    overview_tab, data_tab, analysis_tab, operations_tab = st.tabs([
        "🏠 Overview", "🧾 Data Entry", "📊 Analysis", "📋 Operations & Reports"
    ])

    with data_tab:
        st.caption("Enter or upload production, seismic and optional convergence data. Calculations update automatically.")
        source_mode = st.radio(
            "Production input",
            ["Manual drawpoint entry", "Upload spreadsheet"],
            horizontal=True,
            key="polygon_production_mode",
        )
    
        if source_mode == "Upload spreadsheet":
            tonnes_upload = st.file_uploader(
                "Upload production spreadsheet — drawpoint tonnes",
                type=["xlsx", "xls", "csv"],
                key="polygon_tonnes_upload",
            )
            tonnes_df = read_table(tonnes_upload) if tonnes_upload else pd.DataFrame()
    
            if not tonnes_df.empty:
                detected_drawpoint, detected_tonnes = auto_detect_production_columns(tonnes_df)
                p1, p2 = st.columns(2)
                draw_index = list(tonnes_df.columns).index(detected_drawpoint) if detected_drawpoint in tonnes_df.columns else 0
                tonnes_index = list(tonnes_df.columns).index(detected_tonnes) if detected_tonnes in tonnes_df.columns else min(1, len(tonnes_df.columns)-1)
                drawpoint_col = p1.selectbox("Drawpoint column", tonnes_df.columns, index=draw_index, key="polygon_drawpoint_col")
                tonnes_col = p2.selectbox("Tonnes column", tonnes_df.columns, index=tonnes_index, key="polygon_tonnes_col")
                if detected_drawpoint and detected_tonnes:
                    st.success(f"Detected columns automatically: {detected_drawpoint} and {detected_tonnes}.")
                template = build_mapping_template(tonnes_df[drawpoint_col])
                bell_mapping = build_drawbell_mapping(template, mapping_df)
    
                with st.expander("Drawbell → polygon master mapping", expanded=bool(mapping_df.empty)):
                    edited_bell_mapping = st.data_editor(
                        bell_mapping, width="stretch", hide_index=True, disabled=["Drawbell"],
                        column_config={
                            "Polygon": st.column_config.SelectboxColumn(
                                "Polygon", options=["", "A", "B", "C", "D", "E", "F"]
                            )
                        },
                        key="polygon_drawbell_mapping_editor",
                    )
                    edited_mapping = expand_drawbell_mapping(template, edited_bell_mapping)
                    if st.button("Save Master Mapping", key="save_polygon_mapping"):
                        save_csv(edited_mapping, DRAWPOINT_MAPPING_FILE)
                        st.success("Master mapping saved.")
    
                merged_tonnes, drawbell_totals, polygon_totals = aggregate_polygon_tonnes(
                    tonnes_df, drawpoint_col, tonnes_col, edited_mapping
                )
            else:
                st.caption("Upload a spreadsheet to calculate polygon tonnes automatically.")
    
        else:
            if mapping_df.empty:
                st.warning("Manual entry needs the master drawpoint mapping. Upload one production spreadsheet first, assign drawbells to polygons, and save the mapping.")
            else:
                manual_template = load_manual_tonnes_for_date(polygon_date, mapping_df)
                st.caption("Enter tonnes per drawpoint. Saved values reload automatically for the selected reporting date.")
                manual_input = st.data_editor(
                    manual_template,
                    width="stretch",
                    hide_index=True,
                    column_config={"Tonnes": st.column_config.NumberColumn("Tonnes", min_value=0.0, step=1.0)},
                    key=f"manual_drawpoint_tonnes_{polygon_date}",
                    disabled=["Drawpoint", "Drawbell", "Polygon"] if not is_locked else list(manual_template.columns),
                )
                save_manual_col, clear_manual_col = st.columns(2)
                if save_manual_col.button("Save Manual Tonnes", type="primary", use_container_width=True, disabled=is_locked):
                    save_manual_tonnes(polygon_date, manual_input)
                    st.success(f"Manual tonnes saved for {polygon_date}.")
                if clear_manual_col.button("Clear This Date", use_container_width=True, disabled=is_locked):
                    blank = manual_input.copy()
                    blank["Tonnes"] = 0.0
                    save_manual_tonnes(polygon_date, blank)
                    st.rerun()
                merged_tonnes = manual_input.copy()
                drawbell_totals = (
                    manual_input.groupby(["Polygon", "Drawbell"], as_index=False)["Tonnes"].sum()
                )
                polygon_totals = (
                    manual_input.groupby("Polygon", as_index=False)["Tonnes"].sum()
                    .rename(columns={"Tonnes": "Polygon Tonnes"})
                )
    
        with st.expander("Polygon seismic input", expanded=True):
            ims_pdf = st.file_uploader("Upload polygon IMS PDF to fill A-F automatically", type=["pdf"], key="polygon_ims_pdf")
            seismic_seed = pd.DataFrame({
                "Polygon": list("ABCDEF"), "Events": [0] * 6,
                "Energy J": [0.0] * 6, "Potency m3": [0.0] * 6,
                "Target Tonnes": [0.0] * 6,
            })
            parsed_ims = parse_polygon_ims_pdf(ims_pdf) if ims_pdf else pd.DataFrame()
            if not parsed_ims.empty:
                seismic_seed = seismic_seed.drop(columns=["Events", "Energy J", "Potency m3"]).merge(parsed_ims, on="Polygon", how="left").merge(
                    pd.DataFrame({"Polygon": list("ABCDEF"), "Target Tonnes": [0.0]*6}), on="Polygon", how="left")
                for c in ["Events", "Energy J", "Potency m3"]:
                    seismic_seed[c] = pd.to_numeric(seismic_seed[c], errors="coerce").fillna(0)
                st.success(f"IMS data imported for {len(parsed_ims)} polygon(s).")
            seismic_input = st.data_editor(
                seismic_seed, width="stretch", hide_index=True,
                disabled=list(seismic_seed.columns) if is_locked else ["Polygon"],
                key=f"polygon_seismic_editor_v7_{polygon_date}",
            )
    
        with st.expander("Convergence overlay", expanded=False):
            saved_conv = load_csv(CONVERGENCE_FILE)
            conv_seed = pd.DataFrame({"Polygon": list("ABCDEF"), "Closure mm": [0.0]*6, "Rate mm/day": [0.0]*6, "Convergence Trend": ["Stable"]*6})
            if not saved_conv.empty and "Report Date" in saved_conv.columns:
                day_conv = saved_conv[saved_conv["Report Date"].astype(str) == str(polygon_date)].copy()
                if not day_conv.empty:
                    conv_seed = conv_seed.drop(columns=["Closure mm", "Rate mm/day", "Convergence Trend"]).merge(day_conv[["Polygon", "Closure mm", "Rate mm/day", "Convergence Trend"]], on="Polygon", how="left")
                    conv_seed[["Closure mm", "Rate mm/day"]] = conv_seed[["Closure mm", "Rate mm/day"]].fillna(0)
                    conv_seed["Convergence Trend"] = conv_seed["Convergence Trend"].fillna("Stable")
            convergence_input = st.data_editor(conv_seed, hide_index=True, width="stretch", disabled=list(conv_seed.columns) if is_locked else ["Polygon"], key=f"convergence_editor_{polygon_date}")
            if st.button("Save Convergence", disabled=is_locked, key="save_convergence"):
                conv_save = convergence_input.copy(); conv_save.insert(0, "Report Date", str(polygon_date))
                existing_conv = load_csv(CONVERGENCE_FILE)
                combined_conv = pd.concat([existing_conv, conv_save], ignore_index=True) if not existing_conv.empty else conv_save
                combined_conv = combined_conv.drop_duplicates(["Report Date", "Polygon"], keep="last")
                save_csv(combined_conv, CONVERGENCE_FILE)
                write_audit(current_user, current_role, "Saved convergence data", f"Report date {polygon_date}")
                st.success("Convergence data saved.")

    results = seismic_input.merge(polygon_totals, on="Polygon", how="left").merge(convergence_input, on="Polygon", how="left")
    for col in ["Polygon Tonnes", "Energy J", "Potency m3", "Events", "Target Tonnes"]:
        results[col] = pd.to_numeric(results[col], errors="coerce").fillna(0.0)

    results["E/P"] = results.apply(lambda r: r["Energy J"] / r["Potency m3"] if r["Potency m3"] > 0 else 0.0, axis=1)
    results["E/T"] = results.apply(lambda r: r["Energy J"] / r["Polygon Tonnes"] if r["Polygon Tonnes"] > 0 else None, axis=1)
    results["Status"] = results.apply(
        lambda r: get_status(r["E/P"], r["E/T"], ep_threshold, et_threshold) if r["Potency m3"] > 0 else "Waiting for seismic data",
        axis=1,
    )
    results["Map Rating"] = results["Status"].map(
        lambda s: "Green" if s == "Efficient Draw" else "Red" if s == "High-Risk Draw" else "Amber" if s in ["Brittle but Efficient", "Ductile but Costly"] else "Grey"
    )
    results["Health Score"] = results.apply(
        lambda r: polygon_health_score(r["E/P"], r["E/T"], r["Events"], r["Potency m3"], r["Polygon Tonnes"], r["Target Tonnes"], ep_threshold, et_threshold),
        axis=1,
    )
    results["Health Rating"] = results["Health Score"].map(polygon_health_rating)
    results["Convergence Risk"] = results.apply(lambda r: convergence_risk_label(to_float(r.get("Closure mm",0)), to_float(r.get("Rate mm/day",0)), r.get("Status","")), axis=1)
    results["Combined Risk Index"] = results.apply(lambda r: combined_risk_index(r, risk_engine_config), axis=1)
    results["Risk Inputs Used"] = risk_inputs_used(risk_engine_config)
    results["Combined Risk"] = results["Combined Risk Index"].map(combined_risk_label)
    neighbour_notes = neighbouring_polygon_assessment(results)
    results["Neighbour Influence"] = results["Polygon"].map(neighbour_notes)
    results["Decision Recommendation"] = results.apply(decision_recommendation, axis=1)
    results["Achievement %"] = results.apply(
        lambda r: (100 * r["Polygon Tonnes"] / r["Target Tonnes"]) if r["Target Tonnes"] > 0 else None,
        axis=1,
    )

    polygon_history = load_csv(POLYGON_HISTORY_FILE)

    with analysis_tab:
        st.caption("Review data quality, historical replay, summaries, comparisons and exceptions.")
        quality_issues = polygon_data_quality(results, merged_tonnes, mapping_df)
        with st.expander("Data quality checks", expanded=bool(quality_issues)):
            if quality_issues:
                for issue in quality_issues:
                    st.warning(issue)
            else:
                st.success("All core data-quality checks passed.")
    
        dashboard_mode = st.radio(
            "Dashboard view", ["Current input", "Historical replay"], horizontal=True,
            key="polygon_dashboard_mode",
        )
        dashboard_results = results
        replay_date = None
        if dashboard_mode == "Historical replay":
            if polygon_history.empty or "Report Date" not in polygon_history.columns:
                st.info("Save polygon results first to unlock historical replay.")
            else:
                replay_dates = sorted(polygon_history["Report Date"].dropna().astype(str).unique(), reverse=True)
                replay_date = st.select_slider(
                    "Replay date", options=list(reversed(replay_dates)), value=replay_dates[0],
                    key="polygon_replay_date",
                )
                replay_snapshot = prepare_history_snapshot(polygon_history, replay_date, ep_threshold, et_threshold)
                if not replay_snapshot.empty:
                    dashboard_results = replay_snapshot
                    st.info(f"Historical replay: showing saved polygon conditions for {replay_date}.")
    
    with overview_tab:
        st.caption("At-a-glance status, selected polygon proof, matrix, alerts and trends.")
        rating_counts = dashboard_results["Map Rating"].value_counts()
        g1, g2, g3, g4 = st.columns(4)
        g1.metric("Good", int(rating_counts.get("Green", 0)))
        g2.metric("Average / watch", int(rating_counts.get("Amber", 0)))
        g3.metric("Poor", int(rating_counts.get("Red", 0)))
        g4.metric("Waiting", int(rating_counts.get("Grey", 0)))
    
        left, right = st.columns([1.15, 1], gap="medium")
        with left:
            st.markdown('<div class="map-title">Lift II East — Polygon Status Map</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="legend-row">'
                '<span class="legend-item"><span class="legend-box" style="background:#15803d"></span>Good</span>'
                '<span class="legend-item"><span class="legend-box" style="background:#d97706"></span>Watch</span>'
                '<span class="legend-item"><span class="legend-box" style="background:#dc2626"></span>Poor</span>'
                '<span class="legend-item"><span class="legend-box" style="background:#cbd5e1"></span>No data</span>'
                '</div>', unsafe_allow_html=True
            )
            if DEFAULT_MAP_FILE.exists():
                map_event = st.plotly_chart(
                    polygon_map_figure(dashboard_results, DEFAULT_MAP_FILE, st.session_state.selected_polygon),
                    width="stretch", key="polygon_status_map_v63_clickable",
                    config={"displayModeBar": False, "displaylogo": False},
                    on_select="rerun",
                    selection_mode="points",
                )
                try:
                    selected_points = map_event.selection.points
                except (AttributeError, TypeError):
                    selected_points = []
                if selected_points:
                    clicked = selected_points[0].get("customdata", [None])
                    clicked_poly = clicked[0] if isinstance(clicked, (list, tuple)) else clicked
                    if clicked_poly in list("ABCDEF") and clicked_poly != st.session_state.selected_polygon:
                        st.session_state.selected_polygon = clicked_poly
                        st.rerun()
                st.caption("Click the shaded polygon directly. The fallback buttons remain available if your Streamlit version does not return click events.")
            else:
                st.error("Keep lift2east_polygon_map_clean_v2.png in the same folder as this app.")
    
        with right:
            st.markdown('<div class="section-bar">SELECTED POLYGON</div>', unsafe_allow_html=True)
            bcols = st.columns(3)
            for index, poly in enumerate("ABCDEF"):
                if bcols[index % 3].button(poly, key=f"select_poly_v4_{poly}", use_container_width=True):
                    st.session_state.selected_polygon = poly
                    st.rerun()
    
            selected_poly = st.session_state.selected_polygon
            selected = dashboard_results[dashboard_results["Polygon"] == selected_poly].iloc[0]
            rating = selected["Map Rating"]
            badge_colour = {"Green": "#15803d", "Amber": "#d97706", "Red": "#dc2626", "Grey": "#64748b"}.get(rating, "#64748b")
            st.markdown(f"## POLYGON {selected_poly}")
            st.markdown(f'<span class="status-badge" style="background:{badge_colour}">{selected["Status"]}</span>', unsafe_allow_html=True)
    
            trend_label, trend_change, trend_arrow = polygon_trend(
                polygon_history, selected_poly, selected["Health Score"]
            )
            explanation = polygon_explanation(selected, ep_threshold, et_threshold)
            active_alerts = polygon_alerts(selected, trend_label, trend_change, ep_threshold, et_threshold)
    
            st.markdown(
                f'<div class="proof-card"><div class="proof-title">WHY THIS COLOUR</div>{explanation}</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f'<span class="trend-chip">{trend_arrow} {trend_label} ({trend_change:+.0f})</span>',
                unsafe_allow_html=True,
            )
    
            h1, h2 = st.columns([1, 1])
            h1.metric("Health Score", f"{int(selected['Health Score'])}/100")
            h2.metric("Health Rating", selected["Health Rating"])
            st.progress(int(selected["Health Score"]))
    
            k1, k2 = st.columns(2)
            k1.metric("Energy", f"{selected['Energy J']:,.0f} J")
            k2.metric("Potency", f"{selected['Potency m3']:,.3f} m³")
            k3, k4 = st.columns(2)
            k3.metric("Polygon tonnes", f"{selected['Polygon Tonnes']:,.0f} t")
            k4.metric("Events", f"{to_int(selected['Events'])}")
            k5, k6 = st.columns(2)
            k5.metric("E/P", f"{selected['E/P']:,.2f}")
            et_display = selected["E/T"]
            k6.metric("E/T", f"{et_display:,.2f}" if et_display is not None and not pd.isna(et_display) else "Waiting")
            c1, c2 = st.columns(2)
            c1.metric("Closure", f"{to_float(selected.get('Closure mm',0)):.1f} mm")
            c2.metric("Closure rate", f"{to_float(selected.get('Rate mm/day',0)):.2f} mm/day")
            st.caption(f"Convergence assessment: {selected.get('Convergence Risk', 'No convergence data')}")
    
            if selected["Target Tonnes"] > 0:
                st.metric("Production achievement", f"{selected['Achievement %']:.1f}%", f"Target {selected['Target Tonnes']:,.0f} t")
    
            if active_alerts:
                st.markdown('<div class="section-bar">ACTIVE ALERTS</div>', unsafe_allow_html=True)
                for alert in active_alerts:
                    st.markdown(f'<div class="alert-box">⚠️ {alert}</div>', unsafe_allow_html=True)
            else:
                st.success("No active alerts for the selected polygon.")
    
            insight = management_summary(
                selected_poly,
                to_int(selected["Events"]),
                to_float(selected["Energy J"]),
                to_float(selected["Potency m3"]),
                to_float(selected["Polygon Tonnes"]),
                to_float(selected["E/P"]),
                None if et_display is None or pd.isna(et_display) else to_float(et_display),
                selected["Status"],
                to_int(selected["Health Score"]),
            )
            st.markdown('<div class="section-bar">MANAGEMENT INSIGHT</div>', unsafe_allow_html=True)
            st.info(insight + f" Trend: {trend_arrow} {trend_label}.")
    
            report_file = selected_polygon_report(
                selected, trend_label, trend_change, explanation, active_alerts
            )
            st.download_button(
                "Download Selected Polygon Report",
                data=report_file,
                file_name=f"GDI_Polygon_{selected_poly}_{polygon_date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"download_polygon_report_{selected_poly}",
            )
    
            if selected["Potency m3"] > 0 and et_display is not None and not pd.isna(et_display):
                selected_history = load_csv(POLYGON_HISTORY_FILE)
                if not selected_history.empty and "Polygon" in selected_history.columns:
                    selected_history = selected_history[selected_history["Polygon"].astype(str) == selected_poly]
                st.plotly_chart(
                    plot_matrix(selected["E/P"], et_display, selected_history, ep_threshold, et_threshold),
                    width="stretch", key=f"polygon_matrix_v4_{selected_poly}",
                )
            else:
                st.info("Enter energy, potency and tonnes to display the matrix position.")
    
            st.markdown('<div class="section-bar">POLYGON TREND</div>', unsafe_allow_html=True)
            trend_fig = polygon_trend_figure(polygon_history, selected_poly)
            if trend_fig is not None:
                st.plotly_chart(trend_fig, width="stretch", key=f"polygon_trend_chart_{selected_poly}")
            else:
                st.caption("Save at least two dated polygon records to display a trend chart.")

    with analysis_tab:
        st.markdown(f'<div class="section-bar">DRAWBELLS & DRAWPOINTS — POLYGON {st.session_state.selected_polygon}</div>', unsafe_allow_html=True)
        if not drawbell_totals.empty:
            selected_poly = st.session_state.selected_polygon
            poly_drawbells = drawbell_totals[drawbell_totals["Polygon"] == selected_poly].copy()
            poly_points = merged_tonnes[merged_tonnes["Polygon"] == selected_poly][["Drawpoint", "Drawbell", "Tonnes"]].copy()
            d1, d2 = st.columns([1, 1.3])
            d1.dataframe(poly_drawbells, width="stretch", hide_index=True)
            d2.dataframe(poly_points, width="stretch", hide_index=True)
        else:
            st.caption("Drawbell and drawpoint detail appears after tonnes are entered or uploaded.")
    
        st.markdown('<div class="section-bar">POLYGON SUMMARY</div>', unsafe_allow_html=True)
        summary_cols = ["Polygon", "Status", "Combined Risk", "Combined Risk Index", "Health Score", "Polygon Tonnes", "Achievement %", "E/P", "E/T", "Events", "Closure mm", "Rate mm/day", "Convergence Risk", "Risk Inputs Used", "Neighbour Influence"]
        st.dataframe(dashboard_results[summary_cols], width="stretch", hide_index=True)
    
        st.markdown('<div class="section-bar">DECISION INTELLIGENCE</div>', unsafe_allow_html=True)
        risk_rank = dashboard_results.sort_values("Combined Risk Index", ascending=False).copy()
        r1, r2, r3, r4 = st.columns(4)
        r1.metric("Critical", int((risk_rank["Combined Risk"] == "Critical").sum()))
        r2.metric("High", int((risk_rank["Combined Risk"] == "High").sum()))
        r3.metric("Moderate", int((risk_rank["Combined Risk"] == "Moderate").sum()))
        r4.metric("Low", int((risk_rank["Combined Risk"] == "Low").sum()))
    
        selected_decision = dashboard_results[dashboard_results["Polygon"] == st.session_state.selected_polygon].iloc[0]
        st.markdown(f"**Polygon {st.session_state.selected_polygon} — {selected_decision['Combined Risk']} risk ({int(selected_decision['Combined Risk Index'])}/100)**")
        st.caption(f"Index inputs used: {selected_decision.get('Risk Inputs Used', 'No components enabled')}")
        st.info(selected_decision["Decision Recommendation"])
        st.caption(selected_decision["Neighbour Influence"])
    
        st.markdown('<div class="section-bar">EXCEPTION DASHBOARD</div>', unsafe_allow_html=True)
        exceptions = risk_rank[
            (risk_rank["Combined Risk"].isin(["Critical", "High"]))
            | (risk_rank["Map Rating"] == "Red")
            | (risk_rank["Convergence Risk"].isin(["Combined geotechnical concern", "High convergence"]))
        ][["Polygon", "Combined Risk", "Combined Risk Index", "Status", "Convergence Risk", "Neighbour Influence", "Decision Recommendation"]]
        if exceptions.empty:
            st.success("No high-priority exceptions for this reporting view.")
        else:
            st.dataframe(exceptions, width="stretch", hide_index=True)

    with operations_tab:
        st.caption("Manage handovers, reports, approvals and saved records.")
        st.markdown('<div class="section-bar">SHIFT / DAILY HANDOVER</div>', unsafe_allow_html=True)
        h1, h2 = st.columns([1, 1])
        with h1:
            shift_name = st.selectbox("Shift", ["Day Shift", "Night Shift", "Daily Handover"], key="handover_shift")
            handover_observation = st.text_area("Observation / outstanding issue", key="handover_observation")
            handover_action = st.text_area("Action taken / required", key="handover_action")
        with h2:
            handover_owner = st.text_input("Responsible person", key="handover_owner")
            handover_status = st.selectbox("Item status", ["Open", "In progress", "Closed"], key="handover_status")
            if st.button("Save Handover Item", key="save_handover_item"):
                if not handover_observation.strip():
                    st.warning("Enter an observation or outstanding issue first.")
                else:
                    save_handover(polygon_date, shift_name, handover_observation, handover_action, handover_owner, handover_status, current_user)
                    write_audit(current_user, current_role, "Saved handover item", f"{polygon_date} - {shift_name}")
                    st.success("Handover item saved.")
    
        handover_df = load_csv(HANDOVER_FILE)
        if not handover_df.empty:
            open_items = handover_df[handover_df["Item Status"].astype(str) != "Closed"].copy()
            if not open_items.empty:
                st.caption("Open items carry forward until they are closed.")
                st.dataframe(open_items.tail(20), width="stretch", hide_index=True)

    with overview_tab:
        st.markdown('<div class="section-bar">EXECUTIVE OVERVIEW</div>', unsafe_allow_html=True)
        executive_text = executive_summary_text(dashboard_results)
        st.info(executive_text)
        ex1, ex2, ex3, ex4 = st.columns(4)
        ranked = dashboard_results.copy()
        ranked["Health Score"] = pd.to_numeric(ranked["Health Score"], errors="coerce").fillna(0)
        ex1.metric("Total tonnes", f"{pd.to_numeric(ranked['Polygon Tonnes'], errors='coerce').fillna(0).sum():,.0f} t")
        ex2.metric("Best polygon", str(ranked.sort_values("Health Score", ascending=False).iloc[0]["Polygon"]))
        ex3.metric("Highest concern", str(ranked.sort_values("Health Score", ascending=True).iloc[0]["Polygon"]))
        ex4.metric("Active poor ratings", int((ranked["Map Rating"] == "Red").sum()))

    with analysis_tab:
        st.markdown('<div class="section-bar">POLYGON COMPARISON</div>', unsafe_allow_html=True)
        cmp1, cmp2 = st.columns(2)
        compare_a = cmp1.selectbox("First polygon", list("ABCDEF"), index=0, key="compare_polygon_a")
        compare_b = cmp2.selectbox("Second polygon", list("ABCDEF"), index=1, key="compare_polygon_b")
        comparison = dashboard_results[dashboard_results["Polygon"].isin([compare_a, compare_b])][
            ["Polygon", "Status", "Health Score", "Polygon Tonnes", "Energy J", "Potency m3", "E/P", "E/T", "Events"]
        ].copy()
        st.dataframe(comparison, width="stretch", hide_index=True)
        if len(comparison) == 2:
            comp = comparison.set_index("Polygon")
            healthier = compare_a if to_float(comp.loc[compare_a, "Health Score"]) >= to_float(comp.loc[compare_b, "Health Score"]) else compare_b
            st.caption(f"Polygon {healthier} currently has the stronger health score in this comparison.")

    with operations_tab:
        full_report = build_dashboard_report(
            replay_date or polygon_date, dashboard_results, quality_issues, executive_text
        )
        st.download_button(
            "Download Full Management Report", data=full_report,
            file_name=f"GDI_Lift_II_East_{replay_date or polygon_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="download_full_management_report",
        )
        printable_report = build_printable_html_report(replay_date or polygon_date, dashboard_results, executive_text, quality_issues)
        st.download_button(
            "Download Printable PDF Version", data=printable_report,
            file_name=f"GDI_Lift_II_East_{replay_date or polygon_date}.html",
            mime="text/html", use_container_width=True, key="download_printable_report",
            help="Open the downloaded file in your browser, then Print and choose Save as PDF.",
        )
    
        st.markdown('<div class="section-bar">APPROVAL WORKFLOW</div>', unsafe_allow_html=True)
        ap1, ap2, ap3 = st.columns(3)
        if ap1.button("Save as Draft", disabled=is_locked, use_container_width=True):
            save_approval_status(polygon_date, "Draft", current_user, current_role); st.rerun()
        if ap2.button("Mark Reviewed", disabled=is_locked or current_role not in ["Strata Control Officer", "Geotechnical Specialist", "Geotechnical Superintendent", "Manager"], use_container_width=True):
            save_approval_status(polygon_date, "Reviewed", current_user, current_role); st.rerun()
        if ap3.button("Approve & Lock", disabled=is_locked or current_role not in ["Geotechnical Superintendent", "Manager"], use_container_width=True):
            save_approval_status(polygon_date, "Approved", current_user, current_role); st.rerun()
        if is_locked and current_role in ["Geotechnical Superintendent", "Manager"] and st.button("Reopen Approved Record", use_container_width=True):
            save_approval_status(polygon_date, "Reviewed", current_user, current_role); st.rerun()
    
        if st.button("Save Polygon Results", type="primary", key="save_polygon_results_v4", disabled=is_locked):
            save_df = results.copy()
            save_df.insert(0, "Report Date", str(polygon_date))
            existing = load_csv(POLYGON_HISTORY_FILE)
            combined = pd.concat([existing, save_df], ignore_index=True) if not existing.empty else save_df
            combined = combined.drop_duplicates(subset=["Report Date", "Polygon"], keep="last")
            save_csv(combined, POLYGON_HISTORY_FILE)
            write_audit(current_user, current_role, "Saved polygon results", f"Report date {polygon_date}")
            st.success("Polygon results saved.")

with tab_daily:
    st.header("Daily Draw Intelligence")

    uploaded_pdf = st.file_uploader("Upload Daily IMS PDF", type=["pdf"], key="daily_pdf")
    uploaded_excel = st.file_uploader("Optional: Upload Production Tonnes Excel", type=["xlsx", "xls", "csv"], key="daily_excel")

    tonnes_from_excel = None

    if uploaded_excel:
        tonnes_df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith(".csv") else pd.read_excel(uploaded_excel)
        st.dataframe(tonnes_df.head(), width="stretch")
        tonnes_col = st.selectbox("Select Tonnes Column", tonnes_df.columns, key="daily_tonnes_col")
        tonnes_from_excel = pd.to_numeric(tonnes_df[tonnes_col], errors="coerce").sum()
        st.success(f"Tonnes from Excel: {tonnes_from_excel:,.2f}")

    if uploaded_pdf:
        full_text = pdf_to_text(uploaded_pdf)
        poly_data = parse_poly_sections(full_text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index)
            data = poly_data[target_poly]

            st.success(f"{target_poly} section found.")

            c1, c2, c3 = st.columns(3)
            c1.metric("New Events", data["Events"])
            c1.metric("Total Energy (J)", f"{data['Energy J']:,.2f}")
            c2.metric("Total Potency (m³)", f"{data['Potency m3']:,.3f}")
            c2.metric("Current Activity Rate", data["Current Activity Rate"])
            c3.metric("Medium Activity Rate", data["Medium Activity Rate"])
            c3.metric("Last Event", data["Last Event"])

            st.header("Tonnes Input")
            input_mode = st.radio("Tonnes Source", ["Manual Entry", "Excel Upload"], horizontal=True, key="daily_mode")

            if input_mode == "Excel Upload" and tonnes_from_excel is not None:
                tonnes = tonnes_from_excel
            else:
                tonnes = st.number_input("Tonnes Drawn", min_value=0.0, value=0.0, step=1.0, key="daily_tonnes")

            energy = data["Energy J"]
            potency = data["Potency m3"]

            if potency > 0:
                ep = energy / potency
                et = energy / tonnes if tonnes > 0 else None
                status = get_status(ep, et, ep_threshold, et_threshold)
                health = cave_health_score(status)
                action = recommended_action(status)

                st.header("Draw Intelligence Results")
                r1, r2, r3, r4 = st.columns(4)
                r1.metric("E/P", f"{ep:,.2f}")
                r2.metric("E/T", f"{et:,.2f}" if et is not None else "Waiting")
                r3.metric("Status", status)
                r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                score = draw_compliance_score(ep, et, data["Events"], potency, ep_threshold, et_threshold)
                s1, s2 = st.columns(2)
                s1.metric("Draw Compliance Score", f"{score}/100")
                s2.metric("Score Label", score_label(score))
                st.info(draw_advisor(score, status))

                st.info(action)

                if et is not None:
                    st.plotly_chart(plot_matrix(ep, et, load_csv(DAILY_FILE), ep_threshold, et_threshold), width="stretch")

                report_date = st.text_input("Report Date", value=clean_report_date(data["Last Event"]), key="daily_date")

                comment = management_summary(
                    target_poly, data["Events"], energy, potency, tonnes, ep, et, status, score
                )

                st.text_area("Management Comment", value=comment, height=170)

                if st.button("Save or Update Daily Record"):
                    record = {
                        "Report Date": report_date,
                        "Poly": target_poly,
                        "New Events": data["Events"],
                        "Energy J": energy,
                        "Potency m3": potency,
                        "Tonnes": tonnes,
                        "E/P": ep,
                        "E/T": et if et is not None else "",
                        "Status": status,
                        "Cave Health Score": health,
                        "Draw Compliance Score": score,
                        "Score Label": score_label(score),
                        "Recommended Action": action,
                    }
                    save_or_update(DAILY_FILE, record, ["Report Date", "Poly"])
                    st.success("Daily record saved or updated.")


with tab_weekend:
    st.header("Weekend Cumulative Processor")
    st.write("Save Friday baseline, Fri-Sat, Fri-Sun and Fri-Mon cumulative PDFs. GDI derives each separate day.")

    uploaded_weekend_pdf = st.file_uploader("Upload Weekend Cumulative IMS PDF", type=["pdf"], key="weekend_pdf")

    weekend_start = st.text_input("Weekend Start Date / Friday Date", value=datetime.now().strftime("%Y-%m-%d"))
    stage = st.selectbox(
        "Cumulative Stage",
        ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"]
    )

    stage_order = {
        "Friday Baseline": 0,
        "Friday to Saturday": 1,
        "Friday to Sunday": 2,
        "Friday to Monday": 3,
    }

    stage_daily_name = {
        "Friday Baseline": "Friday",
        "Friday to Saturday": "Saturday",
        "Friday to Sunday": "Sunday",
        "Friday to Monday": "Monday",
    }

    if uploaded_weekend_pdf:
        text = pdf_to_text(uploaded_weekend_pdf)
        poly_data = parse_poly_sections(text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index, key="weekend_poly")
            data = poly_data[target_poly]

            st.subheader("Cumulative Values")
            c1, c2, c3 = st.columns(3)
            c1.metric("Cumulative Events", data["Events"])
            c2.metric("Cumulative Energy (J)", f"{data['Energy J']:,.2f}")
            c3.metric("Cumulative Potency (m³)", f"{data['Potency m3']:,.3f}")

            cumulative_record = {
                "Weekend Start": weekend_start,
                "Poly": target_poly,
                "Stage": stage,
                "Stage Order": stage_order[stage],
                "Events": data["Events"],
                "Energy J": data["Energy J"],
                "Potency m3": data["Potency m3"],
                "Last Event": data["Last Event"],
            }

            if st.button("Save Weekend Cumulative Record"):
                save_or_update(WEEKEND_FILE, cumulative_record, ["Weekend Start", "Poly", "Stage"])
                st.success("Weekend cumulative record saved.")

            weekend_df = load_csv(WEEKEND_FILE)

            if not weekend_df.empty:
                filtered = weekend_df[
                    (weekend_df["Weekend Start"].astype(str) == weekend_start)
                    & (weekend_df["Poly"].astype(str) == target_poly)
                ].copy()

                filtered["Stage Order"] = pd.to_numeric(filtered["Stage Order"], errors="coerce")
                filtered = filtered.sort_values("Stage Order")

                st.subheader("Saved Weekend Cumulative Records")
                st.dataframe(filtered, width="stretch")

                current_order = stage_order[stage]

                if current_order == 0:
                    derived_events = data["Events"]
                    derived_energy = data["Energy J"]
                    derived_potency = data["Potency m3"]
                else:
                    previous = filtered[filtered["Stage Order"] == current_order - 1]

                    if previous.empty:
                        st.warning("Previous cumulative stage is missing. Upload/save the previous stage first.")
                        derived_events = None
                        derived_energy = None
                        derived_potency = None
                    else:
                        prev = previous.iloc[-1]
                        derived_events = data["Events"] - int(prev["Events"])
                        derived_energy = data["Energy J"] - float(prev["Energy J"])
                        derived_potency = data["Potency m3"] - float(prev["Potency m3"])

                if derived_events is not None:
                    derived_day = stage_daily_name[stage]

                    st.subheader("Derived Separate Day Values")
                    derived_df = pd.DataFrame([{
                        "Derived Day": derived_day,
                        "Events": derived_events,
                        "Energy J": derived_energy,
                        "Potency m3": derived_potency,
                    }])
                    st.dataframe(derived_df, width="stretch")

                    if derived_energy < 0 or derived_potency < 0 or derived_events < 0:
                        st.error("Derived values are negative. Check the cumulative order.")
                    else:
                        st.success("Derived day calculated successfully.")

                        st.header("Weekend Derived Draw Compliance")
                        weekend_tonnes = st.number_input(
                            f"{derived_day} Tonnes Drawn",
                            min_value=0.0,
                            value=0.0,
                            step=1.0,
                            key="weekend_tonnes"
                        )

                        if derived_potency > 0:
                            ep = derived_energy / derived_potency
                            et = derived_energy / weekend_tonnes if weekend_tonnes > 0 else None
                            status = get_status(ep, et, ep_threshold, et_threshold)
                            health = cave_health_score(status)
                            action = recommended_action(status)

                            r1, r2, r3, r4 = st.columns(4)
                            r1.metric("Derived E/P", f"{ep:,.2f}")
                            r2.metric("Derived E/T", f"{et:,.2f}" if et is not None else "Waiting")
                            r3.metric("Status", status)
                            r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                            score = draw_compliance_score(ep, et, derived_events, derived_potency, ep_threshold, et_threshold)
                            ws1, ws2 = st.columns(2)
                            ws1.metric("Draw Compliance Score", f"{score}/100")
                            ws2.metric("Score Label", score_label(score))
                            st.info(draw_advisor(score, status))

                            st.info(action)

                            if et is not None:
                                st.plotly_chart(
                                    plot_matrix(ep, et, load_csv(WEEKEND_DERIVED_FILE), ep_threshold, et_threshold),
                                    width="stretch"
                                )

                            if st.button("Save Weekend Derived Daily Record"):
                                actual_report_date = weekend_actual_date(weekend_start, derived_day)

                                derived_record = {
                                    "Weekend Start": weekend_start,
                                    "Derived Day": derived_day,
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Draw Compliance Score": score,
                                    "Score Label": score_label(score),
                                    "Recommended Action": action,
                                }

                                daily_record = {
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "New Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Draw Compliance Score": score,
                                    "Score Label": score_label(score),
                                    "Recommended Action": action,
                                }

                                save_or_update(WEEKEND_DERIVED_FILE, derived_record, ["Weekend Start", "Derived Day", "Poly"])
                                save_or_update(DAILY_FILE, daily_record, ["Report Date", "Poly"])
                                st.success("Weekend derived daily record saved and Daily Draw History updated.")


with tab_weekly:
    st.header("Weekly IMS Dashboard")

    uploaded_weekly_pdf = st.file_uploader("Upload Monday-to-Monday Weekly IMS Summary", type=["pdf"], key="weekly_pdf")

    if uploaded_weekly_pdf:
        text = pdf_to_text(uploaded_weekly_pdf)
        poly_data = parse_poly_sections(text)

        week_range = parse_week_range(text)
        mine_total = parse_mine_total(text)
        largest_ml = parse_largest_magnitude(text)

        st.subheader("Weekly Overview")
        o1, o2, o3 = st.columns(3)
        o1.metric("Reporting Period", week_range or "Not found")
        o2.metric("Mine-wide Events", mine_total)
        o3.metric("Largest ML", largest_ml if largest_ml is not None else "Not found")

        if poly_data:
            weekly_df = pd.DataFrame(poly_data.values())

            st.subheader("Weekly Zone Table")
            st.dataframe(weekly_df, width="stretch")

            weekly_df["Lift Group"] = weekly_df["Poly"].apply(
                lambda x: "Lift II" if str(x).startswith("Lift_II")
                else "Lift I" if str(x).startswith("Lift_I_")
                else "Other"
            )

            lift_split = weekly_df.groupby("Lift Group", as_index=False)["Events"].sum()
            lift_split = lift_split.rename(columns={"Lift Group": "Zone"})

            lift_i = weekly_df[weekly_df["Lift Group"] == "Lift I"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})
            lift_ii = weekly_df[weekly_df["Lift Group"] == "Lift II"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})

            st.plotly_chart(pie_chart(lift_split, "Mine-wide Event Split"), width="stretch")

            c1, c2 = st.columns(2)
            with c1:
                if not lift_i.empty:
                    st.plotly_chart(pie_chart(lift_i, "Lift I Event Distribution"), width="stretch")
            with c2:
                if not lift_ii.empty:
                    st.plotly_chart(pie_chart(lift_ii, "Lift II Event Distribution"), width="stretch")

            ranking = weekly_df.sort_values("Events", ascending=False)[["Poly", "Events", "Energy J", "Potency m3", "Comment"]]
            st.subheader("Hotspot Ranking")
            st.dataframe(ranking, width="stretch")
            st.plotly_chart(bar_chart(ranking.rename(columns={"Poly": "Zone"}), "Weekly Hotspot Ranking"), width="stretch")

            dominant = ranking.iloc[0]["Poly"] if not ranking.empty else "Not found"
            second = ranking.iloc[1]["Poly"] if len(ranking) > 1 else "Not found"

            weekly_comment = (
                f"Weekly seismicity for the period {week_range} recorded {mine_total} mine-wide events. "
                f"The dominant hotspot was {dominant}, followed by {second}. "
                f"The largest recorded event was ML{largest_ml}. Continued monitoring is recommended."
            )

            st.subheader("Weekly Executive Summary")
            st.text_area("Copy-ready weekly summary", value=weekly_comment, height=170)

            st.header("Weekly Lift II East Draw Indicator")

            if "Lift_II_E" in poly_data:
                liie = poly_data["Lift_II_E"]

                w1, w2, w3 = st.columns(3)
                w1.metric("Lift II East Weekly Events", liie["Events"])
                w2.metric("Lift II East Weekly Energy (J)", f"{liie['Energy J']:,.2f}")
                w3.metric("Lift II East Weekly Potency (m³)", f"{liie['Potency m3']:,.3f}")

                weekly_tonnes = st.number_input(
                    "Weekly Lift II East Tonnes Drawn",
                    min_value=0.0,
                    value=0.0,
                    step=1.0,
                    key="weekly_tonnes"
                )

                if liie["Potency m3"] > 0:
                    weekly_ep = liie["Energy J"] / liie["Potency m3"]
                    weekly_et = liie["Energy J"] / weekly_tonnes if weekly_tonnes > 0 else None
                    weekly_status = get_status(weekly_ep, weekly_et, ep_threshold, et_threshold)
                    weekly_health = cave_health_score(weekly_status)
                    weekly_action = recommended_action(weekly_status)

                    q1, q2, q3, q4 = st.columns(4)
                    q1.metric("Weekly E/P", f"{weekly_ep:,.2f}")
                    q2.metric("Weekly E/T", f"{weekly_et:,.2f}" if weekly_et is not None else "Waiting")
                    q3.metric("Weekly Status", weekly_status)
                    q4.metric("Weekly Cave Health", f"{weekly_health}/100" if weekly_health else "Waiting")

                    st.info(weekly_action)

                    if weekly_et is not None:
                        st.plotly_chart(
                            plot_matrix(weekly_ep, weekly_et, load_csv(WEEKLY_FILE), ep_threshold, et_threshold),
                            width="stretch"
                        )

                    weekly_draw_comment = (
                        f"For the weekly period {week_range}, Lift II East recorded {liie['Events']} events, "
                        f"with total energy of {liie['Energy J']:,.0f} J and total potency of {liie['Potency m3']:.3f} m³. "
                        f"The weekly E/P ratio is {weekly_ep:,.2f}. "
                    )

                    if weekly_et is not None:
                        weekly_draw_comment += (
                            f"The weekly E/T ratio is {weekly_et:,.2f}, placing Lift II East in the "
                            f"'{weekly_status}' quadrant. {weekly_action}"
                        )
                    else:
                        weekly_draw_comment += "Weekly E/T cannot be finalised until weekly tonnes are entered."

                    st.text_area("Weekly Lift II East Draw Comment", value=weekly_draw_comment, height=150)

                    if st.button("Save Weekly Summary"):
                        weekly_record = {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Second Hotspot": second,
                            "Lift II East Events": liie["Events"],
                            "Lift II East Energy J": liie["Energy J"],
                            "Lift II East Potency m3": liie["Potency m3"],
                            "Weekly Tonnes": weekly_tonnes,
                            "Weekly E/P": weekly_ep,
                            "Weekly E/T": weekly_et if weekly_et is not None else "",
                            "Weekly Status": weekly_status,
                            "Weekly Cave Health Score": weekly_health,
                            "Weekly Comment": weekly_comment,
                            "Weekly Draw Comment": weekly_draw_comment,
                        }

                        save_or_update(WEEKLY_FILE, weekly_record, ["Week Range"])
                        st.success("Weekly summary saved.")

                    word_file = build_word_report(
                        "GDI Weekly IMS Summary",
                        weekly_comment + "\n\n" + weekly_draw_comment,
                        {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Lift II East Weekly E/P": round(weekly_ep, 2),
                            "Lift II East Weekly E/T": round(weekly_et, 2) if weekly_et is not None else "Waiting",
                            "Weekly Status": weekly_status,
                        }
                    )

                    st.download_button(
                        "Download Weekly Word Report",
                        data=word_file,
                        file_name="GDI_Weekly_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Lift_II_E section was not found in the weekly PDF.")


with tab_insights:
    st.header("GDI Insights")

    daily_insights = add_score_columns(safe_sort_history(load_csv(DAILY_FILE)), ep_threshold, et_threshold)
    weekly_insights = safe_sort_history(load_csv(WEEKLY_FILE))

    if daily_insights.empty:
        st.info("No daily records available yet. Save daily or weekend-derived records first.")
    else:
        latest = daily_insights.iloc[-1]
        latest_score = to_int(latest.get("Draw Compliance Score", 0))
        latest_status = str(latest.get("Status", ""))

        st.subheader("Current Draw Snapshot")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Latest Date", latest.get("Report Date", ""))
        c2.metric("Latest Poly", latest.get("Poly", ""))
        c3.metric("Compliance Score", f"{latest_score}/100")
        c4.metric("Status", latest_status)
        st.info(draw_advisor(latest_score, latest_status))

        st.subheader("30-Day Trend Forecast")
        forecast_ep = simple_forecast(daily_insights.tail(30), "E/P")
        forecast_et = simple_forecast(daily_insights.tail(30), "E/T")
        f1, f2 = st.columns(2)
        f1.metric("Expected E/P", f"{forecast_ep:,.2f}" if forecast_ep is not None else "Waiting")
        f2.metric("Expected E/T", f"{forecast_et:,.2f}" if forecast_et is not None else "Waiting")

        st.subheader("Draw Compliance Calendar")
        calendar_df = daily_insights.copy()
        calendar_df["Report Date"] = pd.to_datetime(calendar_df["Report Date"], errors="coerce")
        calendar_df = calendar_df.dropna(subset=["Report Date"])
        if not calendar_df.empty:
            calendar_df["Day"] = calendar_df["Report Date"].dt.strftime("%Y-%m-%d")
            fig = px.scatter(
                calendar_df,
                x="Day",
                y="Poly",
                size="Draw Compliance Score",
                color="Score Label",
                hover_data=["E/P", "E/T", "Status", "Tonnes"],
                title="Daily Draw Compliance Calendar"
            )
            fig.update_layout(height=500)
            st.plotly_chart(fig, width="stretch")

        st.subheader("Hotspot / Concern Ranking")
        ranking_df = daily_insights.copy()
        ranking_df["Concern Level"] = 100 - pd.to_numeric(ranking_df["Draw Compliance Score"], errors="coerce")
        ranking_df = ranking_df.sort_values("Concern Level", ascending=False)
        st.dataframe(
            ranking_df[[col for col in ["Report Date", "Poly", "New Events", "Energy J", "Potency m3", "Tonnes", "E/P", "E/T", "Status", "Draw Compliance Score", "Score Label"] if col in ranking_df.columns]].head(20),
            width="stretch"
        )

        st.subheader("Automatic Executive Summary")
        summary_text = management_summary(
            latest.get("Poly", "Selected area"),
            to_int(latest.get("New Events", latest.get("Events", 0))),
            to_float(latest.get("Energy J", 0.0)),
            to_float(latest.get("Potency m3", 0.0)),
            to_float(latest.get("Tonnes", 0.0)),
            to_float(latest.get("E/P", 0.0)),
            None if str(latest.get("E/T", "")).strip() == "" else to_float(latest.get("E/T", 0.0)),
            latest_status,
            latest_score,
        )
        st.text_area("Copy-ready management summary", value=summary_text, height=170)

    if not weekly_insights.empty:
        st.subheader("Saved Weekly Summaries")
        st.dataframe(weekly_insights.tail(10), width="stretch")



with tab_red_rating:
    st.header("Red Rating Management")
    st.caption("Operational time-loss, hotspot and corrective-action dashboard. All figures come from saved or uploaded records.")

    capture_tab, excel_tab, dashboard_tab, history_tab = st.tabs([
        "➕ Capture Incident", "📥 Excel Upload", "📊 Management Dashboard", "🗂️ History"
    ])

    with capture_tab:
        st.markdown('<div class="section-bar">INCIDENT CAPTURE</div>', unsafe_allow_html=True)
        with st.form("red_rating_capture_form_v2", clear_on_submit=True):
            r1, r2, r3, r4 = st.columns(4)
            incident_date = r1.date_input("Date", value=datetime.now().date(), key="rr2_date")
            start_time = r2.time_input("Start time", value=datetime.now().replace(second=0, microsecond=0).time(), key="rr2_start")
            end_time = r3.time_input("End time", value=datetime.now().replace(second=0, microsecond=0).time(), key="rr2_end")
            shift = r4.selectbox("Shift", ["Day", "Afternoon", "Night", "Unspecified"], key="rr2_shift")

            locations_text = st.text_area(
                "Affected locations",
                placeholder="Example: II-UXC, II-SI, II-UX14",
                help="Enter one or several locations separated by commas. Each location is analysed individually in the hotspot dashboard.",
                key="rr2_locations",
            )

            c1, c2, c3 = st.columns(3)
            category = c1.selectbox("Root cause / category", ["Seismic", "Ground Support", "Ventilation", "Mechanical", "Electrical", "Access / Travel", "Communication / System", "Other", "Unspecified"], key="rr2_category")
            activity = c2.selectbox("Activity affected", ["Production", "Development", "Support", "Inspection", "Travel / Access", "Other", "Unspecified"], key="rr2_activity")
            department = c3.selectbox("Department / function", ["Geotechnical", "Mining", "Engineering", "Ventilation", "Survey", "Operations", "Other", "Unspecified"], key="rr2_department")

            x1, x2, x3 = st.columns(3)
            status = x1.selectbox("Status", ["Open", "Under Review", "Closed"], key="rr2_status")
            responsible = x2.text_input("Responsible person", value=current_user if 'current_user' in globals() else "", key="rr2_responsible")
            trigger = x3.text_input("Trigger / seismic event", key="rr2_trigger")

            action_taken = st.text_area("Corrective action / comments", height=90, key="rr2_action")
            outcome = st.text_area("Outcome / reopening condition", height=70, placeholder="Example: Inspection completed; workplace reopened after 2 hours.", key="rr2_outcome")
            submit_red = st.form_submit_button("Save red-rating incident", type="primary", use_container_width=True)

        minutes, hours = calculate_red_rating_duration(start_time, end_time)
        st.info(f"Calculated time lost: **{minutes // 60} h {minutes % 60} min** ({hours:.2f} h). Overnight restrictions are handled automatically.")

        if submit_red:
            final_locations = [x.strip() for x in locations_text.split(",") if x.strip()]
            if not final_locations:
                st.error("Enter at least one affected location.")
            elif minutes <= 0:
                st.error("End time must be different from start time.")
            else:
                record = {
                    "Record ID": datetime.now().strftime("RR-%Y%m%d-%H%M%S"),
                    "Date": str(incident_date), "Start Time": start_time.strftime("%H:%M"), "End Time": end_time.strftime("%H:%M"),
                    "Duration Minutes": minutes, "Duration Hours": round(hours, 3), "Locations": ", ".join(final_locations),
                    "Category": category, "Activity Affected": activity, "Department / Function": department, "Shift": shift,
                    "Trigger": trigger, "Action / Comments": action_taken, "Outcome": outcome, "Responsible Person": responsible,
                    "Status": status, "Entered By": current_user if 'current_user' in globals() else "Unknown",
                    "Timestamp": datetime.now().isoformat(timespec="seconds"),
                }
                save_red_rating_record(record)
                if 'write_audit' in globals():
                    write_audit(record["Entered By"], current_role if 'current_role' in globals() else "Unknown", "Saved red-rating record", record["Record ID"])
                st.success("Red-rating incident saved.")

    with excel_tab:
        st.subheader("Upload the approved Excel Red Rating Tracker")
        st.caption("Tailored for the 'Red Rating Delay Register' sheet. Existing incident data is imported; no future or fabricated records are created.")
        red_excel = st.file_uploader("Upload Red_Rating_Tracker_July_Corrected_II_UXC.xlsx", type=["xlsx", "xls"], key="red_rating_excel_upload_v2")
        if red_excel is not None:
            imported_red, import_warnings = read_red_rating_excel(red_excel)
            if imported_red.empty:
                st.error("No valid red-rating records were found in the workbook.")
            else:
                st.success(f"{len(imported_red)} valid incident record(s) found.")
                preview_cols = [c for c in ["Date", "Start Time", "End Time", "Duration Hours", "Locations", "Category", "Status"] if c in imported_red.columns]
                st.dataframe(imported_red[preview_cols], width="stretch", hide_index=True)
                total_import_hours = pd.to_numeric(imported_red["Duration Hours"], errors="coerce").fillna(0).sum()
                p1, p2 = st.columns(2)
                p1.metric("Records ready", len(imported_red)); p2.metric("Total hours", f"{total_import_hours:.2f} h")
                for warning in import_warnings: st.warning(warning)
                if st.button("Import Excel records into GeoIntel", type="primary", use_container_width=True, key="import_red_excel_button_v2"):
                    merged = merge_red_rating_import(imported_red)
                    if 'write_audit' in globals():
                        write_audit(current_user if 'current_user' in globals() else "Unknown", current_role if 'current_role' in globals() else "Unknown", "Imported red-rating Excel workbook", f"{len(imported_red)} records")
                    st.success(f"Import complete. GeoIntel now contains {len(merged)} red-rating record(s).")
                    st.rerun()

    red_df = prepare_red_rating_data(load_csv(RED_RATING_FILE))

    with dashboard_tab:
        if red_df.empty:
            st.info("No red-rating records have been saved or imported yet.")
        else:
            st.markdown('<div class="section-bar">RED RATING KPI DASHBOARD</div>', unsafe_allow_html=True)
            s1, s2 = st.columns(2)
            production_rate = s1.number_input("Production rate (tonnes/hour)", min_value=0.0, value=500.0, step=10.0, key="rr_prod_rate")
            value_per_tonne = s2.number_input("Estimated value per tonne (R)", min_value=0.0, value=0.0, step=1.0, key="rr_value_tonne")

            today = datetime.now().date()
            month_start = today.replace(day=1)
            dated = red_df.dropna(subset=["Date"]).copy()
            today_df = dated[dated["Date"].dt.date == today]
            month_df = dated[(dated["Date"].dt.date >= month_start) & (dated["Date"].dt.date <= today)]
            month_hours = float(month_df["Duration Hours"].sum())
            tonnes_lost = month_hours * production_rate
            revenue_risk = tonnes_lost * value_per_tonne
            active = int(red_df["Status"].isin(["Open", "Under Review"]).sum())

            k1, k2, k3, k4 = st.columns(4)
            k1.metric("Active restrictions", active)
            k2.metric("Hours lost this month", f"{month_hours:.2f} h", f"Today {today_df['Duration Hours'].sum():.2f} h")
            k3.metric("Estimated tonnes lost", f"{tonnes_lost:,.0f} t")
            k4.metric("Revenue at risk", f"R {revenue_risk:,.0f}" if value_per_tonne > 0 else "Set value/t")

            loc = explode_red_locations(red_df)
            loc_summary = loc.groupby("Location", as_index=False).agg(
                Red_Ratings=("Location", "size"), Hours_Lost=("Duration Hours", "sum")
            ).sort_values(["Red_Ratings", "Hours_Lost"], ascending=False)
            loc_summary["Estimated Tonnes Lost"] = loc_summary["Hours_Lost"] * production_rate
            loc_summary["Reliability Score"] = loc_summary.apply(lambda r: reliability_score(r["Red_Ratings"], r["Hours_Lost"]), axis=1)
            loc_summary["Repeat Hotspot"] = loc_summary["Red_Ratings"].map(lambda n: "⚠ Repeat hotspot" if n >= 3 else "")

            left, right = st.columns(2)
            with left:
                st.subheader("Live Hotspot Ranking")
                st.dataframe(loc_summary.head(10), width="stretch", hide_index=True)
            with right:
                st.subheader("Cause Analysis")
                cause = red_df.groupby("Category", as_index=False)["Duration Hours"].sum().sort_values("Duration Hours", ascending=False)
                st.plotly_chart(px.pie(cause, names="Category", values="Duration Hours", hole=0.45, title="Hours lost by cause"), width="stretch")

            c1, c2 = st.columns(2)
            with c1:
                st.subheader("Monthly Time-Loss Trend")
                monthly = dated.assign(Month=dated["Date"].dt.to_period("M").astype(str)).groupby("Month", as_index=False).agg(Incidents=("Record ID", "count"), Hours_Lost=("Duration Hours", "sum"))
                st.plotly_chart(px.bar(monthly, x="Month", y="Hours_Lost", text="Incidents", title="Hours lost by month"), width="stretch")
            with c2:
                st.subheader("Shift Analysis")
                shift_summary = red_df.groupby("Shift", as_index=False)["Duration Hours"].sum().sort_values("Duration Hours", ascending=False)
                st.plotly_chart(px.bar(shift_summary, x="Shift", y="Duration Hours", title="Hours lost by shift"), width="stretch")

            d1, d2 = st.columns(2)
            with d1:
                st.subheader("Department / Function")
                dept = red_df.groupby("Department / Function", as_index=False)["Duration Hours"].sum().sort_values("Duration Hours", ascending=False)
                st.plotly_chart(px.bar(dept, x="Duration Hours", y="Department / Function", orientation="h", title="Hours lost by function"), width="stretch")
            with d2:
                st.subheader("Open and Under-Review Incidents")
                open_df = red_df[red_df["Status"].isin(["Open", "Under Review"])].copy()
                show_cols = [c for c in ["Date", "Locations", "Duration Hours", "Category", "Responsible Person", "Status", "Outcome"] if c in open_df.columns]
                st.dataframe(open_df[show_cols].sort_values("Date", ascending=False), width="stretch", hide_index=True)

            if not loc_summary.empty:
                repeated = loc_summary[loc_summary["Red_Ratings"] >= 3]
                if not repeated.empty:
                    st.warning("Recurring hotspot warning: " + ", ".join(f"{r.Location} ({int(r.Red_Ratings)} ratings)" for r in repeated.itertuples()))

    with history_tab:
        if red_df.empty:
            st.info("No red-rating history is available.")
        else:
            h1, h2, h3 = st.columns(3)
            valid_dates = red_df["Date"].dropna()
            default_from = valid_dates.min().date() if not valid_dates.empty else datetime.now().date()
            date_filter = h1.date_input("From date", value=default_from, key="rr2_from")
            status_filter = h2.multiselect("Status", sorted(red_df["Status"].dropna().unique()), key="rr2_status_filter")
            search_location = h3.text_input("Location contains", key="rr2_location_search")
            filtered = red_df[red_df["Date"].dt.date >= date_filter].copy()
            if status_filter: filtered = filtered[filtered["Status"].isin(status_filter)]
            if search_location.strip(): filtered = filtered[filtered["Locations"].str.contains(search_location.strip(), case=False, na=False)]
            show_cols = [c for c in ["Date", "Start Time", "End Time", "Duration Hours", "Locations", "Category", "Activity Affected", "Department / Function", "Shift", "Status", "Responsible Person", "Action / Comments", "Outcome"] if c in filtered.columns]
            st.dataframe(filtered.sort_values(["Date", "Start Time"], ascending=False)[show_cols], width="stretch", hide_index=True)
            st.download_button("Download filtered CSV", filtered.to_csv(index=False).encode("utf-8"), "GeoIntel_red_rating_history.csv", "text/csv", use_container_width=True)


with tab_history:
    st.header("History / Admin")

    daily_history = safe_sort_history(load_csv(DAILY_FILE))
    weekend_history = safe_sort_history(load_csv(WEEKEND_FILE))
    weekend_derived_history = safe_sort_history(load_csv(WEEKEND_DERIVED_FILE))
    weekly_history = safe_sort_history(load_csv(WEEKLY_FILE))

    st.subheader("Daily Draw History")
    if daily_history.empty:
        st.info("No daily records saved yet.")
    else:
        st.dataframe(daily_history, width="stretch")

        selected_daily_row = st.selectbox(
            "Select Daily Record",
            options=list(range(len(daily_history))),
            format_func=lambda x: f"{daily_history.iloc[x].get('Report Date', '')} - {daily_history.iloc[x].get('Poly', '')}",
            key="daily_selected_row",
        )

        daily_row = daily_history.iloc[selected_daily_row]

        with st.expander("Edit Selected Daily Record", expanded=False):
            with st.form("edit_daily_form"):
                edit_report_date = st.text_input("Report Date", value=str(row_value(daily_row, "Report Date")), key="edit_daily_report_date")
                edit_poly = st.text_input("Poly", value=str(row_value(daily_row, "Poly")), key="edit_daily_poly")
                edit_events = st.number_input("New Events", min_value=0, value=to_int(row_value(daily_row, "New Events", 0)), step=1, key="edit_daily_events")
                edit_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(daily_row, "Energy J", 0.0)), step=1000.0, key="edit_daily_energy")
                edit_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(daily_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_daily_potency")
                edit_tonnes = st.number_input("Tonnes", min_value=0.0, value=to_float(row_value(daily_row, "Tonnes", 0.0)), step=1.0, key="edit_daily_tonnes")
                save_daily_edit = st.form_submit_button("Save Edited Daily Record")

            if save_daily_edit:
                updated_record = {
                    "Report Date": edit_report_date,
                    "Poly": edit_poly,
                    "New Events": edit_events,
                    "Energy J": edit_energy,
                    "Potency m3": edit_potency,
                    "Tonnes": edit_tonnes,
                }
                updated_record = recalc_draw_record(updated_record, ep_threshold, et_threshold)

                daily_history = daily_history.drop(daily_history.index[selected_daily_row])
                daily_history = pd.concat([daily_history, pd.DataFrame([updated_record])], ignore_index=True)
                daily_history = daily_history.drop_duplicates(subset=["Report Date", "Poly"], keep="last")
                daily_history = safe_sort_history(daily_history)
                save_csv(daily_history, DAILY_FILE)
                st.success("Daily record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Daily Record"):
            daily_history = daily_history.drop(daily_history.index[selected_daily_row])
            save_csv(daily_history, DAILY_FILE)
            st.success("Daily record deleted.")
            st.rerun()

        st.download_button(
            "Download Daily History CSV",
            daily_history.to_csv(index=False).encode("utf-8"),
            "gdi_history.csv",
            "text/csv",
        )

    st.subheader("Weekend Cumulative History")
    if weekend_history.empty:
        st.info("No weekend cumulative records saved yet.")
    else:
        st.dataframe(weekend_history, width="stretch")

        selected_weekend_row = st.selectbox(
            "Select Weekend Cumulative Record",
            options=list(range(len(weekend_history))),
            format_func=lambda x: f"{weekend_history.iloc[x].get('Weekend Start', '')} - {weekend_history.iloc[x].get('Poly', '')} - {weekend_history.iloc[x].get('Stage', '')}",
            key="weekend_selected_row",
        )

        weekend_row = weekend_history.iloc[selected_weekend_row]

        with st.expander("Edit Selected Weekend Cumulative Record", expanded=False):
            with st.form("edit_weekend_form"):
                edit_weekend_start = st.text_input("Weekend Start", value=str(row_value(weekend_row, "Weekend Start")), key="edit_weekend_start")
                edit_weekend_poly = st.text_input("Poly", value=str(row_value(weekend_row, "Poly")), key="edit_weekend_poly")
                edit_weekend_stage = st.selectbox(
                    "Stage",
                    ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"],
                    index=["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"].index(row_value(weekend_row, "Stage", "Friday Baseline")) if row_value(weekend_row, "Stage", "Friday Baseline") in ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"] else 0,
                    key="edit_weekend_stage",
                )
                edit_weekend_events = st.number_input("Events", min_value=0, value=to_int(row_value(weekend_row, "Events", 0)), step=1, key="edit_weekend_events")
                edit_weekend_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(weekend_row, "Energy J", 0.0)), step=1000.0, key="edit_weekend_energy")
                edit_weekend_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(weekend_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_weekend_potency")
                edit_weekend_last_event = st.text_input("Last Event", value=str(row_value(weekend_row, "Last Event")), key="edit_weekend_last_event")
                save_weekend_edit = st.form_submit_button("Save Edited Weekend Cumulative Record")

            if save_weekend_edit:
                stage_order_edit = {
                    "Friday Baseline": 0,
                    "Friday to Saturday": 1,
                    "Friday to Sunday": 2,
                    "Friday to Monday": 3,
                }
                updated_weekend_record = {
                    "Weekend Start": edit_weekend_start,
                    "Poly": edit_weekend_poly,
                    "Stage": edit_weekend_stage,
                    "Stage Order": stage_order_edit[edit_weekend_stage],
                    "Events": edit_weekend_events,
                    "Energy J": edit_weekend_energy,
                    "Potency m3": edit_weekend_potency,
                    "Last Event": edit_weekend_last_event,
                }

                weekend_history = weekend_history.drop(weekend_history.index[selected_weekend_row])
                weekend_history = pd.concat([weekend_history, pd.DataFrame([updated_weekend_record])], ignore_index=True)
                weekend_history = weekend_history.drop_duplicates(subset=["Weekend Start", "Poly", "Stage"], keep="last")
                weekend_history = safe_sort_history(weekend_history)
                save_csv(weekend_history, WEEKEND_FILE)
                st.success("Weekend cumulative record edited.")
                st.rerun()

        if st.button("Delete Selected Weekend Cumulative Record"):
            weekend_history = weekend_history.drop(weekend_history.index[selected_weekend_row])
            save_csv(weekend_history, WEEKEND_FILE)
            st.success("Weekend cumulative record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Cumulative CSV",
            weekend_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_cumulative.csv",
            "text/csv",
        )

    st.subheader("Weekend Derived Daily History")
    if weekend_derived_history.empty:
        st.info("No weekend derived records saved yet.")
    else:
        st.dataframe(weekend_derived_history, width="stretch")

        selected_derived_row = st.selectbox(
            "Select Weekend Derived Record",
            options=list(range(len(weekend_derived_history))),
            format_func=lambda x: f"{weekend_derived_history.iloc[x].get('Report Date', weekend_derived_history.iloc[x].get('Weekend Start', ''))} - {weekend_derived_history.iloc[x].get('Derived Day', '')} - {weekend_derived_history.iloc[x].get('Poly', '')}",
            key="derived_selected_row",
        )

        derived_row = weekend_derived_history.iloc[selected_derived_row]

        with st.expander("Edit Selected Weekend Derived Record", expanded=False):
            with st.form("edit_derived_form"):
                edit_derived_start = st.text_input("Weekend Start", value=str(row_value(derived_row, "Weekend Start")), key="edit_derived_start")
                edit_derived_day = st.selectbox(
                    "Derived Day",
                    ["Friday", "Saturday", "Sunday", "Monday"],
                    index=["Friday", "Saturday", "Sunday", "Monday"].index(row_value(derived_row, "Derived Day", "Friday")) if row_value(derived_row, "Derived Day", "Friday") in ["Friday", "Saturday", "Sunday", "Monday"] else 0,
                    key="edit_derived_day",
                )
                edit_derived_report_date = st.text_input("Report Date", value=str(row_value(derived_row, "Report Date", weekend_actual_date(edit_derived_start, edit_derived_day))), key="edit_derived_report_date")
                edit_derived_poly = st.text_input("Poly", value=str(row_value(derived_row, "Poly")), key="edit_derived_poly")
                edit_derived_events = st.number_input("Events", min_value=0, value=to_int(row_value(derived_row, "Events", 0)), step=1, key="edit_derived_events")
                edit_derived_energy = st.number_input("Energy J", min_value=0.0, value=to_float(row_value(derived_row, "Energy J", 0.0)), step=1000.0, key="edit_derived_energy")
                edit_derived_potency = st.number_input("Potency m3", min_value=0.0, value=to_float(row_value(derived_row, "Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_derived_potency")
                edit_derived_tonnes = st.number_input("Tonnes", min_value=0.0, value=to_float(row_value(derived_row, "Tonnes", 0.0)), step=1.0, key="edit_derived_tonnes")
                sync_daily = st.checkbox("Also update Daily Draw History", value=True, key="sync_derived_to_daily")
                save_derived_edit = st.form_submit_button("Save Edited Weekend Derived Record")

            if save_derived_edit:
                updated_derived_record = {
                    "Weekend Start": edit_derived_start,
                    "Derived Day": edit_derived_day,
                    "Report Date": edit_derived_report_date,
                    "Poly": edit_derived_poly,
                    "Events": edit_derived_events,
                    "Energy J": edit_derived_energy,
                    "Potency m3": edit_derived_potency,
                    "Tonnes": edit_derived_tonnes,
                }
                updated_derived_record = recalc_draw_record(updated_derived_record, ep_threshold, et_threshold)

                weekend_derived_history = weekend_derived_history.drop(weekend_derived_history.index[selected_derived_row])
                weekend_derived_history = pd.concat([weekend_derived_history, pd.DataFrame([updated_derived_record])], ignore_index=True)
                weekend_derived_history = weekend_derived_history.drop_duplicates(subset=["Weekend Start", "Derived Day", "Poly"], keep="last")
                weekend_derived_history = safe_sort_history(weekend_derived_history)
                save_csv(weekend_derived_history, WEEKEND_DERIVED_FILE)

                if sync_daily:
                    daily_record = {
                        "Report Date": edit_derived_report_date,
                        "Poly": edit_derived_poly,
                        "New Events": edit_derived_events,
                        "Energy J": edit_derived_energy,
                        "Potency m3": edit_derived_potency,
                        "Tonnes": edit_derived_tonnes,
                    }
                    daily_record = recalc_draw_record(daily_record, ep_threshold, et_threshold)
                    save_or_update(DAILY_FILE, daily_record, ["Report Date", "Poly"])

                st.success("Weekend derived record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Weekend Derived Record"):
            weekend_derived_history = weekend_derived_history.drop(weekend_derived_history.index[selected_derived_row])
            save_csv(weekend_derived_history, WEEKEND_DERIVED_FILE)
            st.success("Weekend derived record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Derived CSV",
            weekend_derived_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_derived_daily.csv",
            "text/csv",
        )

    st.subheader("Weekly Summary History")
    if weekly_history.empty:
        st.info("No weekly records saved yet.")
    else:
        st.dataframe(weekly_history, width="stretch")

        selected_weekly_row = st.selectbox(
            "Select Weekly Record",
            options=list(range(len(weekly_history))),
            format_func=lambda x: f"{weekly_history.iloc[x].get('Week Range', '')}",
            key="weekly_selected_row",
        )

        weekly_row = weekly_history.iloc[selected_weekly_row]

        with st.expander("Edit Selected Weekly Record", expanded=False):
            with st.form("edit_weekly_form"):
                edit_week_range = st.text_input("Week Range", value=str(row_value(weekly_row, "Week Range")), key="edit_week_range")
                edit_mine_events = st.number_input("Mine-wide Events", min_value=0, value=to_int(row_value(weekly_row, "Mine-wide Events", 0)), step=1, key="edit_mine_events")
                edit_largest_ml = st.number_input("Largest ML", value=to_float(row_value(weekly_row, "Largest ML", 0.0)), step=0.1, key="edit_largest_ml")
                edit_dominant = st.text_input("Dominant Hotspot", value=str(row_value(weekly_row, "Dominant Hotspot")), key="edit_dominant")
                edit_second = st.text_input("Second Hotspot", value=str(row_value(weekly_row, "Second Hotspot")), key="edit_second")
                edit_liie_events = st.number_input("Lift II East Events", min_value=0, value=to_int(row_value(weekly_row, "Lift II East Events", 0)), step=1, key="edit_liie_events")
                edit_liie_energy = st.number_input("Lift II East Energy J", min_value=0.0, value=to_float(row_value(weekly_row, "Lift II East Energy J", 0.0)), step=1000.0, key="edit_liie_energy")
                edit_liie_potency = st.number_input("Lift II East Potency m3", min_value=0.0, value=to_float(row_value(weekly_row, "Lift II East Potency m3", 0.0)), step=0.001, format="%.6f", key="edit_liie_potency")
                edit_weekly_tonnes = st.number_input("Weekly Tonnes", min_value=0.0, value=to_float(row_value(weekly_row, "Weekly Tonnes", 0.0)), step=1.0, key="edit_weekly_tonnes")
                edit_weekly_comment = st.text_area("Weekly Comment", value=str(row_value(weekly_row, "Weekly Comment")), height=120, key="edit_weekly_comment")
                edit_weekly_draw_comment = st.text_area("Weekly Draw Comment", value=str(row_value(weekly_row, "Weekly Draw Comment")), height=120, key="edit_weekly_draw_comment")
                save_weekly_edit = st.form_submit_button("Save Edited Weekly Record")

            if save_weekly_edit:
                weekly_ep = edit_liie_energy / edit_liie_potency if edit_liie_potency > 0 else 0.0
                weekly_et = edit_liie_energy / edit_weekly_tonnes if edit_weekly_tonnes > 0 else None
                weekly_status = get_status(weekly_ep, weekly_et, ep_threshold, et_threshold)
                weekly_health = cave_health_score(weekly_status)

                updated_weekly_record = {
                    "Week Range": edit_week_range,
                    "Mine-wide Events": edit_mine_events,
                    "Largest ML": edit_largest_ml,
                    "Dominant Hotspot": edit_dominant,
                    "Second Hotspot": edit_second,
                    "Lift II East Events": edit_liie_events,
                    "Lift II East Energy J": edit_liie_energy,
                    "Lift II East Potency m3": edit_liie_potency,
                    "Weekly Tonnes": edit_weekly_tonnes,
                    "Weekly E/P": weekly_ep,
                    "Weekly E/T": weekly_et if weekly_et is not None else "",
                    "Weekly Status": weekly_status,
                    "Weekly Cave Health Score": weekly_health,
                    "Weekly Comment": edit_weekly_comment,
                    "Weekly Draw Comment": edit_weekly_draw_comment,
                }

                weekly_history = weekly_history.drop(weekly_history.index[selected_weekly_row])
                weekly_history = pd.concat([weekly_history, pd.DataFrame([updated_weekly_record])], ignore_index=True)
                weekly_history = weekly_history.drop_duplicates(subset=["Week Range"], keep="last")
                weekly_history = safe_sort_history(weekly_history)
                save_csv(weekly_history, WEEKLY_FILE)
                st.success("Weekly record edited and recalculated.")
                st.rerun()

        if st.button("Delete Selected Weekly Record"):
            weekly_history = weekly_history.drop(weekly_history.index[selected_weekly_row])
            save_csv(weekly_history, WEEKLY_FILE)
            st.success("Weekly record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekly CSV",
            weekly_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekly_history.csv",
            "text/csv",
        )
