import re
from io import BytesIO
from pathlib import Path
from datetime import datetime, timedelta

import pandas as pd
import pdfplumber
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document


st.set_page_config(page_title="Geotechnical Draw Intelligence", layout="wide")

DAILY_FILE = Path("gdi_history.csv")
WEEKEND_FILE = Path("gdi_weekend_cumulative.csv")
WEEKEND_DERIVED_FILE = Path("gdi_weekend_derived_daily.csv")
WEEKLY_FILE = Path("gdi_weekly_history.csv")

DEFAULT_EP_THRESHOLD = 10000.0
DEFAULT_ET_THRESHOLD = 10.0


def find_value(pattern, text):
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else None


def pdf_to_text(uploaded_pdf):
    text = ""
    with pdfplumber.open(uploaded_pdf) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def load_csv(path):
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


def save_csv(df, path):
    df.to_csv(path, index=False)


def save_or_update(path, record, key_cols):
    df = load_csv(path)
    record_df = pd.DataFrame([record])

    if df.empty:
        df = record_df
    else:
        df = pd.concat([df, record_df], ignore_index=True)
        df = df.drop_duplicates(subset=key_cols, keep="last")

    save_csv(df, path)




def weekend_actual_date(weekend_start, derived_day):
    """Return the real calendar date for Friday/Saturday/Sunday/Monday derived data."""
    offset_map = {
        "Friday": 0,
        "Saturday": 1,
        "Sunday": 2,
        "Monday": 3,
    }

    try:
        base_date = datetime.strptime(str(weekend_start), "%Y-%m-%d")
    except ValueError:
        return str(weekend_start)

    return (base_date + timedelta(days=offset_map.get(derived_day, 0))).strftime("%Y-%m-%d")


def safe_sort_history(df):
    if df.empty:
        return df
    df = df.copy()
    sort_candidates = ["Report Date", "Weekend Start", "Week Range"]
    for col in sort_candidates:
        if col in df.columns:
            return df.sort_values(col, kind="stable")
    return df

def clean_number(value):
    if value is None:
        return None
    value = str(value).replace(",", "").strip()
    try:
        return float(value)
    except ValueError:
        return None


def clean_report_date(last_event):
    if not last_event:
        return datetime.now().strftime("%Y-%m-%d")

    match = re.search(r"(\d{1,2})\s+([A-Za-z]{3})", last_event)
    if not match:
        return datetime.now().strftime("%Y-%m-%d")

    day = int(match.group(1))
    month_text = match.group(2).title()

    month_map = {
        "Jan": 1, "Feb": 2, "Mar": 3, "Apr": 4,
        "May": 5, "Jun": 6, "Jul": 7, "Aug": 8,
        "Sep": 9, "Oct": 10, "Nov": 11, "Dec": 12
    }

    month = month_map.get(month_text, datetime.now().month)
    year = datetime.now().year
    return f"{year}-{month:02d}-{day:02d}"


def parse_poly_sections(text):
    poly_matches = list(re.finditer(r"Poly:\s*([A-Za-z0-9_]+)", text))
    results = {}

    for i, match in enumerate(poly_matches):
        poly = match.group(1)
        start = match.start()
        end = poly_matches[i + 1].start() if i + 1 < len(poly_matches) else len(text)
        block = text[start:end]

        last_event = find_value(r"Last event Date and Time:\s*(.+)", block)
        events = find_value(r"Number of new events:\s*(\d+)", block)
        energy = find_value(r"Total Energy:\s*([0-9.eE+-]+)", block)
        potency = find_value(r"Total Potency:\s*([0-9.eE+-]+)", block)
        current_rate = find_value(r"Current normalized activity rate:\s*([0-9.]+)", block)
        medium_rate = find_value(r"Medium term normalized activity rate:\s*([0-9.]+)", block)

        comment = ""
        comment_match = re.search(
            r"Medium term normalized activity rate:\s*[0-9.]+\s*(.*?)(?:\n\d+\s+of\s+\d+|\Z)",
            block,
            re.DOTALL
        )
        if comment_match:
            comment = " ".join(comment_match.group(1).split())

        if events or energy or potency:
            results[poly] = {
                "Poly": poly,
                "Last Event": last_event,
                "Events": int(events) if events else 0,
                "Energy J": clean_number(energy) or 0.0,
                "Potency m3": clean_number(potency) or 0.0,
                "Current Activity Rate": clean_number(current_rate),
                "Medium Activity Rate": clean_number(medium_rate),
                "Comment": comment,
            }

    return results


def parse_week_range(text):
    match = re.search(r"Seismic Report:\s*(.+)", text)
    if match:
        return match.group(1).strip()
    return ""


def parse_mine_total(text):
    total = find_value(r"Total number of events\s*=\s*(\d+)", text)
    return int(total) if total else 0


def parse_largest_magnitude(text):
    mags = re.findall(r"\b(?:ml|ML|M)(-?\d+(?:\.\d+)?)", text)
    values = [float(m) for m in mags if m is not None]
    return max(values) if values else None


def get_status(ep, et, ep_threshold, et_threshold):
    if et is None:
        return "Waiting for tonnes"
    if ep < ep_threshold and et < et_threshold:
        return "Efficient Draw"
    if ep >= ep_threshold and et < et_threshold:
        return "Brittle but Efficient"
    if ep < ep_threshold and et >= et_threshold:
        return "Ductile but Costly"
    return "High-Risk Draw"


def cave_health_score(status):
    return {
        "Efficient Draw": 90,
        "Brittle but Efficient": 65,
        "Ductile but Costly": 55,
        "High-Risk Draw": 30,
        "Waiting for tonnes": 0,
    }.get(status, 0)


def recommended_action(status):
    return {
        "Efficient Draw": "Maintain current draw strategy and continue routine monitoring.",
        "Brittle but Efficient": "Continue draw, but monitor seismic response closely due to elevated E/P.",
        "Ductile but Costly": "Review draw distribution and tonnes efficiency because seismic cost per tonne is elevated.",
        "High-Risk Draw": "Escalate for geotechnical review. Monitor draw, seismicity and nearby workplaces closely.",
        "Waiting for tonnes": "Enter tonnes drawn to finalise the draw compliance classification.",
    }.get(status, "")


def plot_matrix(ep, et, history, ep_threshold, et_threshold):
    max_x = max(et_threshold * 2, et * 1.4)
    max_y = max(ep_threshold * 2, ep * 1.4)

    fig = go.Figure()

    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=0, y1=ep_threshold,
                  fillcolor="green", opacity=0.12, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=0, y1=ep_threshold,
                  fillcolor="orange", opacity=0.14, line_width=0)
    fig.add_shape(type="rect", x0=0, x1=et_threshold, y0=ep_threshold, y1=max_y,
                  fillcolor="gold", opacity=0.16, line_width=0)
    fig.add_shape(type="rect", x0=et_threshold, x1=max_x, y0=ep_threshold, y1=max_y,
                  fillcolor="red", opacity=0.12, line_width=0)

    fig.add_shape(type="line", x0=et_threshold, x1=et_threshold, y0=0, y1=max_y,
                  line=dict(dash="dash", width=2))
    fig.add_shape(type="line", x0=0, x1=max_x, y0=ep_threshold, y1=ep_threshold,
                  line=dict(dash="dash", width=2))

    if not history.empty and "E/P" in history.columns and "E/T" in history.columns:
        hist = history.copy()
        hist["E/P"] = pd.to_numeric(hist["E/P"], errors="coerce")
        hist["E/T"] = pd.to_numeric(hist["E/T"], errors="coerce")
        hist = hist.dropna(subset=["E/P", "E/T"])

        if not hist.empty:
            fig.add_trace(go.Scatter(
                x=hist["E/T"],
                y=hist["E/P"],
                mode="markers",
                name="Saved Records",
                text=hist["Report Date"] if "Report Date" in hist.columns else None,
                hovertemplate="Date: %{text}<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
            ))

    fig.add_trace(go.Scatter(
        x=[et],
        y=[ep],
        mode="markers+text",
        name="Current",
        text=["Current"],
        textposition="top center",
        marker=dict(size=22),
        hovertemplate="Current<br>E/T: %{x:.2f}<br>E/P: %{y:.2f}<extra></extra>",
    ))

    fig.add_annotation(x=et_threshold * 0.45, y=ep_threshold * 0.5, text="Efficient Draw", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 0.5, text="Ductile but Costly", showarrow=False)
    fig.add_annotation(x=et_threshold * 0.5, y=ep_threshold * 1.45, text="Brittle but Efficient", showarrow=False)
    fig.add_annotation(x=et_threshold * 1.45, y=ep_threshold * 1.45, text="High-Risk Draw", showarrow=False)

    fig.update_layout(
        title="Draw Compliance Matrix",
        xaxis_title="Energy / Tonnes (E/T)",
        yaxis_title="Energy / Potency (E/P)",
        xaxis=dict(range=[0, max_x]),
        yaxis=dict(range=[0, max_y]),
        height=650,
    )

    return fig


def pie_chart(df, title):
    fig = px.pie(df, names="Zone", values="Events", title=title)
    fig.update_traces(textposition="inside", textinfo="percent+label")
    return fig


def bar_chart(df, title):
    fig = px.bar(df.sort_values("Events", ascending=True), x="Events", y="Zone", orientation="h", title=title)
    return fig


def build_word_report(title, comment, record):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("Summary Data", level=1)

    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Comment", level=1)
    doc.add_paragraph(comment)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.sidebar.header("Settings")
ep_threshold = st.sidebar.number_input("E/P Threshold", min_value=0.0, value=DEFAULT_EP_THRESHOLD)
et_threshold = st.sidebar.number_input("E/T Threshold", min_value=0.0, value=DEFAULT_ET_THRESHOLD)

st.title("⛏️ Geotechnical Draw Intelligence (GDI)")
st.markdown("Daily, weekend and weekly seismic draw intelligence platform.")

tab_daily, tab_weekend, tab_weekly, tab_history = st.tabs([
    "Daily Draw Intelligence",
    "Weekend Processor",
    "Weekly Dashboard",
    "History / Admin"
])


with tab_daily:
    st.header("Daily Draw Intelligence")

    uploaded_pdf = st.file_uploader("Upload Daily IMS PDF", type=["pdf"], key="daily_pdf")
    uploaded_excel = st.file_uploader("Optional: Upload Production Tonnes Excel", type=["xlsx", "xls", "csv"], key="daily_excel")

    tonnes_from_excel = None

    if uploaded_excel:
        tonnes_df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith(".csv") else pd.read_excel(uploaded_excel)
        st.dataframe(tonnes_df.head(), width="stretch")
        tonnes_col = st.selectbox("Select Tonnes Column", tonnes_df.columns, key="daily_tonnes_col")
        tonnes_from_excel = pd.to_numeric(tonnes_df[tonnes_col], errors="coerce").sum()
        st.success(f"Tonnes from Excel: {tonnes_from_excel:,.2f}")

    if uploaded_pdf:
        full_text = pdf_to_text(uploaded_pdf)
        poly_data = parse_poly_sections(full_text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index)
            data = poly_data[target_poly]

            st.success(f"{target_poly} section found.")

            c1, c2, c3 = st.columns(3)
            c1.metric("New Events", data["Events"])
            c1.metric("Total Energy (J)", f"{data['Energy J']:,.2f}")
            c2.metric("Total Potency (m³)", f"{data['Potency m3']:,.3f}")
            c2.metric("Current Activity Rate", data["Current Activity Rate"])
            c3.metric("Medium Activity Rate", data["Medium Activity Rate"])
            c3.metric("Last Event", data["Last Event"])

            st.header("Tonnes Input")
            input_mode = st.radio("Tonnes Source", ["Manual Entry", "Excel Upload"], horizontal=True, key="daily_mode")

            if input_mode == "Excel Upload" and tonnes_from_excel is not None:
                tonnes = tonnes_from_excel
            else:
                tonnes = st.number_input("Tonnes Drawn", min_value=0.0, value=0.0, step=1.0, key="daily_tonnes")

            energy = data["Energy J"]
            potency = data["Potency m3"]

            if potency > 0:
                ep = energy / potency
                et = energy / tonnes if tonnes > 0 else None
                status = get_status(ep, et, ep_threshold, et_threshold)
                health = cave_health_score(status)
                action = recommended_action(status)

                st.header("Draw Intelligence Results")
                r1, r2, r3, r4 = st.columns(4)
                r1.metric("E/P", f"{ep:,.2f}")
                r2.metric("E/T", f"{et:,.2f}" if et is not None else "Waiting")
                r3.metric("Status", status)
                r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                st.info(action)

                if et is not None:
                    st.plotly_chart(plot_matrix(ep, et, load_csv(DAILY_FILE), ep_threshold, et_threshold), width="stretch")

                report_date = st.text_input("Report Date", value=clean_report_date(data["Last Event"]), key="daily_date")

                comment = (
                    f"{target_poly} recorded {data['Events']} new seismic events, with total energy of {energy:,.0f} J "
                    f"and total potency of {potency:.3f} m³. The calculated E/P ratio is {ep:,.2f}. "
                )

                if et is not None:
                    comment += f"E/T is {et:,.2f}. {target_poly} plots within the '{status}' quadrant. {action}"
                else:
                    comment += "E/T cannot be finalised until tonnes drawn are entered."

                st.text_area("Management Comment", value=comment, height=170)

                if st.button("Save or Update Daily Record"):
                    record = {
                        "Report Date": report_date,
                        "Poly": target_poly,
                        "New Events": data["Events"],
                        "Energy J": energy,
                        "Potency m3": potency,
                        "Tonnes": tonnes,
                        "E/P": ep,
                        "E/T": et if et is not None else "",
                        "Status": status,
                        "Cave Health Score": health,
                        "Recommended Action": action,
                    }
                    save_or_update(DAILY_FILE, record, ["Report Date", "Poly"])
                    st.success("Daily record saved or updated.")


with tab_weekend:
    st.header("Weekend Cumulative Processor")
    st.write("Save Friday baseline, Fri-Sat, Fri-Sun and Fri-Mon cumulative PDFs. GDI derives each separate day.")

    uploaded_weekend_pdf = st.file_uploader("Upload Weekend Cumulative IMS PDF", type=["pdf"], key="weekend_pdf")

    weekend_start = st.text_input("Weekend Start Date / Friday Date", value=datetime.now().strftime("%Y-%m-%d"))
    stage = st.selectbox(
        "Cumulative Stage",
        ["Friday Baseline", "Friday to Saturday", "Friday to Sunday", "Friday to Monday"]
    )

    stage_order = {
        "Friday Baseline": 0,
        "Friday to Saturday": 1,
        "Friday to Sunday": 2,
        "Friday to Monday": 3,
    }

    stage_daily_name = {
        "Friday Baseline": "Friday",
        "Friday to Saturday": "Saturday",
        "Friday to Sunday": "Sunday",
        "Friday to Monday": "Monday",
    }

    if uploaded_weekend_pdf:
        text = pdf_to_text(uploaded_weekend_pdf)
        poly_data = parse_poly_sections(text)

        if not poly_data:
            st.error("No poly sections found.")
        else:
            poly_list = list(poly_data.keys())
            default_index = poly_list.index("Lift_II_E") if "Lift_II_E" in poly_list else 0
            target_poly = st.selectbox("Select Poly / Zone", poly_list, index=default_index, key="weekend_poly")
            data = poly_data[target_poly]

            st.subheader("Cumulative Values")
            c1, c2, c3 = st.columns(3)
            c1.metric("Cumulative Events", data["Events"])
            c2.metric("Cumulative Energy (J)", f"{data['Energy J']:,.2f}")
            c3.metric("Cumulative Potency (m³)", f"{data['Potency m3']:,.3f}")

            cumulative_record = {
                "Weekend Start": weekend_start,
                "Poly": target_poly,
                "Stage": stage,
                "Stage Order": stage_order[stage],
                "Events": data["Events"],
                "Energy J": data["Energy J"],
                "Potency m3": data["Potency m3"],
                "Last Event": data["Last Event"],
            }

            if st.button("Save Weekend Cumulative Record"):
                save_or_update(WEEKEND_FILE, cumulative_record, ["Weekend Start", "Poly", "Stage"])
                st.success("Weekend cumulative record saved.")

            weekend_df = load_csv(WEEKEND_FILE)

            if not weekend_df.empty:
                filtered = weekend_df[
                    (weekend_df["Weekend Start"].astype(str) == weekend_start)
                    & (weekend_df["Poly"].astype(str) == target_poly)
                ].copy()

                filtered["Stage Order"] = pd.to_numeric(filtered["Stage Order"], errors="coerce")
                filtered = filtered.sort_values("Stage Order")

                st.subheader("Saved Weekend Cumulative Records")
                st.dataframe(filtered, width="stretch")

                current_order = stage_order[stage]

                if current_order == 0:
                    derived_events = data["Events"]
                    derived_energy = data["Energy J"]
                    derived_potency = data["Potency m3"]
                else:
                    previous = filtered[filtered["Stage Order"] == current_order - 1]

                    if previous.empty:
                        st.warning("Previous cumulative stage is missing. Upload/save the previous stage first.")
                        derived_events = None
                        derived_energy = None
                        derived_potency = None
                    else:
                        prev = previous.iloc[-1]
                        derived_events = data["Events"] - int(prev["Events"])
                        derived_energy = data["Energy J"] - float(prev["Energy J"])
                        derived_potency = data["Potency m3"] - float(prev["Potency m3"])

                if derived_events is not None:
                    derived_day = stage_daily_name[stage]

                    st.subheader("Derived Separate Day Values")
                    derived_df = pd.DataFrame([{
                        "Derived Day": derived_day,
                        "Events": derived_events,
                        "Energy J": derived_energy,
                        "Potency m3": derived_potency,
                    }])
                    st.dataframe(derived_df, width="stretch")

                    if derived_energy < 0 or derived_potency < 0 or derived_events < 0:
                        st.error("Derived values are negative. Check the cumulative order.")
                    else:
                        st.success("Derived day calculated successfully.")

                        st.header("Weekend Derived Draw Compliance")
                        weekend_tonnes = st.number_input(
                            f"{derived_day} Tonnes Drawn",
                            min_value=0.0,
                            value=0.0,
                            step=1.0,
                            key="weekend_tonnes"
                        )

                        if derived_potency > 0:
                            ep = derived_energy / derived_potency
                            et = derived_energy / weekend_tonnes if weekend_tonnes > 0 else None
                            status = get_status(ep, et, ep_threshold, et_threshold)
                            health = cave_health_score(status)
                            action = recommended_action(status)

                            r1, r2, r3, r4 = st.columns(4)
                            r1.metric("Derived E/P", f"{ep:,.2f}")
                            r2.metric("Derived E/T", f"{et:,.2f}" if et is not None else "Waiting")
                            r3.metric("Status", status)
                            r4.metric("Cave Health", f"{health}/100" if health else "Waiting")

                            st.info(action)

                            if et is not None:
                                st.plotly_chart(
                                    plot_matrix(ep, et, load_csv(WEEKEND_DERIVED_FILE), ep_threshold, et_threshold),
                                    width="stretch"
                                )

                            if st.button("Save Weekend Derived Daily Record"):
                                actual_report_date = weekend_actual_date(weekend_start, derived_day)

                                derived_record = {
                                    "Weekend Start": weekend_start,
                                    "Derived Day": derived_day,
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Recommended Action": action,
                                }

                                daily_record = {
                                    "Report Date": actual_report_date,
                                    "Poly": target_poly,
                                    "New Events": derived_events,
                                    "Energy J": derived_energy,
                                    "Potency m3": derived_potency,
                                    "Tonnes": weekend_tonnes,
                                    "E/P": ep,
                                    "E/T": et if et is not None else "",
                                    "Status": status,
                                    "Cave Health Score": health,
                                    "Recommended Action": action,
                                }

                                save_or_update(WEEKEND_DERIVED_FILE, derived_record, ["Weekend Start", "Derived Day", "Poly"])
                                save_or_update(DAILY_FILE, daily_record, ["Report Date", "Poly"])
                                st.success("Weekend derived daily record saved and Daily Draw History updated.")


with tab_weekly:
    st.header("Weekly IMS Dashboard")

    uploaded_weekly_pdf = st.file_uploader("Upload Monday-to-Monday Weekly IMS Summary", type=["pdf"], key="weekly_pdf")

    if uploaded_weekly_pdf:
        text = pdf_to_text(uploaded_weekly_pdf)
        poly_data = parse_poly_sections(text)

        week_range = parse_week_range(text)
        mine_total = parse_mine_total(text)
        largest_ml = parse_largest_magnitude(text)

        st.subheader("Weekly Overview")
        o1, o2, o3 = st.columns(3)
        o1.metric("Reporting Period", week_range or "Not found")
        o2.metric("Mine-wide Events", mine_total)
        o3.metric("Largest ML", largest_ml if largest_ml is not None else "Not found")

        if poly_data:
            weekly_df = pd.DataFrame(poly_data.values())

            st.subheader("Weekly Zone Table")
            st.dataframe(weekly_df, width="stretch")

            weekly_df["Lift Group"] = weekly_df["Poly"].apply(
                lambda x: "Lift II" if str(x).startswith("Lift_II")
                else "Lift I" if str(x).startswith("Lift_I_")
                else "Other"
            )

            lift_split = weekly_df.groupby("Lift Group", as_index=False)["Events"].sum()
            lift_split = lift_split.rename(columns={"Lift Group": "Zone"})

            lift_i = weekly_df[weekly_df["Lift Group"] == "Lift I"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})
            lift_ii = weekly_df[weekly_df["Lift Group"] == "Lift II"][["Poly", "Events"]].rename(columns={"Poly": "Zone"})

            st.plotly_chart(pie_chart(lift_split, "Mine-wide Event Split"), width="stretch")

            c1, c2 = st.columns(2)
            with c1:
                if not lift_i.empty:
                    st.plotly_chart(pie_chart(lift_i, "Lift I Event Distribution"), width="stretch")
            with c2:
                if not lift_ii.empty:
                    st.plotly_chart(pie_chart(lift_ii, "Lift II Event Distribution"), width="stretch")

            ranking = weekly_df.sort_values("Events", ascending=False)[["Poly", "Events", "Energy J", "Potency m3", "Comment"]]
            st.subheader("Hotspot Ranking")
            st.dataframe(ranking, width="stretch")
            st.plotly_chart(bar_chart(ranking.rename(columns={"Poly": "Zone"}), "Weekly Hotspot Ranking"), width="stretch")

            dominant = ranking.iloc[0]["Poly"] if not ranking.empty else "Not found"
            second = ranking.iloc[1]["Poly"] if len(ranking) > 1 else "Not found"

            weekly_comment = (
                f"Weekly seismicity for the period {week_range} recorded {mine_total} mine-wide events. "
                f"The dominant hotspot was {dominant}, followed by {second}. "
                f"The largest recorded event was ML{largest_ml}. Continued monitoring is recommended."
            )

            st.subheader("Weekly Executive Summary")
            st.text_area("Copy-ready weekly summary", value=weekly_comment, height=170)

            st.header("Weekly Lift II East Draw Indicator")

            if "Lift_II_E" in poly_data:
                liie = poly_data["Lift_II_E"]

                w1, w2, w3 = st.columns(3)
                w1.metric("Lift II East Weekly Events", liie["Events"])
                w2.metric("Lift II East Weekly Energy (J)", f"{liie['Energy J']:,.2f}")
                w3.metric("Lift II East Weekly Potency (m³)", f"{liie['Potency m3']:,.3f}")

                weekly_tonnes = st.number_input(
                    "Weekly Lift II East Tonnes Drawn",
                    min_value=0.0,
                    value=0.0,
                    step=1.0,
                    key="weekly_tonnes"
                )

                if liie["Potency m3"] > 0:
                    weekly_ep = liie["Energy J"] / liie["Potency m3"]
                    weekly_et = liie["Energy J"] / weekly_tonnes if weekly_tonnes > 0 else None
                    weekly_status = get_status(weekly_ep, weekly_et, ep_threshold, et_threshold)
                    weekly_health = cave_health_score(weekly_status)
                    weekly_action = recommended_action(weekly_status)

                    q1, q2, q3, q4 = st.columns(4)
                    q1.metric("Weekly E/P", f"{weekly_ep:,.2f}")
                    q2.metric("Weekly E/T", f"{weekly_et:,.2f}" if weekly_et is not None else "Waiting")
                    q3.metric("Weekly Status", weekly_status)
                    q4.metric("Weekly Cave Health", f"{weekly_health}/100" if weekly_health else "Waiting")

                    st.info(weekly_action)

                    if weekly_et is not None:
                        st.plotly_chart(
                            plot_matrix(weekly_ep, weekly_et, load_csv(WEEKLY_FILE), ep_threshold, et_threshold),
                            width="stretch"
                        )

                    weekly_draw_comment = (
                        f"For the weekly period {week_range}, Lift II East recorded {liie['Events']} events, "
                        f"with total energy of {liie['Energy J']:,.0f} J and total potency of {liie['Potency m3']:.3f} m³. "
                        f"The weekly E/P ratio is {weekly_ep:,.2f}. "
                    )

                    if weekly_et is not None:
                        weekly_draw_comment += (
                            f"The weekly E/T ratio is {weekly_et:,.2f}, placing Lift II East in the "
                            f"'{weekly_status}' quadrant. {weekly_action}"
                        )
                    else:
                        weekly_draw_comment += "Weekly E/T cannot be finalised until weekly tonnes are entered."

                    st.text_area("Weekly Lift II East Draw Comment", value=weekly_draw_comment, height=150)

                    if st.button("Save Weekly Summary"):
                        weekly_record = {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Second Hotspot": second,
                            "Lift II East Events": liie["Events"],
                            "Lift II East Energy J": liie["Energy J"],
                            "Lift II East Potency m3": liie["Potency m3"],
                            "Weekly Tonnes": weekly_tonnes,
                            "Weekly E/P": weekly_ep,
                            "Weekly E/T": weekly_et if weekly_et is not None else "",
                            "Weekly Status": weekly_status,
                            "Weekly Cave Health Score": weekly_health,
                            "Weekly Comment": weekly_comment,
                            "Weekly Draw Comment": weekly_draw_comment,
                        }

                        save_or_update(WEEKLY_FILE, weekly_record, ["Week Range"])
                        st.success("Weekly summary saved.")

                    word_file = build_word_report(
                        "GDI Weekly IMS Summary",
                        weekly_comment + "\n\n" + weekly_draw_comment,
                        {
                            "Week Range": week_range,
                            "Mine-wide Events": mine_total,
                            "Largest ML": largest_ml,
                            "Dominant Hotspot": dominant,
                            "Lift II East Weekly E/P": round(weekly_ep, 2),
                            "Lift II East Weekly E/T": round(weekly_et, 2) if weekly_et is not None else "Waiting",
                            "Weekly Status": weekly_status,
                        }
                    )

                    st.download_button(
                        "Download Weekly Word Report",
                        data=word_file,
                        file_name="GDI_Weekly_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Lift_II_E section was not found in the weekly PDF.")


with tab_history:
    st.header("History / Admin")

    daily_history = safe_sort_history(load_csv(DAILY_FILE))
    weekend_history = safe_sort_history(load_csv(WEEKEND_FILE))
    weekend_derived_history = safe_sort_history(load_csv(WEEKEND_DERIVED_FILE))
    weekly_history = safe_sort_history(load_csv(WEEKLY_FILE))

    st.subheader("Daily Draw History")
    if daily_history.empty:
        st.info("No daily records saved yet.")
    else:
        st.dataframe(daily_history, width="stretch")

        selected_daily_row = st.selectbox(
            "Select Daily Record to Delete",
            options=list(range(len(daily_history))),
            format_func=lambda x: f"{daily_history.iloc[x].get('Report Date', '')} - {daily_history.iloc[x].get('Poly', '')}",
            key="delete_daily_row",
        )

        if st.button("Delete Selected Daily Record"):
            daily_history = daily_history.drop(daily_history.index[selected_daily_row])
            save_csv(daily_history, DAILY_FILE)
            st.success("Daily record deleted.")
            st.rerun()

        st.download_button(
            "Download Daily History CSV",
            daily_history.to_csv(index=False).encode("utf-8"),
            "gdi_history.csv",
            "text/csv",
        )

    st.subheader("Weekend Cumulative History")
    if weekend_history.empty:
        st.info("No weekend cumulative records saved yet.")
    else:
        st.dataframe(weekend_history, width="stretch")

        selected_weekend_row = st.selectbox(
            "Select Weekend Cumulative Record to Delete",
            options=list(range(len(weekend_history))),
            format_func=lambda x: f"{weekend_history.iloc[x].get('Weekend Start', '')} - {weekend_history.iloc[x].get('Poly', '')} - {weekend_history.iloc[x].get('Stage', '')}",
            key="delete_weekend_row",
        )

        if st.button("Delete Selected Weekend Cumulative Record"):
            weekend_history = weekend_history.drop(weekend_history.index[selected_weekend_row])
            save_csv(weekend_history, WEEKEND_FILE)
            st.success("Weekend cumulative record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Cumulative CSV",
            weekend_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_cumulative.csv",
            "text/csv",
        )

    st.subheader("Weekend Derived Daily History")
    if weekend_derived_history.empty:
        st.info("No weekend derived records saved yet.")
    else:
        st.dataframe(weekend_derived_history, width="stretch")

        selected_derived_row = st.selectbox(
            "Select Weekend Derived Record to Delete",
            options=list(range(len(weekend_derived_history))),
            format_func=lambda x: f"{weekend_derived_history.iloc[x].get('Report Date', weekend_derived_history.iloc[x].get('Weekend Start', ''))} - {weekend_derived_history.iloc[x].get('Derived Day', '')} - {weekend_derived_history.iloc[x].get('Poly', '')}",
            key="delete_derived_row",
        )

        if st.button("Delete Selected Weekend Derived Record"):
            weekend_derived_history = weekend_derived_history.drop(weekend_derived_history.index[selected_derived_row])
            save_csv(weekend_derived_history, WEEKEND_DERIVED_FILE)
            st.success("Weekend derived record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekend Derived CSV",
            weekend_derived_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekend_derived_daily.csv",
            "text/csv",
        )

    st.subheader("Weekly Summary History")
    if weekly_history.empty:
        st.info("No weekly records saved yet.")
    else:
        st.dataframe(weekly_history, width="stretch")

        selected_weekly_row = st.selectbox(
            "Select Weekly Record to Delete",
            options=list(range(len(weekly_history))),
            format_func=lambda x: f"{weekly_history.iloc[x].get('Week Range', '')}",
            key="delete_weekly_row",
        )

        if st.button("Delete Selected Weekly Record"):
            weekly_history = weekly_history.drop(weekly_history.index[selected_weekly_row])
            save_csv(weekly_history, WEEKLY_FILE)
            st.success("Weekly record deleted.")
            st.rerun()

        st.download_button(
            "Download Weekly CSV",
            weekly_history.to_csv(index=False).encode("utf-8"),
            "gdi_weekly_history.csv",
            "text/csv",
        )
